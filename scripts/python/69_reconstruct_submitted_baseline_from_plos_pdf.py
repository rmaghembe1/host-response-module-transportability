#!/usr/bin/env python3

from __future__ import annotations

import csv
import hashlib
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# ============================================================================
# Locked inputs
# ============================================================================

PLOS_PDF = Path(
    "/mnt/c/Users/REUBEN MAGHEMBE/Downloads/"
    "PONE-D-26-30583.pdf"
)

EXPECTED_PDF_SHA = (
    "00b27a753e2a4b117f8063f9649c6c31"
    "de4b15ec38b53323930a3954c63fa0d3"
)

DONOR_DOCX = Path(
    "/mnt/c/Users/REUBEN MAGHEMBE/Documents/"
    "External transportability of bacterial and viral.docx"
)

EXPECTED_DONOR_SHA = (
    "ec8af7db39d9db5442b633a846232c36"
    "cd5d17f029692f318a88e4c2e413f119"
)

CLEAN_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

EXPECTED_CLEAN_SHA = (
    "691f494f5f2335b1a96f5b8d009c815"
    "d733b3c74fab0f230eaa3f73274f6b38f"
)

SOURCE_MD = Path(
    "docs/"
    "complete_manuscript_draft_v2.3_"
    "submission_candidate_metadata_restored.md"
)

EXPECTED_SOURCE_SHA = (
    "f3b61e6ddb9f5d38c6211c6cfe0d869"
    "4e6ca3b761d52a3245d58df844ab5b2ae"
)


# ============================================================================
# Expected submitted metadata
# ============================================================================

SUBMITTED_TITLE = (
    "Cross-cohort transportability of bacterial- and "
    "viral-associated host-response modules in public "
    "infection transcriptomic datasets"
)

SUBMITTED_SHORT_TITLE = (
    "Transportability of infection host-response modules"
)

AUTHOR_ANCHORS = [
    "Reuben S. Maghembe",
    "AfroBiomics Co. Ltd.",
    "Bridge Street",
    "rmaghembe@gmail.com",
    "rmaghembe@sfuchas.ac.tz",
]

BASELINE_ANCHORS = [
    "GSE211567",
    "GSE73461",
    "DefiniteBacterial",
    "DefiniteViral",
]

REVISION_ONLY_ANCHORS = [
    "GSE72810",
    "29,826",
    "plosone_revision_round1_2026",
]


# ============================================================================
# Output paths
# ============================================================================

WORK = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E14B_submitted_baseline_reconstruction"
)

RAW_TEXT = (
    WORK
    / "PONE-D-26-30583_pages7_20_pdftotext_layout.txt"
)

NORMALIZED_TEXT = (
    WORK
    / "PONE-D-26-30583_submitted_baseline_normalized.txt"
)

LINE_TSV = (
    WORK
    / "PONE-D-26-30583_submitted_baseline_lines.tsv"
)

PARAGRAPH_TSV = (
    WORK
    / "PONE-D-26-30583_submitted_baseline_paragraphs.tsv"
)

BASELINE_DOCX = (
    WORK
    / "PONE-D-26-30583_reconstructed_submitted_baseline.docx"
)

BASELINE_PDF = (
    WORK
    / "render/"
    "PONE-D-26-30583_reconstructed_submitted_baseline.pdf"
)

OUT = Path(
    "results/revision_round1/"
    "plosone_submitted_baseline_reconstruction"
)

QUALITY_GATE = (
    OUT
    / "PLOS_ONE_submitted_baseline_reconstruction_quality_gate.tsv"
)

SUMMARY = (
    OUT
    / "PLOS_ONE_submitted_baseline_reconstruction_summary.tsv"
)

ALIGNMENT = (
    OUT
    / "PLOS_ONE_submitted_vs_clean_section_anchor_audit.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_submitted_baseline_reconstruction_report.md"
)


# ============================================================================
# Utilities
# ============================================================================

def die(message: str) -> None:

    print(
        f"ERROR: {message}",
        file=sys.stderr,
    )

    raise SystemExit(1)


def sha256(path: Path) -> str:

    digest = hashlib.sha256()

    with path.open("rb") as handle:

        for chunk in iter(
            lambda: handle.read(
                1024 * 1024
            ),
            b"",
        ):
            digest.update(chunk)

    return digest.hexdigest()


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:

    if not rows:

        die(
            f"No rows available for {path}"
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=list(
                rows[0].keys()
            ),
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()
        writer.writerows(rows)


def normalize_spaces(
    value: str,
) -> str:

    value = value.replace(
        "\u00a0",
        " ",
    )

    return re.sub(
        r"\s+",
        " ",
        value,
    ).strip()


def strip_pdf_line_number(
    raw: str,
) -> str:

    line = raw.rstrip()

    if not line.strip():
        return ""

    stripped = line.strip()

    # Page-footer numbers.
    if re.fullmatch(
        r"\d{1,2}",
        stripped,
    ):
        return ""

    # Normal PLOS continuous line number separated from text
    # by two or more spaces.
    match = re.match(
        r"^\s*(\d{1,4})\s{2,}(.+?)\s*$",
        line,
    )

    if match:

        number = int(
            match.group(1)
        )

        if number <= 2000:

            return normalize_spaces(
                match.group(2)
            )

    # Fallback for extraction where only one separating space
    # survives. Avoid reference labels such as [12].
    match = re.match(
        r"^\s*(\d{1,4})\s+(.+?)\s*$",
        line,
    )

    if match:

        number = int(
            match.group(1)
        )

        text = match.group(2)

        if (
            number <= 2000
            and not text.startswith("[")
        ):

            return normalize_spaces(
                text
            )

    return normalize_spaces(
        stripped
    )


def flatten(
    text: str,
) -> str:

    return normalize_spaces(
        text
    )


def word_count(
    text: str,
) -> int:

    return len(
        re.findall(
            r"[A-Za-z0-9_]+"
            r"(?:[-'][A-Za-z0-9_]+)*",
            text,
        )
    )


def clear_document_body(
    doc: Document,
) -> None:

    body = doc.element.body

    for child in list(
        body
    ):

        if child.tag.endswith(
            "}sectPr"
        ):
            continue

        body.remove(child)


def format_baseline_document(
    doc: Document,
) -> None:

    section = doc.sections[0]

    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]

    normal.font.name = (
        "Times New Roman"
    )

    normal.font.size = Pt(12)

    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)


def add_paragraph_with_style(
    doc: Document,
    text: str,
    index: int,
) -> None:

    paragraph = doc.add_paragraph()

    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)

    normalized = normalize_spaces(
        text
    )

    heading_tokens = {
        "Abstract",
        "Introduction",
        "Results",
        "Discussion",
        "Methods",
        "References",
        "Supporting information",
        "Supporting Information",
        "Acknowledgments",
        "Author contributions",
        "Funding",
        "Competing interests",
        "Data availability",
        "Ethics approval",
    }

    if (
        index == 1
        or SUBMITTED_TITLE
        in normalized
    ):

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            normalized
        )

        run.bold = True
        run.font.name = (
            "Times New Roman"
        )
        run.font.size = Pt(14)

        return

    if normalized in heading_tokens:

        run = paragraph.add_run(
            normalized
        )

        run.bold = True
        run.font.name = (
            "Times New Roman"
        )
        run.font.size = Pt(12)

        return

    run = paragraph.add_run(
        normalized
    )

    run.font.name = (
        "Times New Roman"
    )
    run.font.size = Pt(12)


def inspect_docx_package(
    path: Path,
) -> dict[str, object]:

    with zipfile.ZipFile(
        path
    ) as archive:

        names = archive.namelist()

        document_xml = archive.read(
            "word/document.xml"
        ).decode(
            "utf-8",
            errors="replace",
        )

        settings_xml = archive.read(
            "word/settings.xml"
        ).decode(
            "utf-8",
            errors="replace",
        )

    return {
        "comments": [
            name
            for name in names
            if (
                "comment"
                in name.lower()
                or "people"
                in name.lower()
                or "person"
                in name.lower()
            )
        ],
        "media": [
            name
            for name in names
            if name.startswith(
                "word/media/"
            )
        ],
        "tracked_insertions": len(
            re.findall(
                r"<w:ins\b",
                document_xml,
            )
        ),
        "tracked_deletions": len(
            re.findall(
                r"<w:del\b",
                document_xml,
            )
        ),
        "track_revisions": (
            "<w:trackRevisions"
            in settings_xml
        ),
    }


# ============================================================================
# Preflight locks
# ============================================================================

WORK.mkdir(
    parents=True,
    exist_ok=True,
)

OUT.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)

for required in [
    PLOS_PDF,
    DONOR_DOCX,
    CLEAN_DOCX,
    SOURCE_MD,
]:

    if not required.exists():

        die(
            f"Required input missing: "
            f"{required}"
        )


pdf_sha = sha256(
    PLOS_PDF
)

donor_sha = sha256(
    DONOR_DOCX
)

clean_sha = sha256(
    CLEAN_DOCX
)

source_sha = sha256(
    SOURCE_MD
)


if pdf_sha != EXPECTED_PDF_SHA:

    die(
        "Authoritative PLOS PDF SHA mismatch."
    )


if donor_sha != EXPECTED_DONOR_SHA:

    die(
        "Formatting donor SHA mismatch."
    )


if clean_sha != EXPECTED_CLEAN_SHA:

    die(
        "Locked clean manuscript SHA mismatch."
    )


if source_sha != EXPECTED_SOURCE_SHA:

    die(
        "Scientific source SHA mismatch."
    )


# ============================================================================
# Fresh authoritative PDF extraction
# ============================================================================

if shutil.which(
    "pdftotext"
) is None:

    die(
        "pdftotext is unavailable."
    )


command = [
    "pdftotext",
    "-layout",
    "-f",
    "7",
    "-l",
    "20",
    str(PLOS_PDF),
    str(RAW_TEXT),
]


result = subprocess.run(
    command,
    stdout=subprocess.PIPE,
    stderr=subprocess.STDOUT,
    text=True,
)


if (
    result.returncode != 0
    or not RAW_TEXT.exists()
    or RAW_TEXT.stat().st_size == 0
):

    die(
        "Fresh extraction of PLOS PDF "
        "pages 7-20 failed."
    )


raw_text = RAW_TEXT.read_text(
    encoding="utf-8",
    errors="replace",
)


# ============================================================================
# Reconstruct numbered manuscript lines
# ============================================================================

line_rows = []

normalized_lines = []

source_page = 7

for raw_line in raw_text.splitlines():

    if "\f" in raw_line:

        chunks = raw_line.split(
            "\f"
        )

    else:

        chunks = [
            raw_line
        ]

    for chunk_index, chunk in enumerate(
        chunks
    ):

        if (
            chunk_index > 0
        ):

            source_page += 1

        cleaned = strip_pdf_line_number(
            chunk
        )

        if not cleaned:
            continue

        normalized_lines.append(
            cleaned
        )

        line_rows.append(
            {
                "baseline_line_index": (
                    len(
                        normalized_lines
                    )
                ),
                "source_pdf_page": (
                    source_page
                ),
                "text": cleaned,
            }
        )


if not normalized_lines:

    die(
        "No normalized manuscript lines "
        "were reconstructed."
    )


write_tsv(
    LINE_TSV,
    line_rows,
)


# ============================================================================
# Paragraph reconstruction
# ============================================================================

paragraphs = []

current = []

source_page = 7


for raw_line in raw_text.splitlines():

    page_break_count = raw_line.count(
        "\f"
    )

    cleaned = strip_pdf_line_number(
        raw_line.replace(
            "\f",
            ""
        )
    )

    if not cleaned:

        if current:

            paragraphs.append(
                normalize_spaces(
                    " ".join(current)
                )
            )

            current = []

    else:

        current.append(
            cleaned
        )

    if page_break_count:

        source_page += page_break_count


if current:

    paragraphs.append(
        normalize_spaces(
            " ".join(current)
        )
    )


paragraphs = [
    paragraph
    for paragraph in paragraphs
    if paragraph
]


if not paragraphs:

    die(
        "No reconstructed paragraphs."
    )


paragraph_rows = [
    {
        "paragraph_index": index,
        "text": paragraph,
    }
    for index, paragraph
    in enumerate(
        paragraphs,
        start=1,
    )
]


write_tsv(
    PARAGRAPH_TSV,
    paragraph_rows,
)


normalized_text = (
    "\n\n".join(
        paragraphs
    )
)


NORMALIZED_TEXT.write_text(
    normalized_text,
    encoding="utf-8",
    newline="\n",
)


flat_baseline = flatten(
    normalized_text
)


# ============================================================================
# Metadata/anchor audit
# ============================================================================

checks = []


def check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": (
                description
            ),
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected
            ),
        }
    )


check(
    "Authoritative PDF SHA locked",
    pdf_sha == EXPECTED_PDF_SHA,
    pdf_sha,
    EXPECTED_PDF_SHA,
)

check(
    "Formatting donor SHA locked",
    donor_sha == EXPECTED_DONOR_SHA,
    donor_sha,
    EXPECTED_DONOR_SHA,
)

check(
    "Clean manuscript SHA locked",
    clean_sha == EXPECTED_CLEAN_SHA,
    clean_sha,
    EXPECTED_CLEAN_SHA,
)

check(
    "Scientific source SHA locked",
    source_sha == EXPECTED_SOURCE_SHA,
    source_sha,
    EXPECTED_SOURCE_SHA,
)


title_flat = flatten(
    SUBMITTED_TITLE
)

short_title_flat = flatten(
    SUBMITTED_SHORT_TITLE
)


check(
    "Submitted title recovered after whitespace normalization",
    title_flat in flat_baseline,
    (
        "present"
        if title_flat
        in flat_baseline
        else "absent"
    ),
    "present",
)


check(
    "Submitted short title recovered",
    short_title_flat
    in flat_baseline,
    (
        "present"
        if short_title_flat
        in flat_baseline
        else "absent"
    ),
    "present",
)


for anchor in AUTHOR_ANCHORS:

    check(
        f"Submitted author metadata recovered: {anchor}",
        anchor in flat_baseline,
        (
            "present"
            if anchor
            in flat_baseline
            else "absent"
        ),
        "present",
    )


for anchor in BASELINE_ANCHORS:

    check(
        f"Submitted scientific anchor recovered: {anchor}",
        anchor in flat_baseline,
        (
            "present"
            if anchor
            in flat_baseline
            else "absent"
        ),
        "present",
    )


for anchor in REVISION_ONLY_ANCHORS:

    check(
        f"Revision-only anchor absent from submitted baseline: {anchor}",
        anchor not in flat_baseline,
        (
            flat_baseline.count(
                anchor
            )
        ),
        0,
    )


# ============================================================================
# Build baseline DOCX using donor styles only
# ============================================================================

baseline_doc = Document(
    DONOR_DOCX
)

clear_document_body(
    baseline_doc
)

format_baseline_document(
    baseline_doc
)


for index, paragraph_text in enumerate(
    paragraphs,
    start=1,
):

    add_paragraph_with_style(
        baseline_doc,
        paragraph_text,
        index,
    )


baseline_doc.save(
    BASELINE_DOCX
)


if not BASELINE_DOCX.exists():

    die(
        "Baseline DOCX was not created."
    )


baseline_docx_sha = sha256(
    BASELINE_DOCX
)

reopened = Document(
    BASELINE_DOCX
)


reopened_text = "\n\n".join(
    paragraph.text
    for paragraph
    in reopened.paragraphs
    if paragraph.text.strip()
)


flat_reopened = flatten(
    reopened_text
)


check(
    "Reconstructed DOCX preserves normalized baseline text",
    flat_reopened
    == flat_baseline,
    (
        "identical"
        if flat_reopened
        == flat_baseline
        else "different"
    ),
    "identical",
)


package = inspect_docx_package(
    BASELINE_DOCX
)


check(
    "Baseline reconstruction contains no tracked insertions",
    package[
        "tracked_insertions"
    ] == 0,
    package[
        "tracked_insertions"
    ],
    0,
)


check(
    "Baseline reconstruction contains no tracked deletions",
    package[
        "tracked_deletions"
    ] == 0,
    package[
        "tracked_deletions"
    ],
    0,
)


check(
    "Baseline reconstruction does not enable Track Changes",
    not package[
        "track_revisions"
    ],
    package[
        "track_revisions"
    ],
    False,
)


check(
    "Baseline reconstruction contains no comments",
    len(
        package[
            "comments"
        ]
    ) == 0,
    (
        "|".join(
            package[
                "comments"
            ]
        )
        or "NONE"
    ),
    "NONE",
)


check(
    "Baseline reconstruction contains no embedded figures",
    len(
        package[
            "media"
        ]
    ) == 0,
    len(
        package[
            "media"
        ]
    ),
    0,
)


# ============================================================================
# Clean-document anchor map
# ============================================================================

clean_doc = Document(
    CLEAN_DOCX
)


clean_paragraph_text = [
    normalize_spaces(
        paragraph.text
    )
    for paragraph
    in clean_doc.paragraphs
    if paragraph.text.strip()
]


clean_flat = flatten(
    "\n".join(
        clean_paragraph_text
    )
)


anchor_names = [
    "Abstract",
    "Introduction",
    "Methods",
    "Results",
    "Discussion",
    "References",
    "GSE211567",
    "GSE73461",
    "GSE72810",
    "GSVA",
    "29,826",
    "ChatGPT",
]


alignment_rows = []


for anchor in anchor_names:

    baseline_hits = [
        index
        for index, paragraph
        in enumerate(
            paragraphs,
            start=1,
        )
        if anchor in paragraph
    ]

    clean_hits = [
        index
        for index, paragraph
        in enumerate(
            clean_paragraph_text,
            start=1,
        )
        if anchor in paragraph
    ]

    alignment_rows.append(
        {
            "anchor": anchor,
            "submitted_paragraph_hits": (
                ",".join(
                    map(
                        str,
                        baseline_hits,
                    )
                )
                or "NONE"
            ),
            "clean_paragraph_hits": (
                ",".join(
                    map(
                        str,
                        clean_hits,
                    )
                )
                or "NONE"
            ),
            "submitted_count": len(
                baseline_hits
            ),
            "clean_count": len(
                clean_hits
            ),
        }
    )


write_tsv(
    ALIGNMENT,
    alignment_rows,
)


# ============================================================================
# Smoke render
# ============================================================================

render_dir = (
    WORK
    / "render"
)

render_dir.mkdir(
    parents=True,
    exist_ok=True,
)


if shutil.which(
    "soffice"
) is None:

    die(
        "LibreOffice soffice unavailable."
    )


profile_dir = (
    WORK
    / "libreoffice_profile"
)

profile_dir.mkdir(
    parents=True,
    exist_ok=True,
)


for old_pdf in render_dir.glob(
    "*.pdf"
):

    old_pdf.unlink()


result = subprocess.run(
    [
        "soffice",
        "--headless",
        (
            "-env:UserInstallation="
            f"file://{profile_dir.resolve()}"
        ),
        "--convert-to",
        "pdf",
        "--outdir",
        str(
            render_dir.resolve()
        ),
        str(
            BASELINE_DOCX.resolve()
        ),
    ],
    stdout=subprocess.PIPE,
    stderr=subprocess.STDOUT,
    text=True,
)


expected_render = (
    render_dir
    / (
        BASELINE_DOCX.stem
        + ".pdf"
    )
)


smoke_ok = (
    expected_render.exists()
    and expected_render.stat().st_size
    > 0
)


smoke_bytes = (
    expected_render.stat().st_size
    if expected_render.exists()
    else 0
)


check(
    "Baseline reconstruction smoke-renders in LibreOffice",
    smoke_ok,
    (
        f"status={result.returncode}; "
        f"bytes={smoke_bytes}"
    ),
    "non-empty PDF",
)


# ============================================================================
# Final gate
# ============================================================================

passed = sum(
    row[
        "pass"
    ] == "TRUE"
    for row
    in checks
)


failed = (
    len(checks)
    - passed
)


quality_gate = (
    "PASS"
    if failed == 0
    else "FAIL"
)


final_status = (
    "READY_FOR_CONTROLLED_OOXML_REDLINE_BUILD"
    if failed == 0
    else "SUBMITTED_BASELINE_RECONSTRUCTION_REQUIRES_REVIEW"
)


write_tsv(
    QUALITY_GATE,
    checks,
)


write_tsv(
    SUMMARY,
    [
        {
            "plos_pdf_sha256": (
                pdf_sha
            ),
            "donor_docx_sha256": (
                donor_sha
            ),
            "clean_docx_sha256": (
                clean_sha
            ),
            "scientific_source_sha256": (
                source_sha
            ),
            "raw_extraction_pages": (
                "7-20"
            ),
            "normalized_lines": len(
                normalized_lines
            ),
            "reconstructed_paragraphs": len(
                paragraphs
            ),
            "baseline_word_count": (
                word_count(
                    normalized_text
                )
            ),
            "baseline_docx_sha256": (
                baseline_docx_sha
            ),
            "baseline_docx_paragraphs": len(
                reopened.paragraphs
            ),
            "smoke_render_pdf_bytes": (
                smoke_bytes
            ),
            "quality_checks": len(
                checks
            ),
            "quality_checks_passed": (
                passed
            ),
            "quality_checks_failed": (
                failed
            ),
            "quality_gate": (
                quality_gate
            ),
            "final_status": (
                final_status
            ),
        }
    ],
)


REPORT.write_text(
    "\n".join(
        [
            "# PLOS ONE Submitted-Baseline Reconstruction",
            "",
            "Manuscript: PONE-D-26-30583",
            "",
            "## Authority",
            "",
            (
                "- Submitted-text authority: SHA-verified "
                "PLOS submission PDF."
            ),
            (
                "- June 8 DOCX used only as a style/formatting "
                "donor; its text is not treated as authoritative."
            ),
            "",
            "## Reconstruction",
            "",
            (
                "- PLOS PDF pages 7-20 were freshly extracted "
                "with pdftotext -layout."
            ),
            (
                "- Continuous manuscript line numbers and "
                "page-footer numbers were removed."
            ),
            (
                "- A whitespace-normalized text baseline, "
                "line inventory and paragraph inventory were created."
            ),
            (
                "- A baseline DOCX was reconstructed using donor "
                "styles while preserving the normalized submitted text."
            ),
            "",
            "## Locked revised manuscript",
            "",
            (
                f"- Clean DOCX SHA256: `{clean_sha}`"
            ),
            (
                f"- Scientific source SHA256: `{source_sha}`"
            ),
            "",
            "## Output",
            "",
            (
                f"- Baseline DOCX SHA256: "
                f"`{baseline_docx_sha}`"
            ),
            (
                f"- Reconstructed paragraphs: "
                f"{len(paragraphs)}"
            ),
            (
                f"- Baseline word count: "
                f"{word_count(normalized_text)}"
            ),
            "",
            "## Quality gate",
            "",
            (
                f"- Checks passed: "
                f"{passed}/{len(checks)}"
            ),
            (
                f"- Quality gate: `{quality_gate}`"
            ),
            (
                f"- Final status: `{final_status}`"
            ),
            "",
        ]
    ),
    encoding="utf-8",
    newline="\n",
)


# ============================================================================
# Console
# ============================================================================

print(
    "===== PLOS ONE SUBMITTED BASELINE RECONSTRUCTION ====="
)

print(
    f"plos_pdf_sha256\t{pdf_sha}"
)

print(
    f"donor_docx_sha256\t{donor_sha}"
)

print(
    f"clean_docx_sha256\t{clean_sha}"
)

print(
    f"scientific_source_sha256\t{source_sha}"
)

print(
    f"normalized_lines\t{len(normalized_lines)}"
)

print(
    f"reconstructed_paragraphs\t{len(paragraphs)}"
)

print(
    f"baseline_word_count\t{word_count(normalized_text)}"
)

print(
    f"baseline_docx_sha256\t{baseline_docx_sha}"
)

print(
    f"baseline_docx_paragraphs\t{len(reopened.paragraphs)}"
)

print(
    f"smoke_render_pdf_bytes\t{smoke_bytes}"
)

print(
    f"quality_checks_passed\t{passed}/{len(checks)}"
)

print(
    f"quality_gate\t{quality_gate}"
)

print(
    f"final_status\t{final_status}"
)

print(
    f"baseline_docx\t{BASELINE_DOCX}"
)

print(
    f"line_inventory\t{LINE_TSV}"
)

print(
    f"paragraph_inventory\t{PARAGRAPH_TSV}"
)

print(
    f"anchor_audit\t{ALIGNMENT}"
)

print(
    f"summary\t{SUMMARY}"
)

print(
    f"quality_gate_file\t{QUALITY_GATE}"
)

print(
    f"report\t{REPORT}"
)


if failed:

    die(
        "Submitted-baseline reconstruction failed "
        f"{failed} quality check(s)."
    )
