#!/usr/bin/env python3
from __future__ import annotations

import copy
import csv
import hashlib
import os
import re
import shutil
import subprocess
import zipfile
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NS = {"w": W_NS}


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


CLEAN_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

EXPECTED_CLEAN_SHA = (
    "691f494f5f2335b1a96f5b8d009c815"
    "d733b3c74fab0f230eaa3f73274f6b38f"
)

BASELINE_DOCX = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E14B_submitted_baseline_reconstruction/"
    "PONE-D-26-30583_reconstructed_submitted_baseline.docx"
)

EXPECTED_BASELINE_SHA = (
    "0e0b4a2bcec736d2735ecc903ad72e99"
    "a1a3897bc09e2f991a7249000852dd34"
)

SOURCE_MD = Path(
    "docs/"
    "complete_manuscript_draft_v2.3_"
    "submission_candidate_metadata_restored.md"
)

EXPECTED_SOURCE_SHA = (
    "f3b61e6ddb9f5d38c6211c6cfe0d869"
    "4e6ca3b761d52a3245d58df844ab5b2ae"
)

MARKED_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_marked_up_revised_manuscript.docx"
)

EXPECTED_PRIOR_MARKED_SHA = (
    "9ec494b73f6cefa344333e17a6335a6e"
    "59a9d9710f91409f137ae234527ecc11"
)

OLD_MANIFEST = Path(
    "results/revision_round1/"
    "plosone_targeted_ooxml_redline_v2.3/"
    "PLOS_ONE_targeted_redline_manifest.tsv"
)

SUBMITTED_TITLE = (
    "Cross-cohort transportability of bacterial- and "
    "viral-associated host-response modules in public "
    "infection transcriptomic datasets"
)

SUBMITTED_SHORT_TITLE = (
    "Short title: Transportability of infection host-response modules"
)

REVISION_AUTHOR = "Reuben S. Maghembe"
REVISION_DATE = "2026-08-08T06:55:00Z"


WORK = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E14F_readable_redline_repair"
)

CANDIDATE_DOCX = (
    WORK
    / "PONE-D-26-30583_marked_up_readable_candidate.docx"
)

ACCEPTED_QA_DOCX = (
    WORK
    / "PONE-D-26-30583_marked_up_readable_acceptance_QA.docx"
)

SMOKE_DIR = (
    WORK
    / "libreoffice_smoke_render"
)

PROFILE_DIR = (
    WORK
    / "libreoffice_profile"
)

OUT = Path(
    "results/revision_round1/"
    "plosone_targeted_ooxml_redline_readable_v2.3"
)

MANIFEST = (
    OUT
    / "PLOS_ONE_readable_redline_manifest.tsv"
)

QUALITY_GATE = (
    OUT
    / "PLOS_ONE_readable_redline_quality_gate.tsv"
)

SUMMARY = (
    OUT
    / "PLOS_ONE_readable_redline_summary.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_readable_ooxml_redline_repair_report.md"
)


def fail(message: str) -> None:
    raise RuntimeError(message)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()

    with path.open("rb") as handle:
        for chunk in iter(
            lambda: handle.read(1024 * 1024),
            b"",
        ):
            digest.update(chunk)

    return digest.hexdigest()


def read_tsv(
    path: Path,
) -> list[dict[str, str]]:

    with path.open(
        "r",
        encoding="utf-8",
        newline="",
    ) as handle:

        return list(
            csv.DictReader(
                handle,
                delimiter="\t",
            )
        )


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:

    if not rows:
        fail(
            f"No rows available for {path}"
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=list(
                rows[0].keys()
            ),
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()
        writer.writerows(rows)


def xml_bytes(
    root: etree._Element,
) -> bytes:

    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def write_docx_with_overrides(
    source_docx: Path,
    output_docx: Path,
    overrides: dict[str, bytes],
) -> None:

    output_docx.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with zipfile.ZipFile(
        source_docx,
        "r",
    ) as zin:

        with zipfile.ZipFile(
            output_docx,
            "w",
            zipfile.ZIP_DEFLATED,
        ) as zout:

            seen = set()

            for info in zin.infolist():

                name = info.filename

                payload = overrides.get(
                    name,
                    zin.read(name),
                )

                zout.writestr(
                    name,
                    payload,
                )

                seen.add(name)

            for name, payload in overrides.items():

                if name not in seen:
                    zout.writestr(
                        name,
                        payload,
                    )


def enable_track_revisions(
    settings_root: etree._Element,
) -> None:

    if settings_root.find(
        "w:trackRevisions",
        namespaces=NS,
    ) is None:

        settings_root.insert(
            0,
            etree.Element(
                w("trackRevisions")
            ),
        )


def next_change_id(
    doc_root: etree._Element,
) -> int:

    ids = []

    for element in doc_root.xpath(
        ".//*[@w:id]",
        namespaces=NS,
    ):

        try:
            ids.append(
                int(
                    element.get(
                        w("id")
                    )
                )
            )

        except (TypeError, ValueError):
            pass

    return (
        max(ids) + 1
        if ids
        else 1
    )


def revision_wrapper(
    change_id: int,
) -> etree._Element:

    element = etree.Element(
        w("ins")
    )

    element.set(
        w("id"),
        str(change_id),
    )

    element.set(
        w("author"),
        REVISION_AUTHOR,
    )

    element.set(
        w("date"),
        REVISION_DATE,
    )

    return element


def first_run_properties(
    paragraph: etree._Element,
) -> etree._Element | None:

    result = paragraph.xpath(
        ".//w:r/w:rPr",
        namespaces=NS,
    )

    if not result:
        return None

    return copy.deepcopy(
        result[0]
    )


def make_run(
    text: str,
    run_properties: etree._Element | None,
) -> etree._Element:

    run = etree.Element(
        w("r")
    )

    if run_properties is not None:
        run.append(
            copy.deepcopy(
                run_properties
            )
        )

    text_node = etree.SubElement(
        run,
        w("t"),
    )

    if (
        text.startswith(" ")
        or text.endswith(" ")
        or "  " in text
    ):
        text_node.set(
            f"{{{XML_NS}}}space",
            "preserve",
        )

    text_node.text = text

    return run


def tokenize_with_spacing(
    value: str,
) -> list[str]:

    if not value:
        return []

    leading = re.match(
        r"^\s+",
        value,
    )

    chunks = re.findall(
        r"\S+\s*",
        value,
    )

    if leading and chunks:
        chunks[0] = (
            leading.group(0)
            + chunks[0]
        )

    elif leading and not chunks:
        chunks = [
            leading.group(0)
        ]

    return chunks


def insertion_only_diff(
    paragraph: etree._Element,
    old_text: str,
    new_text: str,
    change_id: int,
) -> tuple[int, int]:

    run_properties = (
        first_run_properties(
            paragraph
        )
    )

    paragraph_properties = (
        paragraph.find(
            "w:pPr",
            namespaces=NS,
        )
    )

    for child in list(
        paragraph
    ):

        if child is paragraph_properties:
            continue

        paragraph.remove(child)

    old_tokens = (
        tokenize_with_spacing(
            old_text
        )
    )

    new_tokens = (
        tokenize_with_spacing(
            new_text
        )
    )

    matcher = SequenceMatcher(
        None,
        old_tokens,
        new_tokens,
        autojunk=False,
    )

    insertion_count = 0

    for (
        tag,
        i1,
        i2,
        j1,
        j2,
    ) in matcher.get_opcodes():

        if tag == "equal":

            text = "".join(
                new_tokens[j1:j2]
            )

            if text:
                paragraph.append(
                    make_run(
                        text,
                        run_properties,
                    )
                )

        elif tag in {
            "insert",
            "replace",
        }:

            text = "".join(
                new_tokens[j1:j2]
            )

            if text:

                wrapper = (
                    revision_wrapper(
                        change_id
                    )
                )

                change_id += 1

                wrapper.append(
                    make_run(
                        text,
                        run_properties,
                    )
                )

                paragraph.append(
                    wrapper
                )

                insertion_count += 1

        elif tag == "delete":

            # Deleted submitted-baseline text is omitted from
            # the visible manuscript. The response document
            # records the original wording.
            continue

    visible_text = "".join(
        paragraph.xpath(
            ".//w:t/text()",
            namespaces=NS,
        )
    )

    if visible_text != new_text:

        fail(
            "Insertion-only diff failed to "
            "preserve clean paragraph text."
        )

    return (
        change_id,
        insertion_count,
    )


def wrap_runs_as_insertions(
    container: etree._Element,
    change_id: int,
) -> tuple[int, int]:

    count = 0

    for child in list(
        container
    ):

        if child.tag == w("pPr"):
            continue

        if child.tag == w("r"):

            wrapper = (
                revision_wrapper(
                    change_id
                )
            )

            change_id += 1

            index = (
                container.index(
                    child
                )
            )

            container.remove(
                child
            )

            wrapper.append(
                child
            )

            container.insert(
                index,
                wrapper,
            )

            count += 1
            continue

        (
            change_id,
            nested_count,
        ) = (
            wrap_runs_as_insertions(
                child,
                change_id,
            )
        )

        count += nested_count

    return (
        change_id,
        count,
    )


def unwrap_element(
    element: etree._Element,
) -> None:

    parent = element.getparent()

    if parent is None:
        return

    index = parent.index(
        element
    )

    for child in list(
        element
    ):

        element.remove(child)

        parent.insert(
            index,
            child,
        )

        index += 1

    parent.remove(element)


def accept_insertions(
    doc_root: etree._Element,
) -> etree._Element:

    root = copy.deepcopy(
        doc_root
    )

    for element in list(
        root.xpath(
            ".//w:ins",
            namespaces=NS,
        )
    ):

        unwrap_element(
            element
        )

    return root


def settings_without_track(
    settings_root: etree._Element,
) -> etree._Element:

    root = copy.deepcopy(
        settings_root
    )

    node = root.find(
        "w:trackRevisions",
        namespaces=NS,
    )

    if node is not None:
        root.remove(node)

    return root


def docx_body_signature(
    doc: Document,
) -> tuple[
    list[str],
    list[list[list[str]]],
]:

    paragraphs = [
        paragraph.text
        for paragraph
        in doc.paragraphs
    ]

    tables = [
        [
            [
                cell.text
                for cell in row.cells
            ]
            for row in table.rows
        ]
        for table in doc.tables
    ]

    return (
        paragraphs,
        tables,
    )


def section_orientations(
    doc: Document,
) -> list[str]:

    return [
        (
            "landscape"
            if section.page_width
            > section.page_height
            else "portrait"
        )
        for section in doc.sections
    ]


WORK.mkdir(
    parents=True,
    exist_ok=True,
)

OUT.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)


for required in [
    CLEAN_DOCX,
    BASELINE_DOCX,
    SOURCE_MD,
    MARKED_DOCX,
    OLD_MANIFEST,
]:

    if not required.exists():
        fail(
            f"Required input missing: {required}"
        )


clean_sha = sha256(
    CLEAN_DOCX
)

baseline_sha = sha256(
    BASELINE_DOCX
)

source_sha = sha256(
    SOURCE_MD
)

prior_marked_sha = sha256(
    MARKED_DOCX
)


if clean_sha != EXPECTED_CLEAN_SHA:
    fail(
        f"Clean DOCX SHA mismatch: {clean_sha}"
    )

if baseline_sha != EXPECTED_BASELINE_SHA:
    fail(
        f"Baseline DOCX SHA mismatch: {baseline_sha}"
    )

if source_sha != EXPECTED_SOURCE_SHA:
    fail(
        f"Scientific source SHA mismatch: {source_sha}"
    )

if (
    prior_marked_sha
    != EXPECTED_PRIOR_MARKED_SHA
):
    fail(
        "Prior marked-up DOCX SHA mismatch: "
        f"{prior_marked_sha}"
    )


old_manifest = read_tsv(
    OLD_MANIFEST
)

selected_indices = sorted(
    int(
        row[
            "clean_paragraph_index"
        ]
    )
    for row in old_manifest
)

manifest_by_index = {
    int(
        row[
            "clean_paragraph_index"
        ]
    ): row
    for row in old_manifest
}


clean_doc = Document(
    CLEAN_DOCX
)

baseline_doc = Document(
    BASELINE_DOCX
)

clean_paragraphs = list(
    clean_doc.paragraphs
)

baseline_paragraphs = list(
    baseline_doc.paragraphs
)

clean_signature = (
    docx_body_signature(
        clean_doc
    )
)

clean_orientations = (
    section_orientations(
        clean_doc
    )
)


with zipfile.ZipFile(
    CLEAN_DOCX,
    "r",
) as zin:

    doc_root = etree.fromstring(
        zin.read(
            "word/document.xml"
        )
    )

    settings_root = etree.fromstring(
        zin.read(
            "word/settings.xml"
        )
    )


body_paragraphs = (
    doc_root.xpath(
        "/w:document/w:body/w:p",
        namespaces=NS,
    )
)


if (
    len(body_paragraphs)
    != len(clean_paragraphs)
):
    fail(
        "Top-level paragraph count mismatch."
    )


enable_track_revisions(
    settings_root
)

change_id = next_change_id(
    doc_root
)

manifest_rows = []


for index in selected_indices:

    paragraph = (
        body_paragraphs[
            index - 1
        ]
    )

    clean_text = (
        clean_paragraphs[
            index - 1
        ].text
    )

    source_row = (
        manifest_by_index[
            index
        ]
    )

    prior_action = (
        source_row.get(
            "redline_action",
            "",
        )
    )

    best_index_text = (
        source_row.get(
            "best_submitted_paragraph_index",
            "NONE",
        )
    )

    best_index = (
        int(best_index_text)
        if best_index_text.isdigit()
        else 0
    )


    if index == 1:

        old_text = (
            SUBMITTED_TITLE
        )

    elif index == 2:

        old_text = (
            SUBMITTED_SHORT_TITLE
        )

    elif (
        1 <= best_index
        <= len(
            baseline_paragraphs
        )
    ):

        old_text = (
            baseline_paragraphs[
                best_index - 1
            ].text
        )

    else:

        old_text = ""


    if (
        index in {
            1,
            2,
        }
        or prior_action
        == "TRACKED_WORD_LEVEL_REVISION"
    ):

        (
            change_id,
            insertion_count,
        ) = insertion_only_diff(
            paragraph,
            old_text,
            clean_text,
            change_id,
        )

        new_action = (
            "READABLE_INSERTION_ONLY_DIFF"
        )

    else:

        (
            change_id,
            insertion_count,
        ) = wrap_runs_as_insertions(
            paragraph,
            change_id,
        )

        new_action = (
            "TRACKED_INSERT_PARAGRAPH"
        )


    manifest_rows.append(
        {
            "clean_paragraph_index": index,
            "response_ids": (
                source_row.get(
                    "response_ids",
                    "",
                )
            ),
            "previous_redline_action": (
                prior_action
            ),
            "readable_redline_action": (
                new_action
            ),
            "tracked_insert_wrappers": (
                insertion_count
            ),
            "tracked_delete_wrappers": 0,
            "clean_excerpt": (
                clean_text[:180]
            ),
        }
    )


write_docx_with_overrides(
    CLEAN_DOCX,
    CANDIDATE_DOCX,
    {
        "word/document.xml": (
            xml_bytes(
                doc_root
            )
        ),
        "word/settings.xml": (
            xml_bytes(
                settings_root
            )
        ),
    },
)


candidate_sha = sha256(
    CANDIDATE_DOCX
)


accepted_root = (
    accept_insertions(
        doc_root
    )
)

accepted_settings = (
    settings_without_track(
        settings_root
    )
)


write_docx_with_overrides(
    CLEAN_DOCX,
    ACCEPTED_QA_DOCX,
    {
        "word/document.xml": (
            xml_bytes(
                accepted_root
            )
        ),
        "word/settings.xml": (
            xml_bytes(
                accepted_settings
            )
        ),
    },
)


accepted_doc = Document(
    ACCEPTED_QA_DOCX
)

accepted_signature = (
    docx_body_signature(
        accepted_doc
    )
)

accepted_orientations = (
    section_orientations(
        accepted_doc
    )
)


with zipfile.ZipFile(
    CANDIDATE_DOCX,
    "r",
) as zin:

    candidate_root = etree.fromstring(
        zin.read(
            "word/document.xml"
        )
    )

    candidate_settings = etree.fromstring(
        zin.read(
            "word/settings.xml"
        )
    )

    package_names = zin.namelist()


tracked_insertions = len(
    candidate_root.xpath(
        ".//w:ins",
        namespaces=NS,
    )
)

tracked_deletions = len(
    candidate_root.xpath(
        ".//w:del",
        namespaces=NS,
    )
)

track_enabled = (
    candidate_settings.find(
        "w:trackRevisions",
        namespaces=NS,
    )
    is not None
)

authors = sorted(
    {
        node.get(
            w("author"),
            "",
        )
        for node
        in candidate_root.xpath(
            ".//w:ins",
            namespaces=NS,
        )
    }
)

comment_parts = [
    name
    for name in package_names
    if (
        "comment" in name.lower()
        or "people" in name.lower()
        or "person" in name.lower()
    )
]

media_parts = [
    name
    for name in package_names
    if name.startswith(
        "word/media/"
    )
]


SMOKE_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

PROFILE_DIR.mkdir(
    parents=True,
    exist_ok=True,
)


for old_pdf in SMOKE_DIR.glob(
    "*.pdf"
):
    old_pdf.unlink()


if shutil.which(
    "soffice"
) is None:
    fail(
        "LibreOffice soffice not found."
    )


environment = os.environ.copy()

environment[
    "HOME"
] = str(
    PROFILE_DIR.resolve()
)


render_result = subprocess.run(
    [
        "soffice",
        "--headless",
        (
            "-env:UserInstallation="
            f"file://{PROFILE_DIR.resolve()}"
        ),
        "--convert-to",
        "pdf",
        "--outdir",
        str(
            SMOKE_DIR.resolve()
        ),
        str(
            CANDIDATE_DOCX.resolve()
        ),
    ],
    stdout=subprocess.PIPE,
    stderr=subprocess.STDOUT,
    text=True,
    env=environment,
)


smoke_pdf = (
    SMOKE_DIR
    / f"{CANDIDATE_DOCX.stem}.pdf"
)

smoke_ok = (
    smoke_pdf.exists()
    and smoke_pdf.stat().st_size
    > 0
)

smoke_bytes = (
    smoke_pdf.stat().st_size
    if smoke_pdf.exists()
    else 0
)


checks = []


def check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": (
                description
            ),
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected
            ),
        }
    )


check(
    "Clean DOCX SHA locked",
    clean_sha == EXPECTED_CLEAN_SHA,
    clean_sha,
    EXPECTED_CLEAN_SHA,
)

check(
    "Baseline DOCX SHA locked",
    baseline_sha == EXPECTED_BASELINE_SHA,
    baseline_sha,
    EXPECTED_BASELINE_SHA,
)

check(
    "Scientific source SHA locked",
    source_sha == EXPECTED_SOURCE_SHA,
    source_sha,
    EXPECTED_SOURCE_SHA,
)

check(
    "Prior marked-up candidate SHA locked",
    prior_marked_sha
    == EXPECTED_PRIOR_MARKED_SHA,
    prior_marked_sha,
    EXPECTED_PRIOR_MARKED_SHA,
)

check(
    "Readable candidate created",
    CANDIDATE_DOCX.exists(),
    CANDIDATE_DOCX.exists(),
    True,
)

check(
    "Readable candidate differs from prior marked-up candidate",
    candidate_sha
    != prior_marked_sha,
    candidate_sha,
    "different",
)

check(
    "Track Changes enabled",
    track_enabled,
    track_enabled,
    True,
)

check(
    "Tracked insertions exist",
    tracked_insertions > 0,
    tracked_insertions,
    ">0",
)

check(
    "No tracked deletions remain",
    tracked_deletions == 0,
    tracked_deletions,
    0,
)

check(
    "Revision author correct",
    authors == [
        REVISION_AUTHOR
    ],
    "|".join(authors),
    REVISION_AUTHOR,
)

check(
    "Same 71 revision-target paragraphs retained",
    len(manifest_rows) == 71,
    len(manifest_rows),
    71,
)

check(
    "Acceptance QA reproduces clean paragraph/table content",
    accepted_signature
    == clean_signature,
    (
        "identical"
        if accepted_signature
        == clean_signature
        else "different"
    ),
    "identical",
)

check(
    "Acceptance QA preserves section orientations",
    accepted_orientations
    == clean_orientations,
    "|".join(
        accepted_orientations
    ),
    "|".join(
        clean_orientations
    ),
)

check(
    "Acceptance QA retains two editable tables",
    len(
        accepted_doc.tables
    ) == 2,
    len(
        accepted_doc.tables
    ),
    2,
)

check(
    "Acceptance QA retains five sections",
    len(
        accepted_doc.sections
    ) == 5,
    len(
        accepted_doc.sections
    ),
    5,
)

check(
    "No comment-related package parts",
    len(
        comment_parts
    ) == 0,
    (
        "|".join(
            comment_parts
        )
        or "NONE"
    ),
    "NONE",
)

check(
    "No embedded manuscript figures",
    len(
        media_parts
    ) == 0,
    len(
        media_parts
    ),
    0,
)

check(
    "LibreOffice smoke render succeeded",
    smoke_ok,
    (
        f"status={render_result.returncode}; "
        f"bytes={smoke_bytes}"
    ),
    "non-empty PDF",
)

clean_sha_after = sha256(
    CLEAN_DOCX
)

check(
    "Canonical clean DOCX remains immutable",
    clean_sha_after
    == EXPECTED_CLEAN_SHA,
    clean_sha_after,
    EXPECTED_CLEAN_SHA,
)


passed = sum(
    row[
        "pass"
    ] == "TRUE"
    for row in checks
)

failed = (
    len(checks)
    - passed
)

quality_gate = (
    "PASS"
    if failed == 0
    else "FAIL"
)

final_status = (
    "READY_FOR_READABLE_MARKED_UP_VISUAL_QA"
    if failed == 0
    else "READABLE_REDLINE_REPAIR_REQUIRES_CORRECTION"
)


write_tsv(
    MANIFEST,
    manifest_rows,
)

write_tsv(
    QUALITY_GATE,
    checks,
)

write_tsv(
    SUMMARY,
    [
        {
            "clean_docx_sha256": (
                clean_sha
            ),
            "baseline_docx_sha256": (
                baseline_sha
            ),
            "scientific_source_sha256": (
                source_sha
            ),
            "prior_marked_up_sha256": (
                prior_marked_sha
            ),
            "candidate_docx_sha256": (
                candidate_sha
            ),
            "selected_paragraphs": (
                len(
                    manifest_rows
                )
            ),
            "tracked_insertions": (
                tracked_insertions
            ),
            "tracked_deletions": (
                tracked_deletions
            ),
            "insertion_only_diff_paragraphs": sum(
                row[
                    "readable_redline_action"
                ]
                == "READABLE_INSERTION_ONLY_DIFF"
                for row in manifest_rows
            ),
            "whole_paragraph_insertions": sum(
                row[
                    "readable_redline_action"
                ]
                == "TRACKED_INSERT_PARAGRAPH"
                for row in manifest_rows
            ),
            "accepted_content_matches_clean": (
                accepted_signature
                == clean_signature
            ),
            "smoke_render_pdf_bytes": (
                smoke_bytes
            ),
            "quality_checks": (
                len(checks)
            ),
            "quality_checks_passed": (
                passed
            ),
            "quality_checks_failed": (
                failed
            ),
            "quality_gate": (
                quality_gate
            ),
            "final_status": (
                final_status
            ),
        }
    ],
)


REPORT.write_text(
    "\n".join(
        [
            "# PLOS ONE Readable OOXML Redline Repair",
            "",
            "Manuscript: PONE-D-26-30583",
            "",
            "R1E14E visual QA identified inline concatenation of tracked deletions and insertions in the title, short title, and two word-level Methods revisions.",
            "",
            "The repaired candidate is rebuilt from the locked clean manuscript. The same 71 reviewer/editor revision targets are retained. Whole additions remain tracked insertions. The title, short title, and prior word-level revisions use insertion-only token diffs so the visible manuscript remains readable.",
            "",
            f"- Clean DOCX SHA256: `{clean_sha}`",
            f"- Prior marked-up SHA256: `{prior_marked_sha}`",
            f"- Candidate SHA256: `{candidate_sha}`",
            f"- Checks passed: {passed}/{len(checks)}",
            f"- Quality gate: `{quality_gate}`",
            f"- Final status: `{final_status}`",
            "",
        ]
    ),
    encoding="utf-8",
    newline="\n",
)


if failed:
    fail(
        "Readable redline repair failed "
        f"{failed} quality check(s)."
    )


shutil.copy2(
    CANDIDATE_DOCX,
    MARKED_DOCX,
)

promoted_sha = sha256(
    MARKED_DOCX
)


print(
    "===== PLOS ONE READABLE OOXML REDLINE REPAIR ====="
)

print(
    f"clean_docx_sha256\t{clean_sha}"
)

print(
    f"baseline_docx_sha256\t{baseline_sha}"
)

print(
    f"scientific_source_sha256\t{source_sha}"
)

print(
    f"prior_marked_up_sha256\t{prior_marked_sha}"
)

print(
    f"candidate_docx_sha256\t{candidate_sha}"
)

print(
    f"promoted_docx_sha256\t{promoted_sha}"
)

print(
    f"selected_paragraphs\t{len(manifest_rows)}"
)

print(
    f"tracked_insertions\t{tracked_insertions}"
)

print(
    f"tracked_deletions\t{tracked_deletions}"
)

print(
    "accepted_content_matches_clean\t"
    f"{accepted_signature == clean_signature}"
)

print(
    f"smoke_render_pdf_bytes\t{smoke_bytes}"
)

print(
    f"quality_checks_passed\t{passed}/{len(checks)}"
)

print(
    f"quality_gate\t{quality_gate}"
)

print(
    f"final_status\t{final_status}"
)

print(
    f"manifest\t{MANIFEST}"
)

print(
    f"summary\t{SUMMARY}"
)

print(
    f"quality_gate_file\t{QUALITY_GATE}"
)

print(
    f"report\t{REPORT}"
)

print()

print(
    "===== LIBREOFFICE SMOKE OUTPUT ====="
)

print(
    render_result.stdout.strip()
)
