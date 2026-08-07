#!/usr/bin/env python3

from __future__ import annotations

import csv
import hashlib
import os
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE_MD = Path(
    "docs/"
    "complete_manuscript_draft_v2.3_submission_candidate_metadata_restored.md"
)

EXPECTED_SOURCE_SHA = (
    "f3b61e6ddb9f5d38c6211c6cfe0d8694e6ca3b761d52a3245d58df844ab5b2ae"
)

FINAL_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

EXPECTED_INPUT_SHA = (
    "a21d3599e32d7479dd95402135eaaa27ee2719107bf417a727f721b082f95c8c"
)

WORK = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E13C_plos_read_order"
)

BACKUP = (
    WORK
    / "PONE-D-26-30583_clean_revised_manuscript_pre_read_order.docx"
)

CANDIDATE = (
    WORK
    / "PONE-D-26-30583_clean_revised_manuscript_plos_read_order_candidate.docx"
)

SMOKE_DIR = (
    WORK
    / "libreoffice_smoke_render"
)

PROFILE_DIR = (
    WORK
    / "libreoffice_profile"
)

OUT = Path(
    "results/revision_round1/"
    "plosone_clean_manuscript_plos_read_order_v2.3"
)

GATE = (
    OUT
    / "PLOS_ONE_clean_manuscript_plos_read_order_quality_gate.tsv"
)

SUMMARY = (
    OUT
    / "PLOS_ONE_clean_manuscript_plos_read_order_summary.tsv"
)

STRUCTURE = (
    OUT
    / "PLOS_ONE_clean_manuscript_plos_read_order_structure_audit.tsv"
)

LABELS = (
    OUT
    / "PLOS_ONE_clean_manuscript_figure_label_audit.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_clean_manuscript_plos_read_order_report.md"
)


def die(message: str) -> None:

    print(
        f"ERROR: {message}",
        file=sys.stderr,
    )

    sys.exit(1)


def sha256(path: Path) -> str:

    digest = hashlib.sha256()

    with path.open(
        "rb"
    ) as handle:

        for chunk in iter(
            lambda: handle.read(
                1024 * 1024
            ),
            b"",
        ):
            digest.update(
                chunk
            )

    return digest.hexdigest()


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:

    if not rows:
        die(
            f"No rows for {path}"
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=list(
                rows[0]
            ),
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()
        writer.writerows(
            rows
        )


def entries(
    doc: Document,
):

    output = []

    for child in (
        doc.element.body.iterchildren()
    ):

        if child.tag == qn(
            "w:p"
        ):

            output.append(
                (
                    "P",
                    Paragraph(
                        child,
                        doc,
                    ),
                    child,
                )
            )

        elif child.tag == qn(
            "w:tbl"
        ):

            output.append(
                (
                    "T",
                    Table(
                        child,
                        doc,
                    ),
                    child,
                )
            )

    return output


def remove(
    element,
) -> None:

    parent = (
        element.getparent()
    )

    if parent is not None:
        parent.remove(
            element
        )


def insert_after(
    anchor,
    elements: list,
):

    current = anchor

    for element in elements:

        current.addnext(
            element
        )

        current = element

    return current


def exact_paragraph(
    doc: Document,
    text: str,
) -> Paragraph:

    hits = [
        paragraph
        for paragraph
        in doc.paragraphs
        if paragraph.text.strip()
        == text
    ]

    if len(hits) != 1:

        die(
            f"Expected one exact paragraph "
            f"{text!r}; observed "
            f"{len(hits)}"
        )

    return hits[0]


def prefix_paragraph(
    doc: Document,
    prefix: str,
) -> Paragraph:

    hits = [
        paragraph
        for paragraph
        in doc.paragraphs
        if paragraph.text.strip().startswith(
            prefix
        )
    ]

    if len(hits) != 1:

        die(
            f"Expected one paragraph "
            f"beginning {prefix!r}; "
            f"observed {len(hits)}"
        )

    return hits[0]


def first_citation(
    doc: Document,
    token: str,
    excluded: tuple[str, ...],
) -> Paragraph:

    for paragraph in (
        doc.paragraphs
    ):

        text = (
            paragraph.text.strip()
        )

        if token not in text:
            continue

        if any(
            text.startswith(
                prefix
            )
            for prefix
            in excluded
        ):
            continue

        return paragraph

    die(
        f"No citation found for "
        f"{token}"
    )

    raise AssertionError


def capture_between(
    doc: Document,
    start_prefix: str,
    end_prefix: str | None = None,
    end_exact: str | None = None,
) -> list:

    body = entries(
        doc
    )

    start = None
    end = None

    for index, entry in enumerate(
        body
    ):

        text = (
            entry[1].text.strip()
            if entry[0] == "P"
            else ""
        )

        if (
            start is None
            and text.startswith(
                start_prefix
            )
        ):

            start = index
            continue

        if start is not None:

            if (
                end_prefix is not None
                and text.startswith(
                    end_prefix
                )
            ):

                end = index
                break

            if (
                end_exact is not None
                and text == end_exact
            ):

                end = index
                break

    if (
        start is None
        or end is None
    ):

        die(
            "Could not capture block "
            f"{start_prefix!r}"
        )

    return [
        item[2]
        for item
        in body[
            start:end
        ]
    ]


def capture_table_section(
    doc: Document,
    label: str,
) -> list:

    body = entries(
        doc
    )

    start = None
    notes = None

    for index, entry in enumerate(
        body
    ):

        if entry[0] != "P":
            continue

        text = (
            entry[1].text.strip()
        )

        if (
            start is None
            and text == label
        ):

            start = index
            continue

        if (
            start is not None
            and notes is None
            and text == "Notes"
        ):

            notes = index
            continue

        if (
            start is not None
            and notes is not None
            and index > notes
        ):

            style = (
                entry[1].style.name
                if entry[1].style
                is not None
                else ""
            )

            if style == "Heading 1":

                return [
                    item[2]
                    for item
                    in body[
                        start:index
                    ]
                ]

    if (
        start is None
        or notes is None
    ):

        die(
            "Could not capture "
            f"complete section for "
            f"{label}"
        )

    return [
        item[2]
        for item
        in body[start:]
    ]


def table_matrix(
    table: Table,
) -> list[list[str]]:

    return [
        [
            cell.text
            for cell
            in row.cells
        ]
        for row
        in table.rows
    ]


def format_caption_title(
    paragraph: Paragraph,
    size: float = 12.0,
) -> None:

    paragraph.style = (
        paragraph.part.document.styles[
            "Normal"
        ]
    )

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
    )

    paragraph.paragraph_format.keep_with_next = (
        True
    )

    paragraph.paragraph_format.space_before = (
        Pt(6)
    )

    paragraph.paragraph_format.space_after = (
        Pt(0)
    )

    paragraph.paragraph_format.line_spacing_rule = (
        WD_LINE_SPACING.DOUBLE
    )

    for run in (
        paragraph.runs
    ):

        run.font.name = (
            "Times New Roman"
        )

        run.font.size = (
            Pt(size)
        )

        run.font.bold = (
            True
        )


def rebuild_table_block(
    doc: Document,
    block: list,
    number: int,
):

    table_elements = [
        element
        for element
        in block
        if element.tag
        == qn("w:tbl")
    ]

    if len(
        table_elements
    ) != 1:

        die(
            f"Table {number}: "
            f"expected one table, "
            f"observed "
            f"{len(table_elements)}"
        )

    table_element = (
        table_elements[0]
    )

    matrix_before = (
        table_matrix(
            Table(
                table_element,
                doc,
            )
        )
    )

    records = []

    for index, element in enumerate(
        block
    ):

        if element.tag == qn(
            "w:p"
        ):

            records.append(
                (
                    index,
                    Paragraph(
                        element,
                        doc,
                    ).text.strip(),
                    element,
                )
            )

    def one(
        text: str,
    ):

        hits = [
            record
            for record
            in records
            if record[1] == text
        ]

        if len(hits) != 1:

            die(
                f"Table {number}: "
                f"expected one "
                f"{text!r}, observed "
                f"{len(hits)}"
            )

        return hits[0]

    label = one(
        f"Table {number}"
    )

    title_marker = one(
        "Title"
    )

    editable = one(
        "Editable table"
    )

    notes = one(
        "Notes"
    )

    title_candidates = [
        record
        for record
        in records
        if (
            title_marker[0]
            < record[0]
            < editable[0]
            and record[1]
        )
    ]

    if len(
        title_candidates
    ) != 1:

        die(
            f"Table {number}: "
            f"expected one "
            f"title-content paragraph, "
            f"observed "
            f"{len(title_candidates)}"
        )

    title_content = (
        title_candidates[0]
    )

    note_body = [
        element
        for index, element
        in enumerate(
            block
        )
        if (
            index > notes[0]
            and (
                element.tag
                != qn("w:p")
                or Paragraph(
                    element,
                    doc,
                ).text.strip()
            )
        )
    ]

    for marker in (
        title_marker[2],
        editable[2],
        notes[2],
    ):

        remove(
            marker
        )

    label_paragraph = (
        Paragraph(
            label[2],
            doc,
        )
    )

    label_paragraph.text = (
        f"Table {number}. "
        f"{title_content[1]}"
    )

    format_caption_title(
        label_paragraph,
        size=11.0,
    )

    remove(
        title_content[2]
    )

    new_block = [
        label[2],
        table_element,
    ]

    new_block.extend(
        element
        for element
        in note_body
        if element not in new_block
    )

    return (
        new_block,
        matrix_before,
        label_paragraph.text,
    )


REPLACEMENTS = [
    (
        "Supplementary Figure S1",
        "S1 Fig",
    ),
    (
        "Figure S1A",
        "S1 Fig, panel A",
    ),
    (
        "Figure S1B",
        "S1 Fig, panel B",
    ),
    (
        "Figure S1C",
        "S1 Fig, panel C",
    ),
    (
        "Figure S1",
        "S1 Fig",
    ),
    (
        "Figure 1",
        "Fig 1",
    ),
    (
        "Figure 2",
        "Fig 2",
    ),
    (
        "Figure 3",
        "Fig 3",
    ),
]


def replace_labels_in_paragraph(
    paragraph: Paragraph,
) -> int:

    original = (
        paragraph.text
    )

    if not any(
        old in original
        for old, _
        in REPLACEMENTS
    ):

        return 0

    for run in (
        paragraph.runs
    ):

        value = (
            run.text
        )

        for old, new in (
            REPLACEMENTS
        ):

            value = (
                value.replace(
                    old,
                    new,
                )
            )

        run.text = value

    if any(
        old in paragraph.text
        for old, _
        in REPLACEMENTS
    ):

        value = original

        for old, new in (
            REPLACEMENTS
        ):

            value = (
                value.replace(
                    old,
                    new,
                )
            )

        paragraph.text = value

    return int(
        paragraph.text
        != original
    )


def normalize_labels(
    doc: Document,
) -> int:

    changed = 0

    for paragraph in (
        doc.paragraphs
    ):

        changed += (
            replace_labels_in_paragraph(
                paragraph
            )
        )

    for table in (
        doc.tables
    ):

        for row in (
            table.rows
        ):

            for cell in (
                row.cells
            ):

                for paragraph in (
                    cell.paragraphs
                ):

                    changed += (
                        replace_labels_in_paragraph(
                            paragraph
                        )
                    )

    return changed


def collect_text(
    doc: Document,
) -> str:

    chunks = [
        paragraph.text
        for paragraph
        in doc.paragraphs
    ]

    chunks.extend(
        cell.text
        for table
        in doc.tables
        for row
        in table.rows
        for cell
        in row.cells
    )

    return "\n".join(
        chunks
    )


def package_state(
    path: Path,
) -> dict[str, object]:

    with zipfile.ZipFile(
        path
    ) as archive:

        names = (
            archive.namelist()
        )

        document_xml = (
            archive.read(
                "word/document.xml"
            ).decode(
                "utf-8",
                errors="replace",
            )
        )

        settings_xml = (
            archive.read(
                "word/settings.xml"
            ).decode(
                "utf-8",
                errors="replace",
            )
        )

        footer_xml = "\n".join(
            archive.read(
                name
            ).decode(
                "utf-8",
                errors="replace",
            )
            for name
            in names
            if (
                name.startswith(
                    "word/footer"
                )
                and name.endswith(
                    ".xml"
                )
            )
        )

    return {
        "comments": [
            name
            for name
            in names
            if (
                "comment"
                in name.lower()
                or "people"
                in name.lower()
                or "person"
                in name.lower()
            )
        ],
        "media": [
            name
            for name
            in names
            if name.startswith(
                "word/media/"
            )
        ],
        "line_numbering": (
            "<w:lnNumType"
            in document_xml
        ),
        "page_number_field": (
            " PAGE "
            in footer_xml
        ),
        "tracked_insertions": len(
            re.findall(
                r"<w:ins\b",
                document_xml,
            )
        ),
        "tracked_deletions": len(
            re.findall(
                r"<w:del\b",
                document_xml,
            )
        ),
        "protection": (
            "<w:documentProtection"
            in settings_xml
        ),
    }


def body_pos(
    doc: Document,
    element,
) -> int:

    children = list(
        doc.element.body.iterchildren()
    )

    try:

        return (
            children.index(
                element
            )
            + 1
        )

    except ValueError:

        return -1


def smoke_render(
    path: Path,
):

    if shutil.which(
        "soffice"
    ) is None:

        die(
            "soffice not found"
        )

    SMOKE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    PROFILE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    for pdf in (
        SMOKE_DIR.glob(
            "*.pdf"
        )
    ):

        pdf.unlink()

    abs_profile = (
        PROFILE_DIR.resolve()
    )

    environment = (
        os.environ.copy()
    )

    environment[
        "HOME"
    ] = str(
        abs_profile
    )

    result = subprocess.run(
        [
            "soffice",
            "--headless",
            (
                "-env:UserInstallation="
                f"file://{abs_profile}"
            ),
            "--convert-to",
            "pdf",
            "--outdir",
            str(
                SMOKE_DIR.resolve()
            ),
            str(
                path.resolve()
            ),
        ],
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        env=environment,
    )

    pdf = (
        SMOKE_DIR
        / f"{path.stem}.pdf"
    )

    success = (
        pdf.exists()
        and pdf.stat().st_size
        > 0
    )

    return (
        success,
        result.returncode,
        pdf,
        result.stdout,
    )


# ============================================================================
# Preflight
# ============================================================================

WORK.mkdir(
    parents=True,
    exist_ok=True,
)

OUT.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)

if (
    not SOURCE_MD.exists()
    or not FINAL_DOCX.exists()
):

    die(
        "Required source or "
        "clean DOCX is missing"
    )

source_sha = (
    sha256(
        SOURCE_MD
    )
)

input_sha = (
    sha256(
        FINAL_DOCX
    )
)

if (
    source_sha
    != EXPECTED_SOURCE_SHA
):

    die(
        "Source SHA mismatch: "
        f"{source_sha}"
    )

if (
    input_sha
    != EXPECTED_INPUT_SHA
):

    die(
        "Input DOCX SHA mismatch: "
        f"{input_sha}"
    )

if BACKUP.exists():

    if (
        sha256(
            BACKUP
        )
        != EXPECTED_INPUT_SHA
    ):

        die(
            "Existing pre-read-order "
            "backup has unexpected SHA"
        )

else:

    shutil.copy2(
        FINAL_DOCX,
        BACKUP,
    )

backup_sha = (
    sha256(
        BACKUP
    )
)

doc = Document(
    BACKUP
)

input_tables = [
    table_matrix(
        table
    )
    for table
    in doc.tables
]

if (
    len(
        input_tables
    ) != 2
    or len(
        doc.sections
    ) != 1
):

    die(
        "Unexpected input "
        "table/section count"
    )


# ============================================================================
# Locate citations and source blocks
# ============================================================================

fig1_citation = (
    first_citation(
        doc,
        "Figure 1",
        (
            "Figure 1.",
            "Alt text for Figure 1",
        ),
    )
)

fig2_citation = (
    first_citation(
        doc,
        "Figure 2",
        (
            "Figure 2.",
            "Alt text for Figure 2",
        ),
    )
)

fig3_citation = (
    first_citation(
        doc,
        "Figure 3",
        (
            "Figure 3.",
            "Alt text for Figure 3",
        ),
    )
)

fig1_block = (
    capture_between(
        doc,
        "Figure 1.",
        end_prefix="Figure 2.",
    )
)

fig2_block = (
    capture_between(
        doc,
        "Figure 2.",
        end_prefix="Figure 3.",
    )
)

fig3_block = (
    capture_between(
        doc,
        "Figure 3.",
        end_prefix="Figure S1.",
    )
)

s1_block = (
    capture_between(
        doc,
        "Figure S1.",
        end_exact="Table 1",
    )
)

(
    table1_block,
    table1_before,
    table1_title,
) = rebuild_table_block(
    doc,
    capture_table_section(
        doc,
        "Table 1",
    ),
    1,
)

(
    table2_block,
    table2_before,
    table2_title,
) = rebuild_table_block(
    doc,
    capture_table_section(
        doc,
        "Table 2",
    ),
    2,
)

if (
    table1_before
    != input_tables[0]
    or table2_before
    != input_tables[1]
):

    die(
        "A table matrix changed "
        "before relocation"
    )


# ============================================================================
# Relocate content
# ============================================================================

remove(
    exact_paragraph(
        doc,
        "Figure captions",
    )._p
)

insert_after(
    fig1_citation._p,
    fig1_block,
)

fig2_last = insert_after(
    fig2_citation._p,
    fig2_block,
)

insert_after(
    fig2_last,
    table1_block,
)

fig3_last = insert_after(
    fig3_citation._p,
    fig3_block,
)

insert_after(
    fig3_last,
    table2_block,
)

supporting_heading = (
    doc.add_paragraph(
        "Supporting information captions"
    )
)

if (
    "Heading 1"
    in doc.styles
):

    supporting_heading.style = (
        doc.styles[
            "Heading 1"
        ]
    )

supporting_heading.paragraph_format.keep_with_next = (
    True
)

insert_after(
    supporting_heading._p,
    s1_block,
)

changed_labels = (
    normalize_labels(
        doc
    )
)

for prefix in (
    "Fig 1.",
    "Fig 2.",
    "Fig 3.",
    "S1 Fig.",
):

    format_caption_title(
        prefix_paragraph(
            doc,
            prefix,
        )
    )


# ============================================================================
# Save and reopen candidate
# ============================================================================

if CANDIDATE.exists():

    CANDIDATE.unlink()

doc.save(
    CANDIDATE
)

candidate = Document(
    CANDIDATE
)

candidate_text = (
    collect_text(
        candidate
    )
)

state = (
    package_state(
        CANDIDATE
    )
)

candidate_matrices = [
    table_matrix(
        table
    )
    for table
    in candidate.tables
]

candidate_sha = (
    sha256(
        CANDIDATE
    )
)


# ============================================================================
# Independent structure audit
# ============================================================================

fig1_c = first_citation(
    candidate,
    "Fig 1",
    (
        "Fig 1.",
        "Alt text for Fig 1",
    ),
)

fig2_c = first_citation(
    candidate,
    "Fig 2",
    (
        "Fig 2.",
        "Alt text for Fig 2",
    ),
)

fig3_c = first_citation(
    candidate,
    "Fig 3",
    (
        "Fig 3.",
        "Alt text for Fig 3",
    ),
)

fig1_t = prefix_paragraph(
    candidate,
    "Fig 1.",
)

fig2_t = prefix_paragraph(
    candidate,
    "Fig 2.",
)

fig3_t = prefix_paragraph(
    candidate,
    "Fig 3.",
)

s1_t = prefix_paragraph(
    candidate,
    "S1 Fig.",
)

t1_t = prefix_paragraph(
    candidate,
    "Table 1.",
)

t2_t = prefix_paragraph(
    candidate,
    "Table 2.",
)

references = exact_paragraph(
    candidate,
    "References",
)

supporting = exact_paragraph(
    candidate,
    "Supporting information captions",
)

table_elements = [
    item[2]
    for item
    in entries(
        candidate
    )
    if item[0]
    == "T"
]

if len(
    table_elements
) != 2:

    die(
        "Candidate does not contain "
        "exactly two Word tables"
    )

structure_rows = [
    {
        "item": "Fig 1",
        "citation_body_position": body_pos(
            candidate,
            fig1_c._p,
        ),
        "title_body_position": body_pos(
            candidate,
            fig1_t._p,
        ),
        "table_body_position": "",
        "status": (
            "PASS"
            if body_pos(
                candidate,
                fig1_t._p,
            )
            == body_pos(
                candidate,
                fig1_c._p,
            ) + 1
            else "FAIL"
        ),
    },
    {
        "item": "Fig 2",
        "citation_body_position": body_pos(
            candidate,
            fig2_c._p,
        ),
        "title_body_position": body_pos(
            candidate,
            fig2_t._p,
        ),
        "table_body_position": "",
        "status": (
            "PASS"
            if body_pos(
                candidate,
                fig2_t._p,
            )
            == body_pos(
                candidate,
                fig2_c._p,
            ) + 1
            else "FAIL"
        ),
    },
    {
        "item": "Table 1",
        "citation_body_position": body_pos(
            candidate,
            fig2_c._p,
        ),
        "title_body_position": body_pos(
            candidate,
            t1_t._p,
        ),
        "table_body_position": body_pos(
            candidate,
            table_elements[0],
        ),
        "status": (
            "PASS"
            if body_pos(
                candidate,
                table_elements[0],
            )
            == body_pos(
                candidate,
                t1_t._p,
            ) + 1
            else "FAIL"
        ),
    },
    {
        "item": "Fig 3",
        "citation_body_position": body_pos(
            candidate,
            fig3_c._p,
        ),
        "title_body_position": body_pos(
            candidate,
            fig3_t._p,
        ),
        "table_body_position": "",
        "status": (
            "PASS"
            if body_pos(
                candidate,
                fig3_t._p,
            )
            == body_pos(
                candidate,
                fig3_c._p,
            ) + 1
            else "FAIL"
        ),
    },
    {
        "item": "Table 2",
        "citation_body_position": body_pos(
            candidate,
            fig3_c._p,
        ),
        "title_body_position": body_pos(
            candidate,
            t2_t._p,
        ),
        "table_body_position": body_pos(
            candidate,
            table_elements[1],
        ),
        "status": (
            "PASS"
            if body_pos(
                candidate,
                table_elements[1],
            )
            == body_pos(
                candidate,
                t2_t._p,
            ) + 1
            else "FAIL"
        ),
    },
    {
        "item": "S1 Fig",
        "citation_body_position": "",
        "title_body_position": body_pos(
            candidate,
            s1_t._p,
        ),
        "table_body_position": "",
        "status": (
            "PASS"
            if (
                body_pos(
                    candidate,
                    supporting._p,
                )
                > body_pos(
                    candidate,
                    references._p,
                )
                and body_pos(
                    candidate,
                    s1_t._p,
                )
                == body_pos(
                    candidate,
                    supporting._p,
                ) + 1
            )
            else "FAIL"
        ),
    },
]

write_tsv(
    STRUCTURE,
    structure_rows,
)


# ============================================================================
# Label audit
# ============================================================================

old_labels = [
    "Figure 1",
    "Figure 2",
    "Figure 3",
    "Figure S1",
    "Supplementary Figure S1",
]

new_labels = [
    "Fig 1",
    "Fig 2",
    "Fig 3",
    "S1 Fig",
]

label_rows = [
    {
        "label_type": "obsolete",
        "token": token,
        "occurrences": (
            candidate_text.count(
                token
            )
        ),
        "expected": 0,
    }
    for token
    in old_labels
]

label_rows.extend(
    {
        "label_type": "PLOS",
        "token": token,
        "occurrences": (
            candidate_text.count(
                token
            )
        ),
        "expected": ">0",
    }
    for token
    in new_labels
)

write_tsv(
    LABELS,
    label_rows,
)


# ============================================================================
# Smoke render
# ============================================================================

(
    smoke_ok,
    soffice_status,
    smoke_pdf,
    soffice_output,
) = smoke_render(
    CANDIDATE
)

smoke_bytes = (
    smoke_pdf.stat().st_size
    if smoke_pdf.exists()
    else 0
)


# ============================================================================
# QA
# ============================================================================

checks: list[
    dict[str, object]
] = []


def check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": (
                description
            ),
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected
            ),
        }
    )


check(
    "Scientific source SHA locked",
    source_sha
    == EXPECTED_SOURCE_SHA,
    source_sha,
    EXPECTED_SOURCE_SHA,
)

check(
    "Input DOCX SHA locked",
    input_sha
    == EXPECTED_INPUT_SHA,
    input_sha,
    EXPECTED_INPUT_SHA,
)

check(
    "Backup SHA locked",
    backup_sha
    == EXPECTED_INPUT_SHA,
    backup_sha,
    EXPECTED_INPUT_SHA,
)

check(
    "Candidate created",
    CANDIDATE.exists(),
    CANDIDATE.exists(),
    True,
)

check(
    "Two Word tables retained",
    len(
        candidate.tables
    ) == 2,
    len(
        candidate.tables
    ),
    2,
)

check(
    "One section retained",
    len(
        candidate.sections
    ) == 1,
    len(
        candidate.sections
    ),
    1,
)

check(
    "Table 1 matrix preserved",
    (
        candidate_matrices[0]
        == table1_before
    ),
    (
        "identical"
        if candidate_matrices[0]
        == table1_before
        else "different"
    ),
    "identical",
)

check(
    "Table 2 matrix preserved",
    (
        candidate_matrices[1]
        == table2_before
    ),
    (
        "identical"
        if candidate_matrices[1]
        == table2_before
        else "different"
    ),
    "identical",
)

check(
    "Figure captions heading removed",
    all(
        paragraph.text.strip()
        != "Figure captions"
        for paragraph
        in candidate.paragraphs
    ),
    sum(
        paragraph.text.strip()
        == "Figure captions"
        for paragraph
        in candidate.paragraphs
    ),
    0,
)

for marker in (
    "Title",
    "Editable table",
    "Notes",
):

    check(
        (
            "Table scaffolding "
            f"{marker!r} removed"
        ),
        all(
            paragraph.text.strip()
            != marker
            for paragraph
            in candidate.paragraphs
        ),
        sum(
            paragraph.text.strip()
            == marker
            for paragraph
            in candidate.paragraphs
        ),
        0,
    )

check(
    "Table 1 title preserved",
    table1_title
    in candidate_text,
    (
        "present"
        if table1_title
        in candidate_text
        else "absent"
    ),
    "present",
)

check(
    "Table 2 title preserved",
    table2_title
    in candidate_text,
    (
        "present"
        if table2_title
        in candidate_text
        else "absent"
    ),
    "present",
)

for row in (
    structure_rows
):

    check(
        (
            "Read-order structure "
            f"{row['item']}"
        ),
        row["status"]
        == "PASS",
        row["status"],
        "PASS",
    )

for token in (
    old_labels
):

    check(
        (
            "Obsolete label "
            f"{token!r} absent"
        ),
        token
        not in candidate_text,
        candidate_text.count(
            token
        ),
        0,
    )

for token in (
    new_labels
):

    check(
        (
            "PLOS label "
            f"{token!r} present"
        ),
        token
        in candidate_text,
        candidate_text.count(
            token
        ),
        ">0",
    )

anchors = [
    "Reuben S. Maghembe¹˒²*",
    "224 samples",
    "101 bacterial and 123 viral",
    "52 DefiniteBacterial and 94 DefiniteViral",
    "23 definite bacterial and 28 definite viral",
    "29,826",
    "0.9940",
    "0.9874",
    "0.9814",
    "ChatGPT",
    "OpenAI",
    "plosone_revision_round1_2026",
    (
        "The GSE72810 and GSE73461 "
        "GEO sample accession sets "
        "were disjoint"
    ),
    (
        "(A) Primary bacterial-versus-viral "
        "discovery analysis in GSE211567."
    ),
    (
        "(C) BH-adjusted Wilcoxon P values "
        "for the main projection"
    ),
    (
        "GSE72810 was analysed as a second "
        "accession-level and sample-level cohort "
        "providing cross-platform validation."
    ),
    (
        "C, Exhaustive leave-one/two-gene "
        "robustness in GSE73461."
    ),
]

for anchor in anchors:

    check(
        (
            "Anchor preserved: "
            f"{anchor[:45]}"
        ),
        anchor
        in candidate_text,
        (
            "present"
            if anchor
            in candidate_text
            else "absent"
        ),
        "present",
    )

check(
    "No embedded figures",
    len(
        state["media"]
    ) == 0,
    len(
        state["media"]
    ),
    0,
)

check(
    "No comment package",
    len(
        state["comments"]
    ) == 0,
    (
        "|".join(
            state["comments"]
        )
        or "NONE"
    ),
    "NONE",
)

check(
    "Line numbering retained",
    bool(
        state[
            "line_numbering"
        ]
    ),
    state[
        "line_numbering"
    ],
    True,
)

check(
    "Page-number field retained",
    bool(
        state[
            "page_number_field"
        ]
    ),
    state[
        "page_number_field"
    ],
    True,
)

check(
    "No tracked insertions",
    int(
        state[
            "tracked_insertions"
        ]
    ) == 0,
    state[
        "tracked_insertions"
    ],
    0,
)

check(
    "No tracked deletions",
    int(
        state[
            "tracked_deletions"
        ]
    ) == 0,
    state[
        "tracked_deletions"
    ],
    0,
)

check(
    "Document unprotected",
    not bool(
        state[
            "protection"
        ]
    ),
    state[
        "protection"
    ],
    False,
)

check(
    "LibreOffice smoke render",
    smoke_ok,
    (
        f"status={soffice_status}; "
        f"bytes={smoke_bytes}"
    ),
    "non-empty PDF",
)


# ============================================================================
# Promote only after gate
# ============================================================================

pre_failures = sum(
    row["pass"] == "FALSE"
    for row
    in checks
)

promoted_sha = (
    input_sha
)

promotion_ok = False

if pre_failures == 0:

    shutil.copy2(
        CANDIDATE,
        FINAL_DOCX,
    )

    promoted_sha = (
        sha256(
            FINAL_DOCX
        )
    )

    promotion_ok = (
        promoted_sha
        == candidate_sha
    )

check(
    (
        "Candidate promoted "
        "to canonical DOCX"
    ),
    (
        pre_failures == 0
        and promotion_ok
    ),
    (
        promoted_sha
        if promotion_ok
        else "not promoted"
    ),
    (
        candidate_sha
        if pre_failures == 0
        else "promotion withheld"
    ),
)

passed = sum(
    row["pass"] == "TRUE"
    for row
    in checks
)

failed = (
    len(checks)
    - passed
)

gate = (
    "PASS"
    if failed == 0
    else "FAIL"
)

status = (
    "READY_FOR_FULL_RENDER_AND_VISUAL_QA"
    if failed == 0
    else "PLOS_READ_ORDER_REQUIRES_CORRECTION"
)


# ============================================================================
# Outputs
# ============================================================================

write_tsv(
    GATE,
    checks,
)

write_tsv(
    SUMMARY,
    [
        {
            "scientific_source_sha256": (
                source_sha
            ),
            "input_docx_sha256": (
                input_sha
            ),
            "candidate_docx_sha256": (
                candidate_sha
            ),
            "promoted_docx_sha256": (
                promoted_sha
            ),
            "candidate_paragraphs": len(
                candidate.paragraphs
            ),
            "tables": len(
                candidate.tables
            ),
            "sections": len(
                candidate.sections
            ),
            "figure_label_paragraphs_changed": (
                changed_labels
            ),
            "embedded_media_files": len(
                state["media"]
            ),
            "comment_related_parts": len(
                state["comments"]
            ),
            "smoke_render_pdf_bytes": (
                smoke_bytes
            ),
            "quality_checks": len(
                checks
            ),
            "quality_checks_passed": (
                passed
            ),
            "quality_checks_failed": (
                failed
            ),
            "quality_gate": (
                gate
            ),
            "final_status": (
                status
            ),
        }
    ],
)

REPORT.write_text(
    "\n".join(
        [
            (
                "# PLOS ONE Clean Manuscript "
                "Read-Order Normalization"
            ),
            "",
            (
                "Scientific source SHA256: "
                f"`{source_sha}`"
            ),
            (
                "Input clean DOCX SHA256: "
                f"`{input_sha}`"
            ),
            (
                "Candidate SHA256: "
                f"`{candidate_sha}`"
            ),
            (
                "Promoted DOCX SHA256: "
                f"`{promoted_sha}`"
            ),
            "",
            (
                "Changes: main figure captions "
                "and Tables 1-2 moved into read "
                "order; S1 Fig caption moved to "
                "the end; figure labels normalized; "
                "internal table scaffolding removed "
                "while table cell matrices were "
                "preserved."
            ),
            "",
            (
                f"Checks passed: "
                f"{passed}/{len(checks)}"
            ),
            (
                f"Quality gate: "
                f"`{gate}`"
            ),
            (
                f"Final status: "
                f"`{status}`"
            ),
            "",
            (
                "Full page rendering and visual "
                "inspection remain required "
                "before acceptance."
            ),
            "",
        ]
    ),
    encoding="utf-8",
    newline="\n",
)


# ============================================================================
# Console
# ============================================================================

print(
    "===== PLOS ONE CLEAN MANUSCRIPT "
    "READ-ORDER NORMALIZATION V2 ====="
)

print(
    f"scientific_source_sha256\t"
    f"{source_sha}"
)

print(
    f"input_docx_sha256\t"
    f"{input_sha}"
)

print(
    f"candidate_docx_sha256\t"
    f"{candidate_sha}"
)

print(
    f"promoted_docx_sha256\t"
    f"{promoted_sha}"
)

print(
    f"candidate_paragraphs\t"
    f"{len(candidate.paragraphs)}"
)

print(
    f"tables\t"
    f"{len(candidate.tables)}"
)

print(
    f"sections\t"
    f"{len(candidate.sections)}"
)

print(
    "figure_label_paragraphs_changed\t"
    f"{changed_labels}"
)

print(
    "embedded_media_files\t"
    f"{len(state['media'])}"
)

print(
    "comment_related_parts\t"
    f"{len(state['comments'])}"
)

print(
    f"smoke_render_pdf_bytes\t"
    f"{smoke_bytes}"
)

print(
    f"soffice_status\t"
    f"{soffice_status}"
)

print(
    f"quality_checks_passed\t"
    f"{passed}/{len(checks)}"
)

print(
    f"quality_gate\t"
    f"{gate}"
)

print(
    f"final_status\t"
    f"{status}"
)

print(
    f"final_docx\t"
    f"{FINAL_DOCX}"
)

print(
    f"structure_audit\t"
    f"{STRUCTURE}"
)

print(
    f"label_audit\t"
    f"{LABELS}"
)

print(
    f"summary\t"
    f"{SUMMARY}"
)

print(
    f"quality_gate_file\t"
    f"{GATE}"
)

print(
    f"report\t"
    f"{REPORT}"
)

print()
print(
    "===== LIBREOFFICE SMOKE OUTPUT ====="
)

print(
    soffice_output.strip()
)

if failed:

    die(
        "PLOS read-order normalization "
        f"failed {failed} "
        "quality check(s)."
    )
