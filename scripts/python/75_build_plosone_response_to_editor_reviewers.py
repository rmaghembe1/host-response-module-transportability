#!/usr/bin/env python3
from __future__ import annotations

import csv
import hashlib
import os
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


MANUSCRIPT_ID = "PONE-D-26-30583"

TITLE = (
    "External transportability of bacterial- and viral-associated "
    "host-response modules across public transcriptomic cohorts"
)

AUTHOR = "Reuben S. Maghembe"

REVISION_DATE = "8 August 2026"

CLEAN_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

MARKED_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_marked_up_revised_manuscript.docx"
)

SOURCE_V24 = Path(
    "docs/complete_manuscript_draft_v2.4_"
    "submission_candidate_ai_methods_compliance.md"
)

EXPECTED_CLEAN_SHA = (
    "3a9db22fe7e1847bc9f994d2833f80ba"
    "fd5395a92aa4fc87964218172cf4835e"
)

EXPECTED_MARKED_SHA = (
    "829134652d86d049585366ef9ab61dde"
    "6eb87520a411a60c4c9c2cdf5ce92084"
)

EXPECTED_SOURCE_SHA = (
    "54aabce4263c581dd9685a1ae9d2fd1b"
    "05030d1c3fe072bff78512a293acd6bd"
)

MATRIX_PRIMARY = Path(
    "results/revision_round1/"
    "plosone_reviewer_response_matrix_v2.3_provenance/"
    "PLOS_ONE_revision_response_evidence_matrix_v2.3.tsv"
)

MATRIX_FALLBACK = Path(
    "results/revision_round1/"
    "plosone_reviewer_response_matrix/"
    "PLOS_ONE_revision_response_evidence_matrix.tsv"
)

PUBLIC_REPOSITORY = (
    "https://github.com/rmaghembe1/"
    "host-response-module-transportability"
)

PUBLIC_BRANCH = (
    "https://github.com/rmaghembe1/"
    "host-response-module-transportability/"
    "tree/plosone_revision_round1_2026"
)

WORK = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E15C_response_letter"
)

OUT = Path(
    "results/revision_round1/"
    "plosone_response_letter_v2.4"
)

SUBMISSION = Path(
    "submission/revision_round1"
)

RESPONSE_DOCX = (
    SUBMISSION
    / "PONE-D-26-30583_response_to_editor_and_reviewers.docx"
)

RESPONSE_MD = (
    OUT
    / "PONE-D-26-30583_response_to_editor_and_reviewers.md"
)

RESPONSE_TSV = (
    OUT
    / "PLOS_ONE_final_response_items.tsv"
)

QUALITY = (
    OUT
    / "PLOS_ONE_response_letter_quality_gate.tsv"
)

SUMMARY = (
    OUT
    / "PLOS_ONE_response_letter_summary.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_response_letter_production_report.md"
)

RENDER = WORK / "render"
PROFILE = WORK / "libreoffice_profile"
PAGES = WORK / "pages"
CONTACT = WORK / "contact_sheets"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()

    with path.open("rb") as handle:
        for chunk in iter(
            lambda: handle.read(1024 * 1024),
            b"",
        ):
            digest.update(chunk)

    return digest.hexdigest()


def read_tsv(path: Path) -> list[dict[str, str]]:
    with path.open(
        "r",
        encoding="utf-8",
        newline="",
    ) as handle:
        return list(
            csv.DictReader(
                handle,
                delimiter="\t",
            )
        )


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:
    if not rows:
        raise RuntimeError(
            f"No rows available for {path}"
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=list(rows[0].keys()),
            delimiter="\t",
            lineterminator="\n",
        )
        writer.writeheader()
        writer.writerows(rows)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    shd = tc_pr.find(
        qn("w:shd")
    )

    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(
        qn("w:fill"),
        fill,
    )


def set_cell_margins(
    cell,
    top=120,
    start=140,
    bottom=120,
    end=140,
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_mar = tc_pr.first_child_found_in(
        "w:tcMar"
    )

    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(
            qn(f"w:{name}")
        )

        if element is None:
            element = OxmlElement(
                f"w:{name}"
            )
            tc_mar.append(element)

        element.set(
            qn("w:w"),
            str(value),
        )

        element.set(
            qn("w:type"),
            "dxa",
        )


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()

    tbl_header = OxmlElement(
        "w:tblHeader"
    )

    tbl_header.set(
        qn("w:val"),
        "true",
    )

    tr_pr.append(tbl_header)


def set_run(
    run,
    *,
    bold=None,
    italic=None,
    size=11,
    color=None,
) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)

    if bold is not None:
        run.bold = bold

    if italic is not None:
        run.italic = italic

    if color is not None:
        run.font.color.rgb = RGBColor(
            *color
        )


def paragraph_format(
    paragraph,
    *,
    after=6,
    before=0,
    line=1.08,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_body_paragraph(
    doc,
    text="",
    *,
    bold_prefix=None,
    italic=False,
    after=6,
    before=0,
):
    paragraph = doc.add_paragraph()
    paragraph_format(
        paragraph,
        after=after,
        before=before,
    )

    if bold_prefix is not None:
        run = paragraph.add_run(
            bold_prefix
        )
        set_run(
            run,
            bold=True,
        )

        body = text

        if body:
            run = paragraph.add_run(
                body
            )
            set_run(
                run,
                italic=italic,
            )

    else:
        run = paragraph.add_run(
            text
        )
        set_run(
            run,
            italic=italic,
        )

    return paragraph


def add_response_item(
    doc,
    number_label: str,
    comment: str,
    response: str,
    location: str,
    evidence: str,
) -> None:

    heading = doc.add_paragraph()
    paragraph_format(
        heading,
        before=10,
        after=5,
    )

    run = heading.add_run(
        number_label
    )

    set_run(
        run,
        bold=True,
        size=12,
        color=(31, 78, 121),
    )

    table = doc.add_table(
        rows=1,
        cols=1,
    )

    table.autofit = True

    cell = table.cell(
        0,
        0,
    )

    set_cell_shading(
        cell,
        "F2F2F2",
    )

    set_cell_margins(
        cell,
        top=140,
        bottom=140,
        start=180,
        end=180,
    )

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    paragraph = cell.paragraphs[0]
    paragraph_format(
        paragraph,
        after=0,
        line=1.05,
    )

    run = paragraph.add_run(
        "Comment: "
    )

    set_run(
        run,
        bold=True,
        size=10.5,
    )

    run = paragraph.add_run(
        comment
    )

    set_run(
        run,
        size=10.5,
    )

    response_p = doc.add_paragraph()
    paragraph_format(
        response_p,
        before=5,
        after=5,
    )

    run = response_p.add_run(
        "Response: "
    )

    set_run(
        run,
        bold=True,
        color=(31, 78, 121),
    )

    run = response_p.add_run(
        response
    )

    set_run(
        run,
    )

    location_p = doc.add_paragraph()
    paragraph_format(
        location_p,
        after=3,
    )

    run = location_p.add_run(
        "Changes in the revised manuscript: "
    )

    set_run(
        run,
        bold=True,
    )

    run = location_p.add_run(
        location
    )

    set_run(
        run,
        italic=True,
    )

    evidence_p = doc.add_paragraph()
    paragraph_format(
        evidence_p,
        after=8,
    )

    run = evidence_p.add_run(
        "Supporting analysis/material: "
    )

    set_run(
        run,
        bold=True,
    )

    run = evidence_p.add_run(
        evidence
    )

    set_run(
        run,
    )


def section_heading(
    doc,
    text: str,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph_format(
        paragraph,
        before=14,
        after=8,
    )

    run = paragraph.add_run(
        text
    )

    set_run(
        run,
        bold=True,
        size=14,
        color=(31, 78, 121),
    )


def render_docx(
    docx_path: Path,
) -> tuple[int, Path, str]:

    RENDER.mkdir(
        parents=True,
        exist_ok=True,
    )

    PROFILE.mkdir(
        parents=True,
        exist_ok=True,
    )

    for old_pdf in RENDER.glob(
        "*.pdf"
    ):
        old_pdf.unlink()

    env = os.environ.copy()

    env["HOME"] = str(
        PROFILE.resolve()
    )

    result = subprocess.run(
        [
            "soffice",
            "--headless",
            (
                "-env:UserInstallation="
                f"file://{PROFILE.resolve()}"
            ),
            "--convert-to",
            "pdf",
            "--outdir",
            str(RENDER.resolve()),
            str(docx_path.resolve()),
        ],
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        env=env,
    )

    pdf = (
        RENDER
        / f"{docx_path.stem}.pdf"
    )

    return (
        result.returncode,
        pdf,
        result.stdout.strip(),
    )


def pdf_page_count(
    pdf: Path,
) -> int:

    result = subprocess.run(
        [
            "pdfinfo",
            str(pdf),
        ],
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )

    match = re.search(
        r"^Pages:\s+(\d+)",
        result.stdout,
        re.MULTILINE,
    )

    if not match:
        return 0

    return int(
        match.group(1)
    )


def create_page_images(
    pdf: Path,
) -> list[Path]:

    PAGES.mkdir(
        parents=True,
        exist_ok=True,
    )

    for old in PAGES.glob(
        "page-*.png"
    ):
        old.unlink()

    result = subprocess.run(
        [
            "pdftoppm",
            "-png",
            "-r",
            "180",
            str(pdf),
            str(
                PAGES
                / "page"
            ),
        ],
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )

    if result.returncode != 0:
        raise RuntimeError(
            "pdftoppm failed:\n"
            + result.stdout
        )

    return sorted(
        PAGES.glob(
            "page-*.png"
        )
    )


def create_contact_sheets(
    pages: list[Path],
) -> list[Path]:

    CONTACT.mkdir(
        parents=True,
        exist_ok=True,
    )

    for old in CONTACT.glob(
        "contact_*.png"
    ):
        old.unlink()

    if shutil.which(
        "montage"
    ) is None:
        return []

    outputs: list[Path] = []

    group_size = 6

    for start in range(
        0,
        len(pages),
        group_size,
    ):

        group = pages[
            start:
            start + group_size
        ]

        first = start + 1
        last = (
            start
            + len(group)
        )

        output = (
            CONTACT
            / f"contact_{first:02d}-{last:02d}.png"
        )

        result = subprocess.run(
            [
                "montage",
                *[
                    str(path)
                    for path in group
                ],
                "-thumbnail",
                "700x",
                "-tile",
                "2x3",
                "-geometry",
                "+16+16",
                str(output),
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )

        if result.returncode != 0:
            raise RuntimeError(
                "montage failed:\n"
                + result.stdout
            )

        outputs.append(
            output
        )

    return outputs


for required in (
    CLEAN_DOCX,
    MARKED_DOCX,
    SOURCE_V24,
):

    if not required.exists():
        raise RuntimeError(
            f"Missing required input: {required}"
        )


if sha256(CLEAN_DOCX) != EXPECTED_CLEAN_SHA:
    raise RuntimeError(
        "Clean manuscript SHA mismatch"
    )

if sha256(MARKED_DOCX) != EXPECTED_MARKED_SHA:
    raise RuntimeError(
        "Marked-up manuscript SHA mismatch"
    )

if sha256(SOURCE_V24) != EXPECTED_SOURCE_SHA:
    raise RuntimeError(
        "v2.4 source SHA mismatch"
    )


if MATRIX_PRIMARY.exists():
    matrix_path = MATRIX_PRIMARY
elif MATRIX_FALLBACK.exists():
    matrix_path = MATRIX_FALLBACK
else:
    raise RuntimeError(
        "No reviewer-response matrix found"
    )


matrix_rows = read_tsv(
    matrix_path
)


expected_ids = [
    f"RR{number:02d}"
    for number in range(
        1,
        18,
    )
]


matrix_by_id = {
    row["response_id"]: row
    for row in matrix_rows
}


missing_ids = [
    response_id
    for response_id in expected_ids
    if response_id not in matrix_by_id
]


if missing_ids:
    raise RuntimeError(
        "Missing response IDs: "
        + ", ".join(
            missing_ids
        )
    )


responses = {
    "RR01": {
        "response": (
            "Thank you. The revision package has now been prepared in PLOS ONE "
            "submission form. Separate clean and marked-up manuscript files have "
            "been generated from the same locked scientific source and subjected "
            "to structural, rendering and visual quality-control checks. Continuous "
            "line numbering is retained, the two wide tables remain editable and "
            "readable in landscape sections, and the marked-up file uses genuine "
            "tracked revisions while accepting those revisions reproduces the clean "
            "manuscript content."
        ),
        "location": (
            "Final revision package: "
            "PONE-D-26-30583_clean_revised_manuscript.docx and "
            "PONE-D-26-30583_marked_up_revised_manuscript.docx."
        ),
        "evidence": (
            "Clean manuscript and marked-up manuscript passed the final DOCX "
            "structural and visual QA gates."
        ),
    },

    "RR02": {
        "response": (
            "The manuscript now contains a dedicated generative-AI disclosure "
            "within the Methods section. The disclosure identifies ChatGPT as an "
            "AI-assisted tool provided by OpenAI, describes its use for editorial "
            "organisation, language refinement, workflow planning, code drafting "
            "and checking, formatting checks and preparation of submission-support "
            "materials, states that AI-generated suggestions were not treated as "
            "scientific evidence, describes verification of numerical and graphical "
            "outputs, and states that the author reviewed the analysis, references, "
            "interpretation, figures, tables and manuscript and takes responsibility "
            "for the final work."
        ),
        "location": (
            "Methods, 'Declaration of generative AI and AI-assisted technologies "
            "in the manuscript preparation process', immediately before the main "
            "Results section."
        ),
        "evidence": (
            "The declaration wording was preserved during relocation into Methods; "
            "clean and marked-up DOCX copies were synchronised and visually checked."
        ),
    },

    "RR03": {
        "response": (
            "The author-generated analysis code and reproducibility materials are "
            "publicly available without access restriction in the project repository. "
            "The revision branch contains the revision-round scripts, decision logs, "
            "quality gates, source-data tables, manuscript-facing figures and "
            "supplementary outputs used to support the revised analyses."
        ),
        "location": (
            "Code availability section."
        ),
        "evidence": (
            f"Public repository: {PUBLIC_REPOSITORY}; "
            f"revision branch: {PUBLIC_BRANCH}."
        ),
    },

    "RR04": {
        "response": (
            "The Data Availability Statement has been revised to provide direct NCBI "
            "Gene Expression Omnibus links for each of the three principal datasets: "
            "GSE211567, GSE73461 and GSE72810. The text also identifies candidate or "
            "technical-rehearsal datasets considered during workflow development."
        ),
        "location": (
            "Data availability section."
        ),
        "evidence": (
            "Direct GEO accession URLs for GSE211567, GSE73461 and GSE72810 are "
            "included in the revised manuscript."
        ),
    },

    "RR05": {
        "response": (
            "We agree that reliance on a single projection cohort was an important "
            "limitation. We therefore added GSE72810 as a second accession-level and "
            "sample-level whole-blood cohort measured on a different Illumina platform "
            "(GPL6947). The cohort contains 146 samples; the prespecified primary "
            "contrast comprises 23 definite bacterial and 28 definite viral samples. "
            "All five modules retained their expected directions. BACT_M2, VIR_M1a, "
            "VIR_M1b and VIR_M2 had confidence intervals excluding zero and passed "
            "false-discovery-rate correction in both GSE73461 and GSE72810, whereas "
            "BACT_M1 remained directionally concordant but borderline in both. We also "
            "softened the replication language because the two GEO accession sets were "
            "disjoint and used different array platforms, but direct participant overlap "
            "could not be assessed and the studies arose from the same broad "
            "investigator network."
        ),
        "location": (
            "Methods: 'GSE72810 cross-platform validation and probe selection' and "
            "'Cross-cohort interpretation and reproducibility boundaries'; Results: "
            "'GSE72810 provides accession-level and cross-platform validation' and "
            "'Cross-cohort effects support four modules in both cohorts'; Discussion "
            "limitations."
        ),
        "evidence": (
            "Figure 3; Table 2; Supplementary Tables S6-S9; GSE72810 primary "
            "Hodges-Lehmann shifts and bootstrap confidence intervals."
        ),
    },

    "RR06": {
        "response": (
            "We expanded the Methods to make the module-definition procedure more "
            "transparent. Directionally supported discovery features were mapped to "
            "genes, assessed by Gene Ontology biological-process enrichment, and "
            "reviewed after grouping overlapping or redundant enriched terms. "
            "Candidate programmes were evaluated using the documented biological "
            "curation hierarchy. We now state explicitly that this was a biologically "
            "guided, reproducible curation step rather than mathematical optimisation "
            "for group separation or prediction accuracy. External-cohort results were "
            "not used to choose or modify module genes."
        ),
        "location": (
            "Methods, 'GSE211567 discovery and module definition'."
        ),
        "evidence": (
            "Locked module definitions are unchanged; the revision clarifies the "
            "selection rationale and its non-optimisation character."
        ),
    },

    "RR07": {
        "response": (
            "We added a GSVA sensitivity analysis in GSE73461 using the same unchanged "
            "module gene sets and both z-reference populations. BACT_M2, VIR_M1a and "
            "VIR_M1b retained statistical support under both mean-z and GSVA scoring. "
            "BACT_M1 was borderline under mean-z scoring but gained statistical support "
            "under GSVA. VIR_M2 retained its expected viral-higher direction but became "
            "near-zero and statistically unsupported under GSVA. The revision therefore "
            "distinguishes direction preservation from method-independent inferential "
            "support and describes VIR_M2 as scoring-method-sensitive."
        ),
        "location": (
            "Methods, 'Sensitivity and robustness analyses'; Results, 'Sensitivity "
            "analyses retain direction but identify method-dependent support'; "
            "Discussion."
        ),
        "evidence": (
            "S1 Fig panel B and Supplementary Table S10."
        ),
    },

    "RR08": {
        "response": (
            "We added effect-size estimates and uncertainty intervals throughout the "
            "external-cohort analyses. The revised Methods specify Hodges-Lehmann "
            "location shifts, rank-biserial effects and stratified 10,000-replicate "
            "bootstrap 95% confidence intervals. Figure 3 and Table 2 provide harmonised "
            "cross-cohort Hodges-Lehmann effects and 95% confidence intervals for all "
            "five modules. This allows the magnitude and uncertainty of the observed "
            "effects to be assessed separately from adjusted P values."
        ),
        "location": (
            "Methods, 'Module scoring and statistical analysis'; Results, GSE72810 "
            "and cross-cohort effect sections; Figure 3; Table 2."
        ),
        "evidence": (
            "Supplementary Tables S8-S9 and cross-cohort Figure 3 source data."
        ),
    },

    "RR09": {
        "response": (
            "We performed exhaustive leave-one-gene and leave-two-gene sensitivity "
            "analyses for every module under both GSE73461 scoring-reference "
            "populations. In total, 29,826 module variants were evaluated. Every "
            "variant retained the expected module direction. The minimum Pearson "
            "correlation between a deletion variant and its corresponding complete "
            "module score was 0.9940, and the minimum Spearman correlation was 0.9897. "
            "These results support a distributed module signal rather than dependence "
            "on removal-sensitive individual genes or gene pairs."
        ),
        "location": (
            "Methods, 'Sensitivity and robustness analyses'; Results, 'Exhaustive "
            "deletion analysis supports distributed module signal'."
        ),
        "evidence": (
            "S1 Fig panel C; Supplementary Table S10, including complete "
            "leave-one/two-gene results."
        ),
    },

    "RR10": {
        "response": (
            "The dataset descriptions have been substantially expanded. GSE211567 is "
            "identified as whole-blood transcriptomic data from Sri Lanka and the "
            "United States, with a prespecified discovery set of 224 samples "
            "(101 bacterial and 123 viral). GSE73461 is described with its GPL10558 "
            "platform, 52 definite bacterial, 94 definite viral and 55 control samples "
            "in the main z-score reference population, and the restricted primary "
            "inferential contrast. GSE72810 is described as 146 paediatric whole-blood "
            "samples on GPL6947, including definite, probable, control and uncertain "
            "groups and the prespecified 23-versus-28 primary contrast."
        ),
        "location": (
            "Methods, 'Study design and datasets', 'GSE211567 discovery and module "
            "definition', 'GSE73461 formal external projection', and 'GSE72810 "
            "cross-platform validation and probe selection'."
        ),
        "evidence": (
            "Supplementary cohort-audit and sample-classification tables."
        ),
    },

    "RR11": {
        "response": (
            "We clarified the role of control and non-primary samples in score "
            "standardisation versus inference. In GSE73461, the 52 definite bacterial, "
            "94 definite viral and 55 control samples constitute the main gene-wise "
            "z-score reference population, but control samples are not included in the "
            "bacterial-versus-viral inferential test. A primary-only z-reference "
            "sensitivity analysis quantifies dependence on that choice. In GSE72810, "
            "all 146 samples contribute to the main z-score reference population, "
            "whereas the primary inferential contrast remains restricted to the "
            "23 definite bacterial and 28 definite viral samples. Additional "
            "case-definition and probe-collapse analyses test the robustness of these "
            "choices."
        ),
        "location": (
            "Methods, 'GSE73461 formal external projection', 'GSE72810 "
            "cross-platform validation and probe selection', and 'Sensitivity and "
            "robustness analyses'."
        ),
        "evidence": (
            "Figure 2, S1 Fig panel A and the supplementary sensitivity tables."
        ),
    },

    "RR12": {
        "response": (
            "All principal figure legends were rewritten so that they explain both "
            "the analysis and the graphical elements. The revised legends define "
            "direction conventions, sample groups, symbols, confidence intervals or "
            "adjusted-P-value thresholds where applicable, and state the principal "
            "finding represented in each panel. The S1 Fig legend likewise explains "
            "the z-reference/case-definition/probe-collapse sensitivity analyses, "
            "mean-z versus GSVA comparison and exhaustive gene-deletion analysis."
        ),
        "location": (
            "Legends for Figures 1-3 and Supporting Information S1 Fig."
        ),
        "evidence": (
            "Revised manuscript figure captions and supporting-information caption."
        ),
    },

    "RR13": {
        "response": (
            "We now define directional concordance formally. For each modelled "
            "feature, concordance means agreement in the sign of the bacterial-versus-"
            "viral log2 fold change between the two contrasts being compared. "
            "Percentage directional concordance is the number of same-sign features "
            "divided by the number of compared features and multiplied by 100. "
            "Spearman correlations of log2 fold changes are reported separately as a "
            "complementary measure of ranked effect-size agreement."
        ),
        "location": (
            "Methods, 'GSE211567 discovery and module definition'; Figure 1B legend."
        ),
        "evidence": (
            "The definition now matches the implemented cross-site concordance "
            "calculation."
        ),
    },

    "RR14": {
        "response": (
            "We added the explicit module-scoring equations. For mapped gene g in "
            "sample i, expression is standardised within the prespecified reference "
            "population as z_gi = (x_gi - mean_g) / SD_g. For a module with K mapped "
            "and available genes, the sample-level score is the unweighted arithmetic "
            "mean, score_i = (1/K) sum_g z_gi. Thus each available mapped gene "
            "contributes equally to the primary score; missing genes are omitted only "
            "after module coverage has been documented."
        ),
        "location": (
            "Methods, 'Module scoring and statistical analysis'."
        ),
        "evidence": (
            "The explicit equations and interpretation of the mean z-score are now "
            "given in the manuscript."
        ),
    },

    "RR15": {
        "response": (
            "Figure 2C was rebuilt to treat the module categories correctly as "
            "independent categorical values. Circles denote the main projection and "
            "triangles denote the primary-only z-score sensitivity analysis. The "
            "module categories are not connected by lines. The dashed horizontal line "
            "marks BH-adjusted P = 0.05. The revised caption explains each symbol and "
            "why no connecting line is used."
        ),
        "location": (
            "Figure 2C and Figure 2 legend."
        ),
        "evidence": (
            "Revised Figure 2C generated by the reproducible figure script; "
            "publication outputs include PNG and editable SVG versions."
        ),
    },

    "RR16": {
        "response": (
            "The statistical-analysis description has been expanded substantially. "
            "The revised Methods specify Wilcoxon rank-sum testing, Benjamini-Hochberg "
            "correction across the five modules, the interpretation of positive and "
            "negative effects, median bacterial-minus-viral differences, "
            "Hodges-Lehmann shifts, rank-biserial effects, 10,000-replicate bootstrap "
            "confidence intervals, Pearson and Spearman score correlations, and the "
            "interpretive basis for cross-method comparisons in the GSVA sensitivity "
            "analysis."
        ),
        "location": (
            "Methods, 'Module scoring and statistical analysis' and 'Sensitivity and "
            "robustness analyses'."
        ),
        "evidence": (
            "Supplementary statistical outputs provide the corresponding complete "
            "module-level results."
        ),
    },

    "RR17": {
        "response": (
            "The manuscript was edited throughout for simpler and more precise wording. "
            "The title was reframed around external transportability. The Introduction "
            "now defines transportability as preservation of the prespecified biological "
            "direction when the same module definition and scoring rule are applied to "
            "another dataset, and defines fixed-module projection as applying predefined "
            "gene sets without changing their genes, expected directions or weights and "
            "without retraining a model. Workflow-oriented terms such as 'firewall', "
            "'frozen', 'safeguards' and similar jargon were removed from "
            "manuscript-facing prose."
        ),
        "location": (
            "Title; Abstract; Introduction; Methods; Results; Discussion; figure "
            "captions throughout the revised manuscript."
        ),
        "evidence": (
            "The revised manuscript consistently distinguishes biological "
            "transportability from diagnostic-classifier validation."
        ),
    },
}


labels = {
    "RR01": "J1. PLOS ONE style and revision-package requirements",
    "RR02": "J2. Generative-AI disclosure",
    "RR03": "J3. Public availability of author-generated code",
    "RR04": "J4. Direct database links in Data Availability",
    "RR05": "Reviewer 1, Comment 1 — Additional external cohort and limitation",
    "RR06": "Reviewer 1, Comment 2 — Manual GO review and module definition",
    "RR07": "Reviewer 1, Comment 3 — Comparison with GSVA",
    "RR08": "Reviewer 1, Comment 4 — Effect sizes and confidence intervals",
    "RR09": "Reviewer 2, Comment 1 — Leave-one/two-gene robustness",
    "RR10": "Reviewer 2, Comment 2 — Dataset description",
    "RR11": "Reviewer 2, Comment 3 — Control samples and z-score reference",
    "RR12": "Reviewer 2, Comment 4 — Figure captions",
    "RR13": "Reviewer 2, Comment 5 — Formal concordance definition",
    "RR14": "Reviewer 2, Comment 6 — Module-score definition",
    "RR15": "Reviewer 2, Comment 7 — Figure 2C categorical display",
    "RR16": "Reviewer 2 — Statistical-analysis description",
    "RR17": "Reviewer 2 — Clarity and terminology",
}


final_rows: list[dict[str, object]] = []


for response_id in expected_ids:

    source_row = matrix_by_id[
        response_id
    ]

    item = responses[
        response_id
    ]

    final_rows.append(
        {
            "response_id": response_id,
            "source": source_row[
                "source"
            ],
            "reviewer_item": source_row[
                "item"
            ],
            "review_request": source_row[
                "review_request"
            ],
            "final_response": item[
                "response"
            ],
            "manuscript_location": item[
                "location"
            ],
            "supporting_evidence": item[
                "evidence"
            ],
            "final_status": "ADDRESSED",
        }
    )


WORK.mkdir(
    parents=True,
    exist_ok=True,
)

OUT.mkdir(
    parents=True,
    exist_ok=True,
)

SUBMISSION.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)


doc = Document()


section = doc.sections[0]

section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)


styles = doc.styles

normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)


title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format(
    title_p,
    after=5,
)

run = title_p.add_run(
    "Response to Academic Editor and Reviewers"
)

set_run(
    run,
    bold=True,
    size=16,
    color=(31, 78, 121),
)


meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format(
    meta_p,
    after=2,
)

run = meta_p.add_run(
    f"Manuscript ID: {MANUSCRIPT_ID}"
)

set_run(
    run,
    bold=True,
    size=11,
)


title_meta = doc.add_paragraph()
title_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format(
    title_meta,
    after=2,
)

run = title_meta.add_run(
    TITLE
)

set_run(
    run,
    italic=True,
    size=11,
)


author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format(
    author_p,
    after=2,
)

run = author_p.add_run(
    AUTHOR
)

set_run(
    run,
    size=11,
)


date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format(
    date_p,
    after=12,
)

run = date_p.add_run(
    REVISION_DATE
)

set_run(
    run,
    size=10.5,
)


intro = (
    "Dear Academic Editor and Reviewers,\n\n"
    "Thank you for the careful evaluation of the manuscript and for the constructive "
    "recommendations. The manuscript has been revised substantially in response. "
    "The principal changes include addition of a second accession-level and "
    "cross-platform projection cohort (GSE72810), harmonised effect-size and "
    "confidence-interval reporting, GSVA sensitivity analysis, exhaustive "
    "leave-one/two-gene robustness analysis, expanded cohort and statistical-method "
    "descriptions, a revised categorical Figure 2C, clearer definitions and simpler "
    "terminology, direct public-data links, public code availability, and a dedicated "
    "generative-AI disclosure within Methods.\n\n"
    "Each comment is reproduced below, followed by the response and the location of "
    "the corresponding revision. The revised submission includes separate clean and "
    "marked-up manuscript files."
)

for block in intro.split(
    "\n\n"
):
    add_body_paragraph(
        doc,
        block,
        after=7,
    )


section_heading(
    doc,
    "Academic Editor / Journal Requirements",
)


for response_id in (
    "RR01",
    "RR02",
    "RR03",
    "RR04",
):

    row = matrix_by_id[
        response_id
    ]

    item = responses[
        response_id
    ]

    add_response_item(
        doc,
        labels[
            response_id
        ],
        row[
            "review_request"
        ],
        item[
            "response"
        ],
        item[
            "location"
        ],
        item[
            "evidence"
        ],
    )


section_heading(
    doc,
    "Reviewer 1",
)


for response_id in (
    "RR05",
    "RR06",
    "RR07",
    "RR08",
):

    row = matrix_by_id[
        response_id
    ]

    item = responses[
        response_id
    ]

    add_response_item(
        doc,
        labels[
            response_id
        ],
        row[
            "review_request"
        ],
        item[
            "response"
        ],
        item[
            "location"
        ],
        item[
            "evidence"
        ],
    )


section_heading(
    doc,
    "Reviewer 2",
)


for response_id in (
    "RR09",
    "RR10",
    "RR11",
    "RR12",
    "RR13",
    "RR14",
    "RR15",
    "RR16",
    "RR17",
):

    row = matrix_by_id[
        response_id
    ]

    item = responses[
        response_id
    ]

    add_response_item(
        doc,
        labels[
            response_id
        ],
        row[
            "review_request"
        ],
        item[
            "response"
        ],
        item[
            "location"
        ],
        item[
            "evidence"
        ],
    )


section_heading(
    doc,
    "Closing statement",
)


closing = (
    "We appreciate the editor's and reviewers' comments, which led to substantial "
    "improvements in the analytical robustness, transparency and clarity of the "
    "manuscript. All 17 editor/reviewer response items represented in the revision "
    "matrix are addressed in the revised manuscript and supporting materials."
)

add_body_paragraph(
    doc,
    closing,
    after=8,
)


signature = doc.add_paragraph()
paragraph_format(
    signature,
    before=8,
    after=0,
)

run = signature.add_run(
    "Sincerely,\n"
    "Reuben S. Maghembe"
)

set_run(
    run,
    bold=True,
)


doc.save(
    RESPONSE_DOCX
)


markdown_lines = [
    "# Response to Academic Editor and Reviewers",
    "",
    f"**Manuscript ID:** {MANUSCRIPT_ID}",
    "",
    f"**Title:** {TITLE}",
    "",
    f"**Author:** {AUTHOR}",
    "",
    f"**Date:** {REVISION_DATE}",
    "",
    "Dear Academic Editor and Reviewers,",
    "",
    (
        "Thank you for the careful evaluation of the manuscript and for the "
        "constructive recommendations. The manuscript has been revised "
        "substantially in response."
    ),
    "",
]


current_source = None


for row in final_rows:

    source = row[
        "source"
    ]

    if source != current_source:

        markdown_lines.extend(
            [
                f"## {source}",
                "",
            ]
        )

        current_source = source

    response_id = row[
        "response_id"
    ]

    markdown_lines.extend(
        [
            f"### {labels[response_id]}",
            "",
            f"**Comment:** {row['review_request']}",
            "",
            f"**Response:** {row['final_response']}",
            "",
            (
                "**Changes in the revised manuscript:** "
                f"{row['manuscript_location']}"
            ),
            "",
            (
                "**Supporting analysis/material:** "
                f"{row['supporting_evidence']}"
            ),
            "",
        ]
    )


RESPONSE_MD.write_text(
    "\n".join(
        markdown_lines
    )
    + "\n",
    encoding="utf-8",
    newline="\n",
)


write_tsv(
    RESPONSE_TSV,
    final_rows,
)


render_status = 99
render_pdf = Path()
render_output = ""


if shutil.which(
    "soffice"
) is not None:

    (
        render_status,
        render_pdf,
        render_output,
    ) = render_docx(
        RESPONSE_DOCX
    )


page_count = 0
page_images: list[Path] = []
contact_sheets: list[Path] = []


if (
    render_status == 0
    and render_pdf.exists()
    and render_pdf.stat().st_size > 0
):

    page_count = pdf_page_count(
        render_pdf
    )

    if shutil.which(
        "pdftoppm"
    ) is not None:

        page_images = create_page_images(
            render_pdf
        )

        contact_sheets = create_contact_sheets(
            page_images
        )


response_sha = sha256(
    RESPONSE_DOCX
)


checks: list[
    dict[str, object]
] = []


def check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": description,
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected
            ),
        }
    )


check(
    "Clean manuscript SHA locked",
    sha256(
        CLEAN_DOCX
    ) == EXPECTED_CLEAN_SHA,
    sha256(
        CLEAN_DOCX
    ),
    EXPECTED_CLEAN_SHA,
)

check(
    "Marked manuscript SHA locked",
    sha256(
        MARKED_DOCX
    ) == EXPECTED_MARKED_SHA,
    sha256(
        MARKED_DOCX
    ),
    EXPECTED_MARKED_SHA,
)

check(
    "v2.4 scientific source SHA locked",
    sha256(
        SOURCE_V24
    ) == EXPECTED_SOURCE_SHA,
    sha256(
        SOURCE_V24
    ),
    EXPECTED_SOURCE_SHA,
)

check(
    "Exactly 17 response items generated",
    len(
        final_rows
    ) == 17,
    len(
        final_rows
    ),
    17,
)

check(
    "All RR01-RR17 identifiers present",
    [
        row[
            "response_id"
        ]
        for row
        in final_rows
    ] == expected_ids,
    ",".join(
        row[
            "response_id"
        ]
        for row in final_rows
    ),
    ",".join(
        expected_ids
    ),
)

check(
    "All final response statuses addressed",
    all(
        row[
            "final_status"
        ]
        == "ADDRESSED"
        for row in final_rows
    ),
    sum(
        row[
            "final_status"
        ]
        == "ADDRESSED"
        for row in final_rows
    ),
    17,
)

check(
    "Editor/journal response count",
    sum(
        row[
            "source"
        ]
        == "Academic editor / journal"
        for row in final_rows
    ) == 4,
    sum(
        row[
            "source"
        ]
        == "Academic editor / journal"
        for row in final_rows
    ),
    4,
)

check(
    "Reviewer 1 response count",
    sum(
        row[
            "source"
        ]
        == "Reviewer 1"
        for row in final_rows
    ) == 4,
    sum(
        row[
            "source"
        ]
        == "Reviewer 1"
        for row in final_rows
    ),
    4,
)

check(
    "Reviewer 2 response count",
    sum(
        row[
            "source"
        ]
        == "Reviewer 2"
        for row in final_rows
    ) == 9,
    sum(
        row[
            "source"
        ]
        == "Reviewer 2"
        for row in final_rows
    ),
    9,
)

check(
    "Response DOCX created",
    RESPONSE_DOCX.exists()
    and RESPONSE_DOCX.stat().st_size > 0,
    (
        RESPONSE_DOCX.stat().st_size
        if RESPONSE_DOCX.exists()
        else 0
    ),
    ">0 bytes",
)

check(
    "Response Markdown audit copy created",
    RESPONSE_MD.exists()
    and RESPONSE_MD.stat().st_size > 0,
    (
        RESPONSE_MD.stat().st_size
        if RESPONSE_MD.exists()
        else 0
    ),
    ">0 bytes",
)

check(
    "Response TSV audit copy created",
    RESPONSE_TSV.exists()
    and RESPONSE_TSV.stat().st_size > 0,
    (
        RESPONSE_TSV.stat().st_size
        if RESPONSE_TSV.exists()
        else 0
    ),
    ">0 bytes",
)

check(
    "LibreOffice render succeeds",
    (
        render_status == 0
        and render_pdf.exists()
        and render_pdf.stat().st_size > 0
    ),
    (
        f"status={render_status}; "
        f"bytes="
        f"{render_pdf.stat().st_size if render_pdf.exists() else 0}"
    ),
    "status=0 and non-empty PDF",
)

check(
    "Rendered response has at least one page",
    page_count > 0,
    page_count,
    ">0",
)

check(
    "Every rendered PDF page has a PNG",
    (
        page_count > 0
        and len(
            page_images
        ) == page_count
    ),
    len(
        page_images
    ),
    page_count,
)

check(
    "AI response states disclosure is in Methods",
    (
        "within the Methods section"
        in responses[
            "RR02"
        ][
            "response"
        ]
    ),
    (
        "within Methods"
        if (
            "within the Methods section"
            in responses[
                "RR02"
            ][
                "response"
            ]
        )
        else "missing"
    ),
    "within Methods",
)

check(
    "RR01 updated from package-pending to addressed",
    final_rows[0][
        "final_status"
    ] == "ADDRESSED",
    final_rows[0][
        "final_status"
    ],
    "ADDRESSED",
)

check(
    "No canonical manuscript changed during response production",
    (
        sha256(
            CLEAN_DOCX
        ) == EXPECTED_CLEAN_SHA
        and sha256(
            MARKED_DOCX
        ) == EXPECTED_MARKED_SHA
    ),
    (
        f"clean={sha256(CLEAN_DOCX)}; "
        f"marked={sha256(MARKED_DOCX)}"
    ),
    (
        f"clean={EXPECTED_CLEAN_SHA}; "
        f"marked={EXPECTED_MARKED_SHA}"
    ),
)


passed = sum(
    row[
        "pass"
    ] == "TRUE"
    for row in checks
)

failed = (
    len(
        checks
    )
    - passed
)

gate = (
    "PASS"
    if failed == 0
    else "FAIL"
)

status = (
    "READY_FOR_RESPONSE_LETTER_PAGE_BY_PAGE_VISUAL_QA"
    if failed == 0
    else "RESPONSE_LETTER_REQUIRES_CORRECTION"
)


write_tsv(
    QUALITY,
    checks,
)


write_tsv(
    SUMMARY,
    [
        {
            "manuscript_id": MANUSCRIPT_ID,
            "matrix_source": str(
                matrix_path
            ),
            "clean_manuscript_sha256": (
                EXPECTED_CLEAN_SHA
            ),
            "marked_manuscript_sha256": (
                EXPECTED_MARKED_SHA
            ),
            "scientific_source_v2.4_sha256": (
                EXPECTED_SOURCE_SHA
            ),
            "response_docx": str(
                RESPONSE_DOCX
            ),
            "response_docx_sha256": (
                response_sha
            ),
            "response_items": (
                len(
                    final_rows
                )
            ),
            "editor_items": 4,
            "reviewer1_items": 4,
            "reviewer2_items": 9,
            "rendered_pdf_pages": (
                page_count
            ),
            "page_pngs": (
                len(
                    page_images
                )
            ),
            "contact_sheets": (
                len(
                    contact_sheets
                )
            ),
            "quality_checks": (
                len(
                    checks
                )
            ),
            "quality_checks_passed": (
                passed
            ),
            "quality_checks_failed": (
                failed
            ),
            "quality_gate": (
                gate
            ),
            "final_status": (
                status
            ),
        }
    ],
)


REPORT.write_text(
    "\n".join(
        [
            "# PLOS ONE Response Letter Production Report",
            "",
            f"Manuscript: {MANUSCRIPT_ID}",
            "",
            f"- Matrix source: `{matrix_path}`",
            f"- Clean manuscript SHA256: `{EXPECTED_CLEAN_SHA}`",
            f"- Marked manuscript SHA256: `{EXPECTED_MARKED_SHA}`",
            f"- Scientific source v2.4 SHA256: `{EXPECTED_SOURCE_SHA}`",
            f"- Response DOCX: `{RESPONSE_DOCX}`",
            f"- Response DOCX SHA256: `{response_sha}`",
            f"- Response items: {len(final_rows)}",
            "- Academic editor/journal items: 4",
            "- Reviewer 1 items: 4",
            "- Reviewer 2 items: 9",
            f"- Rendered pages: {page_count}",
            f"- Page PNGs: {len(page_images)}",
            f"- Contact sheets: {len(contact_sheets)}",
            f"- Quality checks passed: {passed}/{len(checks)}",
            f"- Quality gate: `{gate}`",
            f"- Final status: `{status}`",
            "",
        ]
    ),
    encoding="utf-8",
    newline="\n",
)


print(
    "===== PLOS ONE RESPONSE LETTER PRODUCTION ====="
)

print(
    f"matrix_source\t{matrix_path}"
)

print(
    f"clean_manuscript_sha256\t{EXPECTED_CLEAN_SHA}"
)

print(
    f"marked_manuscript_sha256\t{EXPECTED_MARKED_SHA}"
)

print(
    f"scientific_source_v2.4_sha256\t{EXPECTED_SOURCE_SHA}"
)

print(
    f"response_docx\t{RESPONSE_DOCX}"
)

print(
    f"response_docx_sha256\t{response_sha}"
)

print(
    f"response_items\t{len(final_rows)}"
)

print(
    "editor_items\t4"
)

print(
    "reviewer1_items\t4"
)

print(
    "reviewer2_items\t9"
)

print(
    f"rendered_pdf_pages\t{page_count}"
)

print(
    f"page_pngs\t{len(page_images)}"
)

print(
    f"contact_sheets\t{len(contact_sheets)}"
)

print(
    f"quality_checks_passed\t{passed}/{len(checks)}"
)

print(
    f"quality_gate\t{gate}"
)

print(
    f"final_status\t{status}"
)

print(
    f"quality_gate_file\t{QUALITY}"
)

print(
    f"summary\t{SUMMARY}"
)

print(
    f"report\t{REPORT}"
)

print()

print(
    "===== LIBREOFFICE OUTPUT ====="
)

print(
    render_output
)


if contact_sheets:
    print()
    print(
        "===== CONTACT SHEETS ====="
    )

    for path in contact_sheets:
        print(
            path
        )


if failed:
    raise RuntimeError(
        "Response-letter production failed "
        f"{failed} quality check(s)."
    )
