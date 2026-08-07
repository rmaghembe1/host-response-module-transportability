#!/usr/bin/env python3

"""
64_build_plosone_clean_manuscript_docx.py

CORRECTED_BUILD_V2

Build the clean PLOS ONE revision manuscript from the scientifically
locked v2.3 Markdown source.

Corrections in this implementation
----------------------------------
1. Only the FRONT-MATTER "## Title" heading is removed.
   Table 1 and Table 2 local "## Title" headings are preserved.

2. Empty Word comments infrastructure created during DOCX conversion is
   distinguished from real comments. Empty comments.xml/package metadata
   is removed only after confirming that there are zero actual comments
   and zero comment anchors.

3. Main figures and tables are audited for PLOS read-order placement but
   are NOT moved in this phase. Placement correction follows visual QA.

Authoritative source
--------------------
docs/complete_manuscript_draft_v2.3_submission_candidate_metadata_restored.md

Locked source SHA256
--------------------
f3b61e6ddb9f5d38c6211c6cfe0d8694e6ca3b761d52a3245d58df844ab5b2ae
"""

from __future__ import annotations

import csv
import hashlib
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# =============================================================================
# Locked manuscript
# =============================================================================

SOURCE = Path(
    "docs/complete_manuscript_draft_v2.3_submission_candidate_metadata_restored.md"
)

EXPECTED_SOURCE_SHA256 = (
    "f3b61e6ddb9f5d38c6211c6cfe0d8694e6ca3b761d52a3245d58df844ab5b2ae"
)

ARTICLE_TITLE = (
    "External transportability of bacterial- and viral-associated "
    "host-response modules across public transcriptomic cohorts"
)

SHORT_TITLE = "Transportable infection modules"

AUTHOR_LINE = "Reuben S. Maghembe¹˒²*"

AFFILIATION_1 = (
    "¹Department of Microbiology and Parasitology, Faculty of Medicine, "
    "St. Francis University College of Health and Allied Sciences "
    "(SFUCHAS), Ifakara, Tanzania"
)

AFFILIATION_2 = (
    "²Department of Omics and Computational Biology, AfroBiomics Co. Ltd., "
    "Kivukoni, Bridge Street, Dar es Salaam, Tanzania"
)

EMAIL_LINE = (
    "Email: rmaghembe@gmail.com; rmaghembe@sfuchas.ac.tz"
)


# =============================================================================
# Paths
# =============================================================================

WORK_DIR = Path(
    "work/plosone_revision_round1_2026/phaseR1E13_clean_docx"
)

NORMALIZED_MD = (
    WORK_DIR / "PONE-D-26-30583_clean_v2.3_normalized_source.md"
)

INTERMEDIATE_DOCX = (
    WORK_DIR / "PONE-D-26-30583_clean_v2.3_pandoc_intermediate.docx"
)

FINAL_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

OUT_DIR = Path(
    "results/revision_round1/plosone_clean_manuscript_docx_v2.3"
)

QUALITY_GATE = (
    OUT_DIR / "PLOS_ONE_clean_manuscript_docx_v2.3_quality_gate.tsv"
)

QUALITY_SUMMARY = (
    OUT_DIR / "PLOS_ONE_clean_manuscript_docx_v2.3_summary.tsv"
)

STRUCTURE_AUDIT = (
    OUT_DIR / "PLOS_ONE_clean_manuscript_docx_v2.3_structure_audit.tsv"
)

TRANSFORM_MANIFEST = (
    OUT_DIR / "PLOS_ONE_clean_manuscript_docx_v2.3_format_transformations.tsv"
)

COMMENT_AUDIT = (
    OUT_DIR / "PLOS_ONE_clean_manuscript_docx_v2.3_comment_audit.tsv"
)

REPORT = Path(
    "docs/revision_round1/PLOS_ONE_clean_manuscript_docx_v2.3_report.md"
)


# =============================================================================
# Namespaces
# =============================================================================

W_NS = (
    "http://schemas.openxmlformats.org/"
    "wordprocessingml/2006/main"
)


# =============================================================================
# Utilities
# =============================================================================

def fail(message: str) -> None:
    print(f"ERROR: {message}", file=sys.stderr)
    sys.exit(1)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()

    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)

    return digest.hexdigest()


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:

    if not rows:
        fail(f"No rows supplied for {path}")

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    fields = list(rows[0].keys())

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=fields,
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()
        writer.writerows(rows)


def add_transform(
    rows: list[dict[str, str]],
    action: str,
    detail: str,
) -> None:

    rows.append(
        {
            "transformation_id": f"T{len(rows) + 1:02d}",
            "action": action,
            "detail": detail,
        }
    )


def heading_level(line: str) -> int | None:

    match = re.match(
        r"^(#{1,6})\s+",
        line,
    )

    if not match:
        return None

    return len(match.group(1))


# =============================================================================
# Normalize Markdown
# =============================================================================

def normalize_markdown(
    source_text: str,
) -> tuple[str, list[dict[str, str]]]:

    lines = source_text.splitlines()

    if not lines:
        fail("Source manuscript is empty.")

    if not lines[0].startswith(
        "# Complete Manuscript Draft v2.3"
    ):
        fail(
            "Expected v2.3 internal working heading "
            "was not found on line 1."
        )

    transformations: list[dict[str, str]] = []
    output: list[str] = []

    add_transform(
        transformations,
        "REMOVE_INTERNAL_VERSION_HEADING",
        lines[0],
    )

    abstract_index = None

    for index, line in enumerate(lines):
        if re.match(
            r"^##\s+Abstract\s*$",
            line.strip(),
            flags=re.IGNORECASE,
        ):
            abstract_index = index
            break

    if abstract_index is None:
        fail("Could not locate ## Abstract.")

    i = 1

    while i < len(lines) and not lines[i].strip():
        i += 1

    if i < len(lines) and lines[i].strip() == "---":
        add_transform(
            transformations,
            "REMOVE_INTERNAL_HORIZONTAL_RULE",
            "---",
        )
        i += 1

    article_category_removed = False
    front_title_removed = False
    short_title_normalized = False
    author_label_removed = False
    post_abstract_title_count = 0

    while i < len(lines):

        line = lines[i]
        stripped = line.strip()

        # -----------------------------------------------------------------
        # Front-matter Article category block only.
        # -----------------------------------------------------------------

        if (
            i < abstract_index
            and stripped == "## Article category"
        ):
            add_transform(
                transformations,
                "REMOVE_INTERNAL_ARTICLE_CATEGORY_BLOCK",
                "Article category / Original Article",
            )

            article_category_removed = True

            i += 1

            while i < len(lines):
                if heading_level(lines[i]) is not None:
                    break
                i += 1

            continue

        # -----------------------------------------------------------------
        # Only the first/front-matter Title heading is an internal label.
        # Later ## Title headings belong to Table 1 and Table 2.
        # -----------------------------------------------------------------

        if stripped == "## Title":

            if (
                i < abstract_index
                and not front_title_removed
            ):
                add_transform(
                    transformations,
                    "REMOVE_FRONT_MATTER_TITLE_LABEL",
                    "## Title",
                )

                front_title_removed = True
                i += 1
                continue

            post_abstract_title_count += 1

        # -----------------------------------------------------------------
        # Front-matter short title.
        # -----------------------------------------------------------------

        if (
            i < abstract_index
            and stripped == "## Short title"
        ):

            j = i + 1

            while (
                j < len(lines)
                and not lines[j].strip()
            ):
                j += 1

            if j >= len(lines):
                fail(
                    "Short-title heading has no value."
                )

            value = lines[j].strip()

            if value != SHORT_TITLE:
                fail(
                    "Unexpected short title: "
                    + value
                )

            output.append(
                f"**Short title:** {value}"
            )
            output.append("")

            add_transform(
                transformations,
                "NORMALIZE_FRONT_MATTER_SHORT_TITLE",
                value,
            )

            short_title_normalized = True

            i = j + 1
            continue

        # -----------------------------------------------------------------
        # Front-matter author label only.
        # -----------------------------------------------------------------

        if (
            i < abstract_index
            and stripped == "## Author and affiliations"
        ):
            add_transform(
                transformations,
                "REMOVE_FRONT_MATTER_AUTHOR_SECTION_LABEL",
                "## Author and affiliations",
            )

            author_label_removed = True

            i += 1
            continue

        # -----------------------------------------------------------------
        # Promote remaining headings because the internal manuscript-root
        # heading has been removed.
        # -----------------------------------------------------------------

        match = re.match(
            r"^(#{2,6})\s+(.*)$",
            line,
        )

        if match:
            hashes = match.group(1)
            title = match.group(2)

            output.append(
                "#" * (len(hashes) - 1)
                + " "
                + title
            )

            i += 1
            continue

        output.append(line)
        i += 1

    if not article_category_removed:
        fail(
            "Front-matter Article category block "
            "was not removed."
        )

    if not front_title_removed:
        fail(
            "Front-matter Title label "
            "was not removed."
        )

    if not short_title_normalized:
        fail(
            "Front-matter Short title "
            "was not normalized."
        )

    if not author_label_removed:
        fail(
            "Front-matter author label "
            "was not removed."
        )

    if post_abstract_title_count != 2:
        fail(
            "Expected exactly two post-Abstract "
            "## Title headings; observed "
            f"{post_abstract_title_count}."
        )

    add_transform(
        transformations,
        "PRESERVE_TABLE_LOCAL_TITLE_HEADINGS",
        (
            f"{post_abstract_title_count} "
            "post-Abstract Title headings preserved"
        ),
    )

    add_transform(
        transformations,
        "PROMOTE_REMAINING_HEADING_LEVELS",
        "##->#, ###->##, ####->###, etc.",
    )

    normalized = (
        "\n".join(output).strip()
        + "\n"
    )

    if ARTICLE_TITLE not in normalized:
        fail(
            "Article title missing after normalization."
        )

    if AUTHOR_LINE not in normalized:
        fail(
            "Author line missing after normalization."
        )

    return normalized, transformations


# =============================================================================
# Word helpers
# =============================================================================

def set_run_font(
    run,
    name: str,
    size: float | None = None,
) -> None:

    run.font.name = name

    if size is not None:
        run.font.size = Pt(size)

    rpr = run._element.get_or_add_rPr()

    rfonts = rpr.find(
        qn("w:rFonts")
    )

    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)

    for attr in (
        "w:ascii",
        "w:hAnsi",
        "w:eastAsia",
        "w:cs",
    ):
        rfonts.set(
            qn(attr),
            name,
        )


def add_page_number(paragraph) -> None:

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(
        qn("w:fldCharType"),
        "begin",
    )

    instruction = OxmlElement("w:instrText")
    instruction.set(
        qn("xml:space"),
        "preserve",
    )
    instruction.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(
        qn("w:fldCharType"),
        "separate",
    )

    text = OxmlElement("w:t")
    text.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(
        qn("w:fldCharType"),
        "end",
    )

    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)

    set_run_font(
        run,
        "Times New Roman",
        10,
    )


def set_line_numbering(section) -> None:

    sect_pr = section._sectPr

    previous = sect_pr.find(
        qn("w:lnNumType")
    )

    if previous is not None:
        sect_pr.remove(previous)

    element = OxmlElement(
        "w:lnNumType"
    )

    element.set(
        qn("w:countBy"),
        "1",
    )

    element.set(
        qn("w:start"),
        "1",
    )

    element.set(
        qn("w:restart"),
        "continuous",
    )

    element.set(
        qn("w:distance"),
        "360",
    )

    sect_pr.append(element)


def force_single_column(section) -> None:

    sect_pr = section._sectPr

    cols = sect_pr.find(
        qn("w:cols")
    )

    if cols is None:
        cols = OxmlElement(
            "w:cols"
        )
        sect_pr.append(cols)

    cols.set(
        qn("w:num"),
        "1",
    )


def format_docx(
    source_docx: Path,
    output_docx: Path,
) -> None:

    doc = Document(source_docx)

    doc.core_properties.title = (
        ARTICLE_TITLE
    )

    doc.core_properties.subject = (
        "PLOS ONE revised manuscript PONE-D-26-30583"
    )

    doc.core_properties.author = (
        "Reuben S. Maghembe"
    )

    for section in doc.sections:

        section.orientation = (
            WD_ORIENT.PORTRAIT
        )

        section.page_width = (
            Inches(8.5)
        )

        section.page_height = (
            Inches(11)
        )

        section.top_margin = (
            Inches(1.0)
        )

        section.bottom_margin = (
            Inches(1.0)
        )

        section.left_margin = (
            Inches(1.0)
        )

        section.right_margin = (
            Inches(1.0)
        )

        section.header_distance = (
            Inches(0.5)
        )

        section.footer_distance = (
            Inches(0.5)
        )

        force_single_column(section)
        set_line_numbering(section)

        footer = section.footer

        paragraph = (
            footer.paragraphs[0]
            if footer.paragraphs
            else footer.add_paragraph()
        )

        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)

        add_page_number(paragraph)

    normal = doc.styles["Normal"]

    normal.font.name = (
        "Times New Roman"
    )

    normal.font.size = Pt(12)

    normal.paragraph_format.line_spacing_rule = (
        WD_LINE_SPACING.DOUBLE
    )

    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name in (
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ):

        if style_name not in doc.styles:
            continue

        style = doc.styles[style_name]

        style.font.name = (
            "Times New Roman"
        )

        style.font.size = Pt(12)
        style.font.bold = True

        style.paragraph_format.line_spacing_rule = (
            WD_LINE_SPACING.DOUBLE
        )

        style.paragraph_format.space_before = (
            Pt(12)
        )

        style.paragraph_format.space_after = (
            Pt(0)
        )

        style.paragraph_format.keep_with_next = (
            True
        )

    if "Title" in doc.styles:

        style = doc.styles["Title"]

        style.font.name = (
            "Times New Roman"
        )

        style.font.size = Pt(14)
        style.font.bold = True

        style.paragraph_format.line_spacing = (
            1.15
        )

        style.paragraph_format.space_after = (
            Pt(12)
        )

    title_found = False
    author_found = False

    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()

        paragraph.paragraph_format.widow_control = (
            True
        )

        if text == ARTICLE_TITLE:

            title_found = True

            paragraph.style = (
                doc.styles["Title"]
            )

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            paragraph.paragraph_format.keep_with_next = (
                True
            )

            for run in paragraph.runs:
                set_run_font(
                    run,
                    "Times New Roman",
                    14,
                )
                run.font.bold = True

            continue

        if text == AUTHOR_LINE:

            author_found = True

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

        if paragraph.style.name in {
            "Heading 1",
            "Heading 2",
            "Heading 3",
        }:

            paragraph.paragraph_format.keep_with_next = (
                True
            )

            paragraph.paragraph_format.line_spacing_rule = (
                WD_LINE_SPACING.DOUBLE
            )

            for run in paragraph.runs:
                set_run_font(
                    run,
                    "Times New Roman",
                    12,
                )

            continue

        paragraph.paragraph_format.line_spacing_rule = (
            WD_LINE_SPACING.DOUBLE
        )

        paragraph.paragraph_format.space_before = (
            Pt(0)
        )

        paragraph.paragraph_format.space_after = (
            Pt(0)
        )

        for run in paragraph.runs:
            set_run_font(
                run,
                "Times New Roman",
                12,
            )

    if not title_found:
        fail("Article title was not found in DOCX.")

    if not author_found:
        fail("Author line was not found in DOCX.")

    for table in doc.tables:

        table.autofit = True

        for row_index, row in enumerate(
            table.rows
        ):

            for cell in row.cells:

                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

                for paragraph in cell.paragraphs:

                    paragraph.paragraph_format.line_spacing = (
                        1.0
                    )

                    paragraph.paragraph_format.space_before = (
                        Pt(0)
                    )

                    paragraph.paragraph_format.space_after = (
                        Pt(0)
                    )

                    for run in paragraph.runs:

                        set_run_font(
                            run,
                            "Times New Roman",
                            9.5,
                        )

                        if row_index == 0:
                            run.font.bold = True

    output_docx.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    doc.save(output_docx)


# =============================================================================
# Comments package
# =============================================================================

def inspect_comments(
    path: Path,
) -> dict[str, object]:

    related_parts: list[str] = []
    comment_count = 0
    range_start = 0
    range_end = 0
    reference_count = 0

    with zipfile.ZipFile(path, "r") as archive:

        names = archive.namelist()

        related_parts = sorted(
            name
            for name in names
            if (
                "comment" in name.lower()
                or "people" in name.lower()
                or "person" in name.lower()
            )
        )

        if "word/comments.xml" in names:

            root = ET.fromstring(
                archive.read(
                    "word/comments.xml"
                )
            )

            comment_count = len(
                root.findall(
                    f"{{{W_NS}}}comment"
                )
            )

        for name in names:

            if (
                not name.startswith("word/")
                or not name.endswith(".xml")
                or name.startswith("word/comments")
            ):
                continue

            text = archive.read(name).decode(
                "utf-8",
                errors="replace",
            )

            range_start += len(
                re.findall(
                    r"<w:commentRangeStart\b",
                    text,
                )
            )

            range_end += len(
                re.findall(
                    r"<w:commentRangeEnd\b",
                    text,
                )
            )

            reference_count += len(
                re.findall(
                    r"<w:commentReference\b",
                    text,
                )
            )

    return {
        "comment_related_parts": related_parts,
        "comment_count": comment_count,
        "comment_range_start": range_start,
        "comment_range_end": range_end,
        "comment_reference": reference_count,
    }


def strip_empty_comment_package(
    path: Path,
) -> int:

    state = inspect_comments(path)

    actual_markers = (
        int(state["comment_count"])
        + int(state["comment_range_start"])
        + int(state["comment_range_end"])
        + int(state["comment_reference"])
    )

    if actual_markers != 0:
        fail(
            "Actual Word comments/comment anchors "
            "were detected; automatic cleanup refused."
        )

    related = list(
        state["comment_related_parts"]
    )

    if not related:
        return 0

    temp = path.with_name(
        path.stem
        + ".comment_cleanup.tmp.docx"
    )

    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(
            temp,
            "w",
            compression=zipfile.ZIP_DEFLATED,
        ) as zout:

            for item in zin.infolist():

                name = item.filename
                lower_name = name.lower()

                if (
                    lower_name.startswith(
                        "word/comments"
                    )
                    or lower_name.startswith(
                        "word/people"
                    )
                    or lower_name.startswith(
                        "word/person"
                    )
                ):
                    continue

                data = zin.read(name)

                if name == "[Content_Types].xml":

                    root = ET.fromstring(data)

                    for child in list(root):

                        part_name = (
                            child.attrib.get(
                                "PartName",
                                "",
                            ).lower()
                        )

                        content_type = (
                            child.attrib.get(
                                "ContentType",
                                "",
                            ).lower()
                        )

                        if (
                            "comment" in part_name
                            or "comment" in content_type
                            or "/people" in part_name
                            or "/person" in part_name
                        ):
                            root.remove(child)

                    data = ET.tostring(
                        root,
                        encoding="utf-8",
                        xml_declaration=True,
                    )

                elif name.endswith(
                    ".rels"
                ):

                    try:
                        root = ET.fromstring(data)
                    except ET.ParseError:
                        root = None

                    if root is not None:

                        changed = False

                        for child in list(root):

                            rel_type = (
                                child.attrib.get(
                                    "Type",
                                    "",
                                ).lower()
                            )

                            target = (
                                child.attrib.get(
                                    "Target",
                                    "",
                                ).lower()
                            )

                            if (
                                "comment" in rel_type
                                or "comment" in target
                                or "people" in rel_type
                                or "people" in target
                                or "person" in rel_type
                                or "person" in target
                            ):
                                root.remove(child)
                                changed = True

                        if changed:
                            data = ET.tostring(
                                root,
                                encoding="utf-8",
                                xml_declaration=True,
                            )

                zout.writestr(
                    item,
                    data,
                )

    temp.replace(path)

    return len(related)


# =============================================================================
# Document/package inspection
# =============================================================================

def collect_text(doc: Document) -> str:

    chunks: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text:
            chunks.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    chunks.append(cell.text)

    return "\n".join(chunks)


def inspect_package(
    path: Path,
) -> dict[str, object]:

    with zipfile.ZipFile(path, "r") as archive:

        names = archive.namelist()

        document_xml = archive.read(
            "word/document.xml"
        ).decode(
            "utf-8",
            errors="replace",
        )

        settings_xml = archive.read(
            "word/settings.xml"
        ).decode(
            "utf-8",
            errors="replace",
        )

        footer_xml = "\n".join(
            archive.read(name).decode(
                "utf-8",
                errors="replace",
            )
            for name in names
            if (
                name.startswith("word/footer")
                and name.endswith(".xml")
            )
        )

        media = [
            name
            for name in names
            if name.startswith(
                "word/media/"
            )
        ]

    return {
        "document_xml": document_xml,
        "settings_xml": settings_xml,
        "footer_xml": footer_xml,
        "media": media,
        "tracked_insertions": len(
            re.findall(
                r"<w:ins\b",
                document_xml,
            )
        ),
        "tracked_deletions": len(
            re.findall(
                r"<w:del\b",
                document_xml,
            )
        ),
        "line_numbering": (
            "<w:lnNumType" in document_xml
            and 'w:restart="continuous"'
            in document_xml
        ),
        "page_number_field": (
            " PAGE " in footer_xml
        ),
        "document_protection": (
            "<w:documentProtection"
            in settings_xml
        ),
    }


# =============================================================================
# Placement audit
# =============================================================================

def first_citation(
    paragraphs: list[str],
    token: str,
    excluded_prefixes: tuple[str, ...],
) -> int | None:

    for index, text in enumerate(
        paragraphs,
        start=1,
    ):

        stripped = text.strip()

        if token not in stripped:
            continue

        if any(
            stripped.startswith(prefix)
            for prefix in excluded_prefixes
        ):
            continue

        return index

    return None


def first_prefix(
    paragraphs: list[str],
    prefix: str,
) -> int | None:

    for index, text in enumerate(
        paragraphs,
        start=1,
    ):

        if text.strip().startswith(prefix):
            return index

    return None


def first_exact(
    paragraphs: list[str],
    value: str,
) -> int | None:

    for index, text in enumerate(
        paragraphs,
        start=1,
    ):

        if text.strip() == value:
            return index

    return None


def build_structure_audit(
    doc: Document,
) -> list[dict[str, object]]:

    paragraphs = [
        p.text.strip()
        for p in doc.paragraphs
    ]

    rows: list[dict[str, object]] = []

    for item in (
        "Figure 1",
        "Figure 2",
        "Figure 3",
    ):

        citation = first_citation(
            paragraphs,
            item,
            (
                item + ".",
                "Alt text for " + item,
            ),
        )

        caption = first_prefix(
            paragraphs,
            item + ".",
        )

        gap = (
            caption - citation
            if (
                citation is not None
                and caption is not None
            )
            else ""
        )

        rows.append(
            {
                "item_type": "MAIN_FIGURE",
                "item": item,
                "first_citation_paragraph": (
                    citation or ""
                ),
                "caption_or_label_paragraph": (
                    caption or ""
                ),
                "paragraph_gap": gap,
                "placement_status": (
                    "PLOS_RELOCATION_REQUIRED"
                    if isinstance(gap, int)
                    and gap > 1
                    else "REVIEW"
                ),
            }
        )

    s1_citation = first_citation(
        paragraphs,
        "Figure S1",
        ("Figure S1.",),
    )

    s1_caption = first_prefix(
        paragraphs,
        "Figure S1.",
    )

    rows.append(
        {
            "item_type": "SUPPORTING_FIGURE",
            "item": "Figure S1",
            "first_citation_paragraph": (
                s1_citation or ""
            ),
            "caption_or_label_paragraph": (
                s1_caption or ""
            ),
            "paragraph_gap": (
                s1_caption - s1_citation
                if (
                    s1_citation is not None
                    and s1_caption is not None
                )
                else ""
            ),
            "placement_status": (
                "SUPPORTING_CAPTION_FINAL_PLACEMENT_PENDING"
            ),
        }
    )

    for table_number in (1, 2):

        item = f"Table {table_number}"

        citation = first_citation(
            paragraphs,
            item,
            (
                item + ".",
                item + " reports",
            ),
        )

        label = first_exact(
            paragraphs,
            item,
        )

        gap = (
            label - citation
            if (
                citation is not None
                and label is not None
            )
            else ""
        )

        rows.append(
            {
                "item_type": "MAIN_TABLE",
                "item": item,
                "first_citation_paragraph": (
                    citation or ""
                ),
                "caption_or_label_paragraph": (
                    label or ""
                ),
                "paragraph_gap": gap,
                "placement_status": (
                    "PLOS_RELOCATION_REQUIRED"
                    if isinstance(gap, int)
                    and gap > 1
                    else "REVIEW"
                ),
            }
        )

    return rows


# =============================================================================
# Preflight
# =============================================================================

if not SOURCE.exists():
    fail(
        f"Missing source manuscript: {SOURCE}"
    )

if shutil.which("pandoc") is None:
    fail("pandoc is not available on PATH.")

source_sha_before = sha256_file(SOURCE)

if source_sha_before != EXPECTED_SOURCE_SHA256:
    fail(
        "Authoritative v2.3 SHA mismatch. "
        f"Observed {source_sha_before}; "
        f"expected {EXPECTED_SOURCE_SHA256}."
    )

source_text = SOURCE.read_text(
    encoding="utf-8"
)

WORK_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

OUT_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

FINAL_DOCX.parent.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)


# =============================================================================
# Normalize Markdown
# =============================================================================

normalized_text, transformations = (
    normalize_markdown(source_text)
)

NORMALIZED_MD.write_text(
    normalized_text,
    encoding="utf-8",
    newline="\n",
)

write_tsv(
    TRANSFORM_MANIFEST,
    transformations,
)


# =============================================================================
# Pandoc
# =============================================================================

result = subprocess.run(
    [
        "pandoc",
        str(NORMALIZED_MD),
        "--from=gfm",
        "--to=docx",
        "--standalone",
        "--metadata=lang:en-US",
        "--output",
        str(INTERMEDIATE_DOCX),
    ],
    stdout=subprocess.PIPE,
    stderr=subprocess.PIPE,
    text=True,
)

if result.returncode != 0:
    fail(
        "Pandoc conversion failed:\n"
        + result.stderr
    )

if not INTERMEDIATE_DOCX.exists():
    fail(
        "Pandoc intermediate DOCX "
        "was not created."
    )


# =============================================================================
# Word formatting
# =============================================================================

format_docx(
    INTERMEDIATE_DOCX,
    FINAL_DOCX,
)


# =============================================================================
# Empty-comment cleanup
# =============================================================================

comment_before = inspect_comments(
    FINAL_DOCX
)

actual_comment_markers = (
    int(comment_before["comment_count"])
    + int(comment_before["comment_range_start"])
    + int(comment_before["comment_range_end"])
    + int(comment_before["comment_reference"])
)

if actual_comment_markers != 0:
    fail(
        "Actual Word comments were unexpectedly present."
    )

comment_parts_removed = (
    strip_empty_comment_package(
        FINAL_DOCX
    )
)

comment_after = inspect_comments(
    FINAL_DOCX
)

write_tsv(
    COMMENT_AUDIT,
    [
        {
            "stage": "before_cleanup",
            "comment_related_parts": (
                "|".join(
                    comment_before[
                        "comment_related_parts"
                    ]
                )
                or "NONE"
            ),
            "comment_count": (
                comment_before["comment_count"]
            ),
            "comment_range_start": (
                comment_before[
                    "comment_range_start"
                ]
            ),
            "comment_range_end": (
                comment_before[
                    "comment_range_end"
                ]
            ),
            "comment_reference": (
                comment_before[
                    "comment_reference"
                ]
            ),
        },
        {
            "stage": "after_cleanup",
            "comment_related_parts": (
                "|".join(
                    comment_after[
                        "comment_related_parts"
                    ]
                )
                or "NONE"
            ),
            "comment_count": (
                comment_after["comment_count"]
            ),
            "comment_range_start": (
                comment_after[
                    "comment_range_start"
                ]
            ),
            "comment_range_end": (
                comment_after[
                    "comment_range_end"
                ]
            ),
            "comment_reference": (
                comment_after[
                    "comment_reference"
                ]
            ),
        },
    ],
)


# =============================================================================
# Inspect final DOCX
# =============================================================================

doc = Document(FINAL_DOCX)
full_text = collect_text(doc)
package = inspect_package(FINAL_DOCX)

structure_rows = build_structure_audit(
    doc
)

write_tsv(
    STRUCTURE_AUDIT,
    structure_rows,
)

source_sha_after = sha256_file(SOURCE)
final_docx_sha = sha256_file(FINAL_DOCX)


# =============================================================================
# QA
# =============================================================================

checks: list[dict[str, object]] = []


def add_check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": f"Q{len(checks) + 1:02d}",
            "check_description": description,
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(observed),
            "expected": str(expected),
        }
    )


add_check(
    "Authoritative v2.3 source SHA matched",
    source_sha_before == EXPECTED_SOURCE_SHA256,
    source_sha_before,
    EXPECTED_SOURCE_SHA256,
)

add_check(
    "Authoritative v2.3 source remained unchanged",
    source_sha_after == EXPECTED_SOURCE_SHA256,
    source_sha_after,
    EXPECTED_SOURCE_SHA256,
)

add_check(
    "Clean DOCX exists",
    FINAL_DOCX.exists(),
    FINAL_DOCX.exists(),
    True,
)

add_check(
    "Internal working heading is absent",
    "Complete Manuscript Draft v2.3"
    not in full_text,
    full_text.count(
        "Complete Manuscript Draft v2.3"
    ),
    0,
)

add_check(
    "Internal Article category is absent",
    "Article category" not in full_text,
    full_text.count("Article category"),
    0,
)

first_nonempty = [
    p.text.strip()
    for p in doc.paragraphs
    if p.text.strip()
][:20]

add_check(
    "Front-matter Title label is absent",
    "Title" not in first_nonempty,
    (
        "absent"
        if "Title" not in first_nonempty
        else "present"
    ),
    "absent",
)

table_title_count = sum(
    p.text.strip() == "Title"
    for p in doc.paragraphs
)

add_check(
    "Two table-local Title headings are preserved",
    table_title_count == 2,
    table_title_count,
    2,
)

editable_table_count = sum(
    p.text.strip() == "Editable table"
    for p in doc.paragraphs
)

add_check(
    "Two Editable table labels are preserved",
    editable_table_count == 2,
    editable_table_count,
    2,
)

add_check(
    "Article title is present",
    ARTICLE_TITLE in full_text,
    "present",
    "present",
)

add_check(
    "Short title is present",
    SHORT_TITLE in full_text,
    "present",
    "present",
)

add_check(
    "Author line is present",
    AUTHOR_LINE in full_text,
    "present",
    "present",
)

add_check(
    "SFUCHAS affiliation is present",
    AFFILIATION_1 in full_text,
    "present",
    "present",
)

add_check(
    "AfroBiomics affiliation is present",
    AFFILIATION_2 in full_text,
    "present",
    "present",
)

add_check(
    "Both correspondence emails are present",
    EMAIL_LINE in full_text,
    "present",
    "present",
)

scientific_anchors = [
    (
        "Discovery cohort n",
        "prespecified primary discovery set contained 224 samples",
    ),
    (
        "Discovery group counts",
        "101 bacterial and 123 viral",
    ),
    (
        "GSE73461 reference population",
        "55 Control samples (n = 201)",
    ),
    (
        "GSE72810 primary contrast",
        "23 definite bacterial versus 28 definite viral samples",
    ),
    (
        "Deletion variant count",
        "29,826",
    ),
    (
        "Minimum Pearson correlation",
        "0.9940",
    ),
    (
        "Z-score equation",
        "z_gi = (x_gi - mean_g) / SD_g",
    ),
    (
        "Module-score equation",
        "score_i = (1/K) sum_g z_gi",
    ),
    (
        "Figure 2C correction",
        "independent points for each categorical module",
    ),
]

for label, anchor in scientific_anchors:
    add_check(
        label + " is preserved",
        anchor in full_text,
        "present" if anchor in full_text else "absent",
        "present",
    )

add_check(
    "AI disclosure is preserved",
    (
        "ChatGPT" in full_text
        and "OpenAI" in full_text
    ),
    "present",
    "present",
)

add_check(
    "Public revision branch is preserved",
    "plosone_revision_round1_2026"
    in full_text,
    "present",
    "present",
)

geo_urls = [
    "https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE211567",
    "https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE73461",
    "https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE72810",
]

geo_count = sum(
    url in full_text
    for url in geo_urls
)

add_check(
    "All direct GEO URLs are preserved",
    geo_count == 3,
    geo_count,
    3,
)

add_check(
    "Exactly two editable tables are present",
    len(doc.tables) == 2,
    len(doc.tables),
    2,
)

add_check(
    "No figures are embedded",
    len(package["media"]) == 0,
    len(package["media"]),
    0,
)

add_check(
    "Continuous line numbering is encoded",
    bool(package["line_numbering"]),
    package["line_numbering"],
    True,
)

add_check(
    "Page-number field is encoded",
    bool(package["page_number_field"]),
    package["page_number_field"],
    True,
)

add_check(
    "Pre-cleanup comments package contained zero real comments",
    actual_comment_markers == 0,
    actual_comment_markers,
    0,
)

add_check(
    "Empty comment-related package parts were removed",
    len(
        comment_after[
            "comment_related_parts"
        ]
    ) == 0,
    (
        "|".join(
            comment_after[
                "comment_related_parts"
            ]
        )
        or "NONE"
    ),
    "NONE",
)

add_check(
    "Final DOCX has zero comments",
    int(
        comment_after[
            "comment_count"
        ]
    ) == 0,
    comment_after[
        "comment_count"
    ],
    0,
)

add_check(
    "Final DOCX has zero comment anchors",
    (
        int(
            comment_after[
                "comment_range_start"
            ]
        ) == 0
        and int(
            comment_after[
                "comment_range_end"
            ]
        ) == 0
        and int(
            comment_after[
                "comment_reference"
            ]
        ) == 0
    ),
    (
        f"{comment_after['comment_range_start']}/"
        f"{comment_after['comment_range_end']}/"
        f"{comment_after['comment_reference']}"
    ),
    "0/0/0",
)

add_check(
    "No tracked insertions are present",
    int(package["tracked_insertions"]) == 0,
    package["tracked_insertions"],
    0,
)

add_check(
    "No tracked deletions are present",
    int(package["tracked_deletions"]) == 0,
    package["tracked_deletions"],
    0,
)

add_check(
    "Document is not protected",
    not bool(
        package["document_protection"]
    ),
    package["document_protection"],
    False,
)

add_check(
    "Document has one section",
    len(doc.sections) == 1,
    len(doc.sections),
    1,
)

deep_headings = [
    p.style.name
    for p in doc.paragraphs
    if p.style.name in {
        "Heading 4",
        "Heading 5",
        "Heading 6",
        "Heading 7",
        "Heading 8",
        "Heading 9",
    }
]

add_check(
    "Heading hierarchy is limited to Heading 1-3",
    len(deep_headings) == 0,
    len(deep_headings),
    0,
)

placeholder_hits = [
    token
    for token in (
        "[[",
        "TODO",
        "TBD",
        "PLACEHOLDER",
    )
    if token in full_text
]

add_check(
    "No obvious internal placeholders remain",
    len(placeholder_hits) == 0,
    (
        "none"
        if not placeholder_hits
        else "|".join(placeholder_hits)
    ),
    "none",
)


# =============================================================================
# Final status
# =============================================================================

quality_passed = sum(
    row["pass"] == "TRUE"
    for row in checks
)

quality_failed = (
    len(checks) - quality_passed
)

quality_gate = (
    "PASS"
    if quality_failed == 0
    else "FAIL"
)

relocation_count = sum(
    row["placement_status"]
    == "PLOS_RELOCATION_REQUIRED"
    for row in structure_rows
)

final_status = (
    "READY_FOR_RENDER_AND_PLOS_PLACEMENT_CORRECTION"
    if quality_failed == 0
    else "CLEAN_DOCX_REQUIRES_CORRECTION"
)


# =============================================================================
# Write QA
# =============================================================================

write_tsv(
    QUALITY_GATE,
    checks,
)

write_tsv(
    QUALITY_SUMMARY,
    [
        {
            "source_manuscript": str(SOURCE),
            "source_sha256": source_sha_before,
            "clean_docx": str(FINAL_DOCX),
            "clean_docx_sha256": final_docx_sha,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
            "embedded_media_files": len(
                package["media"]
            ),
            "format_transformations": len(
                transformations
            ),
            "empty_comment_parts_removed": (
                comment_parts_removed
            ),
            "placement_items_requiring_relocation": (
                relocation_count
            ),
            "quality_checks": len(checks),
            "quality_checks_passed": (
                quality_passed
            ),
            "quality_checks_failed": (
                quality_failed
            ),
            "quality_gate": quality_gate,
            "final_status": final_status,
        }
    ],
)


# =============================================================================
# Report
# =============================================================================

report_lines = [
    "# PLOS ONE Clean Revised Manuscript DOCX v2.3",
    "",
    "Manuscript: PONE-D-26-30583",
    "",
    f"Source SHA256: `{source_sha_before}`",
    "",
    f"Clean DOCX: `{FINAL_DOCX}`",
    "",
    f"Clean DOCX SHA256: `{final_docx_sha}`",
    "",
    "## Corrected build",
    "",
    "- Only the front-matter Title label was removed.",
    "- Both table-local Title headings were preserved.",
    "- Both Editable table labels were preserved for placement review.",
    (
        "- Empty Word-comment package infrastructure was removed "
        "only after confirming that it contained no real comments."
    ),
    "- No scientific manuscript content was altered.",
    "",
    "## Placement audit",
    "",
    (
        f"- Main figure/table items currently requiring PLOS "
        f"read-order relocation: {relocation_count}."
    ),
    (
        "- These elements were intentionally not moved in this build; "
        "the move will be performed after render review."
    ),
    "",
    "## Quality gate",
    "",
    (
        f"- Checks passed: "
        f"{quality_passed}/{len(checks)}."
    ),
    f"- Quality gate: `{quality_gate}`.",
    f"- Final status: `{final_status}`.",
    "",
]

REPORT.write_text(
    "\n".join(report_lines),
    encoding="utf-8",
    newline="\n",
)


# =============================================================================
# Console
# =============================================================================

print(
    "===== PLOS ONE CLEAN MANUSCRIPT DOCX V2.3 - CORRECTED BUILD V2 ====="
)

print(
    f"source_sha256\t{source_sha_before}"
)

print(
    f"clean_docx_sha256\t{final_docx_sha}"
)

print(
    f"format_transformations\t{len(transformations)}"
)

print(
    f"paragraphs\t{len(doc.paragraphs)}"
)

print(
    f"tables\t{len(doc.tables)}"
)

print(
    f"sections\t{len(doc.sections)}"
)

print(
    f"embedded_media_files\t{len(package['media'])}"
)

print(
    f"empty_comment_parts_removed\t{comment_parts_removed}"
)

print(
    "placement_items_requiring_relocation\t"
    f"{relocation_count}"
)

print(
    f"quality_checks_passed\t"
    f"{quality_passed}/{len(checks)}"
)

print(
    f"quality_gate\t{quality_gate}"
)

print(
    f"final_status\t{final_status}"
)

print(
    f"clean_docx\t{FINAL_DOCX}"
)

print(
    f"comment_audit\t{COMMENT_AUDIT}"
)

print(
    f"structure_audit\t{STRUCTURE_AUDIT}"
)

print(
    f"transform_manifest\t{TRANSFORM_MANIFEST}"
)

print(
    f"summary\t{QUALITY_SUMMARY}"
)

print(
    f"report\t{REPORT}"
)

if quality_failed:
    fail(
        f"Corrected build failed "
        f"{quality_failed} quality check(s)."
    )
