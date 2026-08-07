#!/usr/bin/env python3

from __future__ import annotations

import csv
import hashlib
import os
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


SOURCE_MD = Path(
    "docs/"
    "complete_manuscript_draft_v2.3_submission_candidate_metadata_restored.md"
)

EXPECTED_SOURCE_SHA = (
    "f3b61e6ddb9f5d38c6211c6cfe0d8694e6ca3b761d52a3245d58df844ab5b2ae"
)

FINAL_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

EXPECTED_INPUT_SHA = (
    "b6537d79201d6a750a6de0479ac756861ec5b826cb8573d9b43f66c539df65f5"
)

WORK = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E13G_continuous_line_numbering"
)

BACKUP = (
    WORK
    / "PONE-D-26-30583_clean_revised_manuscript_pre_line_number_fix.docx"
)

CANDIDATE = (
    WORK
    / "PONE-D-26-30583_clean_revised_manuscript_continuous_lines_candidate.docx"
)

SMOKE_DIR = (
    WORK
    / "libreoffice_smoke_render"
)

PROFILE_DIR = (
    WORK
    / "libreoffice_profile"
)

OUT = Path(
    "results/revision_round1/"
    "plosone_clean_manuscript_continuous_line_numbering_v2.3"
)

GATE = (
    OUT
    / "PLOS_ONE_continuous_line_numbering_quality_gate.tsv"
)

SUMMARY = (
    OUT
    / "PLOS_ONE_continuous_line_numbering_summary.tsv"
)

SECTION_AUDIT = (
    OUT
    / "PLOS_ONE_continuous_line_numbering_section_audit.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_continuous_line_numbering_report.md"
)


TABLE1_WIDTHS = [
    0.80,
    1.55,
    0.72,
    0.52,
    0.72,
    0.85,
    0.95,
    0.72,
    0.78,
    1.40,
    0.75,
]

TABLE2_WIDTHS = [
    0.85,
    1.30,
    0.80,
    1.30,
    0.80,
    0.85,
    0.80,
    0.80,
    2.20,
]


def die(message: str) -> None:

    print(
        f"ERROR: {message}",
        file=sys.stderr,
    )

    sys.exit(1)


def sha256(path: Path) -> str:

    digest = hashlib.sha256()

    with path.open(
        "rb"
    ) as handle:

        for chunk in iter(
            lambda: handle.read(
                1024 * 1024
            ),
            b"",
        ):

            digest.update(
                chunk
            )

    return digest.hexdigest()


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:

    if not rows:

        die(
            f"No rows for {path}"
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=list(
                rows[0]
            ),
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()
        writer.writerows(
            rows
        )


def table_matrix(
    table,
) -> list[list[str]]:

    return [
        [
            cell.text
            for cell
            in row.cells
        ]
        for row
        in table.rows
    ]


def set_cell_width(
    cell,
    width_inches: float,
) -> None:

    twips = int(
        round(
            width_inches
            * 1440
        )
    )

    tc_pr = (
        cell._tc.get_or_add_tcPr()
    )

    tc_w = tc_pr.find(
        qn("w:tcW")
    )

    if tc_w is None:

        tc_w = OxmlElement(
            "w:tcW"
        )

        tc_pr.append(
            tc_w
        )

    tc_w.set(
        qn("w:w"),
        str(twips),
    )

    tc_w.set(
        qn("w:type"),
        "dxa",
    )

    cell.width = Inches(
        width_inches
    )


def set_table_widths(
    table,
    widths: list[float],
) -> None:

    if len(
        table.columns
    ) != len(
        widths
    ):

        die(
            "Table width specification has "
            f"{len(widths)} entries but table has "
            f"{len(table.columns)} columns"
        )

    table.autofit = False

    tbl_pr = (
        table._tbl.tblPr
    )

    layout = tbl_pr.find(
        qn("w:tblLayout")
    )

    if layout is None:

        layout = OxmlElement(
            "w:tblLayout"
        )

        tbl_pr.append(
            layout
        )

    layout.set(
        qn("w:type"),
        "fixed",
    )

    tbl_w = tbl_pr.find(
        qn("w:tblW")
    )

    if tbl_w is None:

        tbl_w = OxmlElement(
            "w:tblW"
        )

        tbl_pr.append(
            tbl_w
        )

    tbl_w.set(
        qn("w:w"),
        str(
            int(
                round(
                    sum(widths)
                    * 1440
                )
            )
        ),
    )

    tbl_w.set(
        qn("w:type"),
        "dxa",
    )

    grid_cols = list(
        table._tbl.tblGrid.gridCol_lst
    )

    for index, width in enumerate(
        widths
    ):

        if index < len(
            grid_cols
        ):

            grid_cols[
                index
            ].set(
                qn("w:w"),
                str(
                    int(
                        round(
                            width
                            * 1440
                        )
                    )
                ),
            )

    for row in table.rows:

        for index, cell in enumerate(
            row.cells
        ):

            set_cell_width(
                cell,
                widths[index],
            )


def fix_section_line_numbering(
    doc: Document,
) -> list[dict[str, object]]:

    rows = []

    for index, section in enumerate(
        doc.sections,
        start=1,
    ):

        sect_pr = (
            section._sectPr
        )

        ln = sect_pr.find(
            qn("w:lnNumType")
        )

        if ln is None:

            ln = OxmlElement(
                "w:lnNumType"
            )

            sect_pr.append(
                ln
            )

        ln.set(
            qn("w:countBy"),
            "1",
        )

        ln.set(
            qn("w:restart"),
            "continuous",
        )

        if index == 1:

            ln.set(
                qn("w:start"),
                "1",
            )

        else:

            key = qn(
                "w:start"
            )

            if key in ln.attrib:

                del ln.attrib[
                    key
                ]

        pg_sz = sect_pr.find(
            qn("w:pgSz")
        )

        orientation = (
            "portrait"
        )

        if (
            pg_sz is not None
            and pg_sz.get(
                qn("w:orient")
            )
            == "landscape"
        ):

            orientation = (
                "landscape"
            )

        rows.append(
            {
                "section_index": index,
                "orientation": orientation,
                "count_by": ln.get(
                    qn("w:countBy"),
                    "",
                ),
                "restart": ln.get(
                    qn("w:restart"),
                    "",
                ),
                "start": ln.get(
                    qn("w:start"),
                    "NONE",
                ),
            }
        )

    return rows


def audit_saved_sections(
    doc: Document,
) -> list[dict[str, object]]:

    rows = []

    for index, section in enumerate(
        doc.sections,
        start=1,
    ):

        sect_pr = (
            section._sectPr
        )

        ln = sect_pr.find(
            qn("w:lnNumType")
        )

        pg_sz = sect_pr.find(
            qn("w:pgSz")
        )

        orientation = (
            "portrait"
        )

        if (
            pg_sz is not None
            and pg_sz.get(
                qn("w:orient")
            )
            == "landscape"
        ):

            orientation = (
                "landscape"
            )

        rows.append(
            {
                "section_index": index,
                "orientation": orientation,
                "count_by": (
                    "NONE"
                    if ln is None
                    else ln.get(
                        qn("w:countBy"),
                        "NONE",
                    )
                ),
                "restart": (
                    "NONE"
                    if ln is None
                    else ln.get(
                        qn("w:restart"),
                        "NONE",
                    )
                ),
                "start": (
                    "NONE"
                    if ln is None
                    else ln.get(
                        qn("w:start"),
                        "NONE",
                    )
                ),
            }
        )

    return rows


def package_state(
    path: Path,
) -> dict[str, object]:

    with zipfile.ZipFile(
        path
    ) as archive:

        names = (
            archive.namelist()
        )

        document_xml = (
            archive.read(
                "word/document.xml"
            ).decode(
                "utf-8",
                errors="replace",
            )
        )

        settings_xml = (
            archive.read(
                "word/settings.xml"
            ).decode(
                "utf-8",
                errors="replace",
            )
        )

        footer_xml = "\n".join(
            archive.read(
                name
            ).decode(
                "utf-8",
                errors="replace",
            )
            for name in names
            if (
                name.startswith(
                    "word/footer"
                )
                and name.endswith(
                    ".xml"
                )
            )
        )

    return {
        "comments": [
            name
            for name in names
            if (
                "comment"
                in name.lower()
                or "people"
                in name.lower()
                or "person"
                in name.lower()
            )
        ],
        "media": [
            name
            for name in names
            if name.startswith(
                "word/media/"
            )
        ],
        "tracked_insertions": len(
            re.findall(
                r"<w:ins\b",
                document_xml,
            )
        ),
        "tracked_deletions": len(
            re.findall(
                r"<w:del\b",
                document_xml,
            )
        ),
        "section_count": len(
            re.findall(
                r"<w:sectPr\b",
                document_xml,
            )
        ),
        "landscape_section_count": len(
            re.findall(
                (
                    r'<w:pgSz[^>]*'
                    r'w:orient="landscape"'
                ),
                document_xml,
            )
        ),
        "line_numbering_count": len(
            re.findall(
                r"<w:lnNumType\b",
                document_xml,
            )
        ),
        "page_field": (
            " PAGE "
            in footer_xml
        ),
        "protection": (
            "<w:documentProtection"
            in settings_xml
        ),
    }


def smoke_render(
    path: Path,
):

    if shutil.which(
        "soffice"
    ) is None:

        die(
            "soffice not found"
        )

    SMOKE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    PROFILE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    for pdf in (
        SMOKE_DIR.glob(
            "*.pdf"
        )
    ):

        pdf.unlink()

    abs_profile = (
        PROFILE_DIR.resolve()
    )

    environment = (
        os.environ.copy()
    )

    environment[
        "HOME"
    ] = str(
        abs_profile
    )

    result = subprocess.run(
        [
            "soffice",
            "--headless",
            (
                "-env:UserInstallation="
                f"file://{abs_profile}"
            ),
            "--convert-to",
            "pdf",
            "--outdir",
            str(
                SMOKE_DIR.resolve()
            ),
            str(
                path.resolve()
            ),
        ],
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        env=environment,
    )

    pdf = (
        SMOKE_DIR
        / f"{path.stem}.pdf"
    )

    success = (
        pdf.exists()
        and pdf.stat().st_size
        > 0
    )

    return (
        success,
        result.returncode,
        pdf,
        result.stdout,
    )


WORK.mkdir(
    parents=True,
    exist_ok=True,
)

OUT.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)


if (
    not SOURCE_MD.exists()
    or not FINAL_DOCX.exists()
):

    die(
        "Required scientific source or canonical DOCX is missing"
    )


source_sha = (
    sha256(
        SOURCE_MD
    )
)

input_sha = (
    sha256(
        FINAL_DOCX
    )
)


if (
    source_sha
    != EXPECTED_SOURCE_SHA
):

    die(
        "Scientific source SHA mismatch: "
        f"{source_sha}"
    )


if (
    input_sha
    != EXPECTED_INPUT_SHA
):

    die(
        "Canonical DOCX SHA mismatch: "
        f"{input_sha}"
    )


if BACKUP.exists():

    if (
        sha256(
            BACKUP
        )
        != EXPECTED_INPUT_SHA
    ):

        die(
            "Existing pre-fix backup has unexpected SHA"
        )

else:

    shutil.copy2(
        FINAL_DOCX,
        BACKUP,
    )


backup_sha = (
    sha256(
        BACKUP
    )
)


doc = Document(
    BACKUP
)


if len(
    doc.sections
) != 5:

    die(
        "Expected 5 sections; observed "
        f"{len(doc.sections)}"
    )


if len(
    doc.tables
) != 2:

    die(
        "Expected 2 tables; observed "
        f"{len(doc.tables)}"
    )


input_t1 = (
    table_matrix(
        doc.tables[0]
    )
)

input_t2 = (
    table_matrix(
        doc.tables[1]
    )
)


fix_section_line_numbering(
    doc
)

set_table_widths(
    doc.tables[0],
    TABLE1_WIDTHS,
)

set_table_widths(
    doc.tables[1],
    TABLE2_WIDTHS,
)


if CANDIDATE.exists():

    CANDIDATE.unlink()


doc.save(
    CANDIDATE
)


candidate = Document(
    CANDIDATE
)

candidate_sha = (
    sha256(
        CANDIDATE
    )
)

candidate_t1 = (
    table_matrix(
        candidate.tables[0]
    )
)

candidate_t2 = (
    table_matrix(
        candidate.tables[1]
    )
)

state = (
    package_state(
        CANDIDATE
    )
)

section_rows = (
    audit_saved_sections(
        candidate
    )
)


write_tsv(
    SECTION_AUDIT,
    section_rows,
)


(
    smoke_ok,
    soffice_status,
    smoke_pdf,
    soffice_output,
) = smoke_render(
    CANDIDATE
)


smoke_bytes = (
    smoke_pdf.stat().st_size
    if smoke_pdf.exists()
    else 0
)


checks = []


def check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": (
                description
            ),
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected
            ),
        }
    )


check(
    "Scientific source SHA locked",
    source_sha
    == EXPECTED_SOURCE_SHA,
    source_sha,
    EXPECTED_SOURCE_SHA,
)

check(
    "Input DOCX SHA locked",
    input_sha
    == EXPECTED_INPUT_SHA,
    input_sha,
    EXPECTED_INPUT_SHA,
)

check(
    "Backup SHA locked",
    backup_sha
    == EXPECTED_INPUT_SHA,
    backup_sha,
    EXPECTED_INPUT_SHA,
)

check(
    "Candidate created",
    CANDIDATE.exists(),
    CANDIDATE.exists(),
    True,
)

check(
    "Five sections retained",
    len(
        candidate.sections
    ) == 5,
    len(
        candidate.sections
    ),
    5,
)

check(
    "Two tables retained",
    len(
        candidate.tables
    ) == 2,
    len(
        candidate.tables
    ),
    2,
)

check(
    "Table 1 cell matrix unchanged",
    candidate_t1
    == input_t1,
    (
        "identical"
        if candidate_t1
        == input_t1
        else "different"
    ),
    "identical",
)

check(
    "Table 2 cell matrix unchanged",
    candidate_t2
    == input_t2,
    (
        "identical"
        if candidate_t2
        == input_t2
        else "different"
    ),
    "identical",
)

check(
    "Five section properties encoded",
    state[
        "section_count"
    ] == 5,
    state[
        "section_count"
    ],
    5,
)

check(
    "Two landscape sections retained",
    state[
        "landscape_section_count"
    ] == 2,
    state[
        "landscape_section_count"
    ],
    2,
)

check(
    "Five line-numbering definitions encoded",
    state[
        "line_numbering_count"
    ] >= 5,
    state[
        "line_numbering_count"
    ],
    ">=5",
)

check(
    "First section starts line numbering at 1",
    section_rows[
        0
    ][
        "start"
    ] == "1",
    section_rows[
        0
    ][
        "start"
    ],
    1,
)

check(
    "Sections 2-5 do not force line-number restart values",
    all(
        row[
            "start"
        ] == "NONE"
        for row
        in section_rows[1:]
    ),
    "|".join(
        str(
            row[
                "start"
            ]
        )
        for row
        in section_rows[1:]
    ),
    "NONE|NONE|NONE|NONE",
)

check(
    "Every section uses continuous line numbering",
    all(
        row[
            "restart"
        ] == "continuous"
        for row
        in section_rows
    ),
    "|".join(
        str(
            row[
                "restart"
            ]
        )
        for row
        in section_rows
    ),
    "continuous in all sections",
)

check(
    "Every section counts every line",
    all(
        row[
            "count_by"
        ] == "1"
        for row
        in section_rows
    ),
    "|".join(
        str(
            row[
                "count_by"
            ]
        )
        for row
        in section_rows
    ),
    "1 in all sections",
)

check(
    "No embedded figures",
    len(
        state[
            "media"
        ]
    ) == 0,
    len(
        state[
            "media"
        ]
    ),
    0,
)

check(
    "No comment-related package parts",
    len(
        state[
            "comments"
        ]
    ) == 0,
    (
        "|".join(
            state[
                "comments"
            ]
        )
        or "NONE"
    ),
    "NONE",
)

check(
    "No tracked insertions",
    state[
        "tracked_insertions"
    ] == 0,
    state[
        "tracked_insertions"
    ],
    0,
)

check(
    "No tracked deletions",
    state[
        "tracked_deletions"
    ] == 0,
    state[
        "tracked_deletions"
    ],
    0,
)

check(
    "Page-number field retained",
    bool(
        state[
            "page_field"
        ]
    ),
    state[
        "page_field"
    ],
    True,
)

check(
    "Document remains unprotected",
    not state[
        "protection"
    ],
    state[
        "protection"
    ],
    False,
)


candidate_text = "\n".join(
    [
        paragraph.text
        for paragraph
        in candidate.paragraphs
    ]
    +
    [
        cell.text
        for table
        in candidate.tables
        for row
        in table.rows
        for cell
        in row.cells
    ]
)


for anchor in [
    "Reuben S. Maghembe¹˒²*",
    "Fig 1.",
    "Fig 2.",
    "Fig 3.",
    "S1 Fig.",
    (
        "Table 1. External evaluation of the five "
        "predefined GSE211567 discovery modules in GSE73461"
    ),
    (
        "Table 2. Cross-cohort Hodges-Lehmann effects "
        "for the five locked host-response modules"
    ),
    "29,826",
    "0.9940",
    "ChatGPT",
    "plosone_revision_round1_2026",
]:

    check(
        "Anchor preserved: "
        f"{anchor[:50]}",
        anchor
        in candidate_text,
        (
            "present"
            if anchor
            in candidate_text
            else "absent"
        ),
        "present",
    )


check(
    "LibreOffice smoke render succeeded",
    smoke_ok,
    (
        f"status={soffice_status}; "
        f"bytes={smoke_bytes}"
    ),
    "non-empty PDF",
)


pre_failures = sum(
    row[
        "pass"
    ] == "FALSE"
    for row
    in checks
)


promoted_sha = (
    input_sha
)

promotion_ok = False


if pre_failures == 0:

    shutil.copy2(
        CANDIDATE,
        FINAL_DOCX,
    )

    promoted_sha = (
        sha256(
            FINAL_DOCX
        )
    )

    promotion_ok = (
        promoted_sha
        == candidate_sha
    )


check(
    "Candidate promoted to canonical clean DOCX",
    (
        pre_failures == 0
        and promotion_ok
    ),
    (
        promoted_sha
        if promotion_ok
        else "not promoted"
    ),
    (
        candidate_sha
        if pre_failures == 0
        else "promotion withheld"
    ),
)


passed = sum(
    row[
        "pass"
    ] == "TRUE"
    for row
    in checks
)

failed = (
    len(checks)
    - passed
)

gate = (
    "PASS"
    if failed == 0
    else "FAIL"
)

status = (
    "READY_FOR_FINAL_CONTINUOUS_LINE_NUMBER_VISUAL_QA"
    if failed == 0
    else "CONTINUOUS_LINE_NUMBER_FIX_REQUIRES_CORRECTION"
)


write_tsv(
    GATE,
    checks,
)


write_tsv(
    SUMMARY,
    [
        {
            "scientific_source_sha256": (
                source_sha
            ),
            "input_docx_sha256": (
                input_sha
            ),
            "candidate_docx_sha256": (
                candidate_sha
            ),
            "promoted_docx_sha256": (
                promoted_sha
            ),
            "sections": len(
                candidate.sections
            ),
            "landscape_sections": (
                state[
                    "landscape_section_count"
                ]
            ),
            "table1_columns": len(
                candidate.tables[
                    0
                ].columns
            ),
            "table2_columns": len(
                candidate.tables[
                    1
                ].columns
            ),
            "embedded_media_files": len(
                state[
                    "media"
                ]
            ),
            "comment_related_parts": len(
                state[
                    "comments"
                ]
            ),
            "smoke_render_pdf_bytes": (
                smoke_bytes
            ),
            "quality_checks": len(
                checks
            ),
            "quality_checks_passed": (
                passed
            ),
            "quality_checks_failed": (
                failed
            ),
            "quality_gate": (
                gate
            ),
            "final_status": (
                status
            ),
        }
    ],
)


REPORT.write_text(
    "\n".join(
        [
            "# PLOS ONE Continuous Line-Numbering Repair",
            "",
            (
                "Input clean DOCX SHA256: "
                f"`{input_sha}`"
            ),
            (
                "Candidate SHA256: "
                f"`{candidate_sha}`"
            ),
            (
                "Promoted DOCX SHA256: "
                f"`{promoted_sha}`"
            ),
            "",
            (
                "Visual QA showed that the landscape "
                "table sections rendered correctly, but "
                "line numbering restarted at each section "
                "boundary."
            ),
            "",
            (
                "The repair retains start=1 only for the "
                "first section and removes forced start "
                "values from sections 2-5 while retaining "
                "restart=continuous throughout."
            ),
            "",
            (
                "The two landscape tables retain their "
                "data matrices. Column widths were "
                "modestly refined for module identifiers "
                "and compact numeric fields; no table text "
                "or scientific value was changed."
            ),
            "",
            (
                f"Checks passed: "
                f"{passed}/{len(checks)}"
            ),
            (
                f"Quality gate: "
                f"`{gate}`"
            ),
            (
                f"Final status: "
                f"`{status}`"
            ),
            "",
            (
                "A focused render spanning the landscape "
                "section boundaries is still required to "
                "confirm visible line-number continuity."
            ),
            "",
        ]
    ),
    encoding="utf-8",
    newline="\n",
)


print(
    "===== PLOS ONE CONTINUOUS LINE-NUMBERING REPAIR ====="
)

print(
    f"scientific_source_sha256\t{source_sha}"
)

print(
    f"input_docx_sha256\t{input_sha}"
)

print(
    f"candidate_docx_sha256\t{candidate_sha}"
)

print(
    f"promoted_docx_sha256\t{promoted_sha}"
)

print(
    f"sections\t{len(candidate.sections)}"
)

print(
    "landscape_sections\t"
    f"{state['landscape_section_count']}"
)

print(
    f"table1_columns\t"
    f"{len(candidate.tables[0].columns)}"
)

print(
    f"table2_columns\t"
    f"{len(candidate.tables[1].columns)}"
)

print(
    "embedded_media_files\t"
    f"{len(state['media'])}"
)

print(
    "comment_related_parts\t"
    f"{len(state['comments'])}"
)

print(
    f"smoke_render_pdf_bytes\t"
    f"{smoke_bytes}"
)

print(
    f"soffice_status\t"
    f"{soffice_status}"
)

print(
    f"quality_checks_passed\t"
    f"{passed}/{len(checks)}"
)

print(
    f"quality_gate\t"
    f"{gate}"
)

print(
    f"final_status\t"
    f"{status}"
)

print(
    f"section_audit\t"
    f"{SECTION_AUDIT}"
)

print(
    f"summary\t"
    f"{SUMMARY}"
)

print(
    f"quality_gate_file\t"
    f"{GATE}"
)

print(
    f"report\t"
    f"{REPORT}"
)

print()

print(
    "===== LIBREOFFICE SMOKE OUTPUT ====="
)

print(
    soffice_output.strip()
)


if failed:

    die(
        "Continuous line-numbering repair failed "
        f"{failed} quality check(s)."
    )
