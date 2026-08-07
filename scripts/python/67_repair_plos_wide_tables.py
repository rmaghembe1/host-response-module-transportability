#!/usr/bin/env python3

from __future__ import annotations

import copy
import csv
import hashlib
import os
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE_MD = Path(
    "docs/"
    "complete_manuscript_draft_v2.3_submission_candidate_metadata_restored.md"
)

EXPECTED_SOURCE_SHA = (
    "f3b61e6ddb9f5d38c6211c6cfe0d8694e6ca3b761d52a3245d58df844ab5b2ae"
)

FINAL_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

EXPECTED_INPUT_SHA = (
    "a82a36dac6d69fa642c0a749a1514ca2db795e8c5297418fde8d5feb69e8768e"
)

WORK = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E13E_table_visual_repair"
)

BACKUP = (
    WORK
    / "PONE-D-26-30583_clean_revised_manuscript_pre_table_visual_repair.docx"
)

CANDIDATE = (
    WORK
    / "PONE-D-26-30583_clean_revised_manuscript_table_visual_repair_candidate.docx"
)

SMOKE_DIR = (
    WORK
    / "libreoffice_smoke_render"
)

PROFILE_DIR = (
    WORK
    / "libreoffice_profile"
)

OUT = Path(
    "results/revision_round1/"
    "plosone_clean_manuscript_table_visual_repair_v2.3"
)

GATE = (
    OUT
    / "PLOS_ONE_table_visual_repair_quality_gate.tsv"
)

SUMMARY = (
    OUT
    / "PLOS_ONE_table_visual_repair_summary.tsv"
)

AUDIT = (
    OUT
    / "PLOS_ONE_table_visual_repair_audit.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_table_visual_repair_report.md"
)


TABLE1_TITLE = (
    "Table 1. External evaluation of the five predefined "
    "GSE211567 discovery modules in GSE73461"
)

TABLE1_LEGEND = (
    "Module scores were calculated using the pre-specified "
    "unweighted mean z-score rule without gene reselection, "
    "module redefinition, reweighting or diagnostic model "
    "training. The primary projection contrast compared "
    "DefiniteBacterial and DefiniteViral samples."
)

TABLE2_TITLE = (
    "Table 2. Cross-cohort Hodges-Lehmann effects for "
    "the five locked host-response modules"
)


TABLE1_HEADERS = [
    "Module",
    "Biological label",
    "Discovery\ndirection",
    "Locked\ngenes",
    "GSE73461\ngenes scored",
    "Main projection\nmedian difference;\nBH-adjusted P",
    (
        "Primary-only z-score\n"
        "sensitivity\n"
        "median difference;\n"
        "BH-adjusted P"
    ),
    "Expected direction\nretained (main)",
    "Expected direction\nretained (sensitivity)",
    "Interpretation",
    "Missing\ngenes",
]


TABLE2_HEADERS = [
    "Module",
    "GSE73461 effect\n(95% CI)",
    "GSE73461\nBH-adjusted P",
    "GSE72810 effect\n(95% CI)",
    "GSE72810\nBH-adjusted P",
    "Expected direction\nretained in both",
    "CI excludes zero\nin both",
    "FDR significant\nin both",
    "Interpretation",
]


TABLE1_WIDTHS = [
    0.62,
    1.55,
    0.72,
    0.48,
    0.60,
    0.85,
    0.95,
    0.72,
    0.78,
    1.40,
    0.50,
]


TABLE2_WIDTHS = [
    0.65,
    1.20,
    0.75,
    1.20,
    0.75,
    0.85,
    0.80,
    0.80,
    2.20,
]


def die(message: str) -> None:

    print(
        f"ERROR: {message}",
        file=sys.stderr,
    )

    sys.exit(1)


def sha256(path: Path) -> str:

    digest = hashlib.sha256()

    with path.open(
        "rb"
    ) as handle:

        for chunk in iter(
            lambda: handle.read(
                1024 * 1024
            ),
            b"",
        ):

            digest.update(
                chunk
            )

    return digest.hexdigest()


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:

    if not rows:

        die(
            f"No rows for {path}"
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=list(
                rows[0]
            ),
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()

        writer.writerows(
            rows
        )


def prefix_paragraph(
    doc: Document,
    prefix: str,
):

    hits = [
        paragraph
        for paragraph
        in doc.paragraphs
        if paragraph.text.strip().startswith(
            prefix
        )
    ]

    if len(hits) != 1:

        die(
            f"Expected one paragraph beginning "
            f"{prefix!r}; observed {len(hits)}"
        )

    return hits[0]


def table_matrix(
    table,
) -> list[list[str]]:

    return [
        [
            cell.text
            for cell
            in row.cells
        ]
        for row
        in table.rows
    ]


def remove_first_column(
    table,
) -> None:

    if len(
        table.columns
    ) < 2:

        die(
            "Cannot remove first column "
            "from a one-column table"
        )

    grid = (
        table._tbl.tblGrid
    )

    grid_cols = list(
        grid.gridCol_lst
    )

    if grid_cols:

        grid.remove(
            grid_cols[0]
        )

    for row in table.rows:

        cells = list(
            row._tr.tc_lst
        )

        if not cells:

            die(
                "Encountered table row "
                "without cells"
            )

        row._tr.remove(
            cells[0]
        )


def set_repeat_header(
    row,
) -> None:

    tr_pr = (
        row._tr.get_or_add_trPr()
    )

    existing = tr_pr.find(
        qn("w:tblHeader")
    )

    if existing is None:

        existing = OxmlElement(
            "w:tblHeader"
        )

        tr_pr.append(
            existing
        )

    existing.set(
        qn("w:val"),
        "true",
    )


def set_cell_margin(
    cell,
    twips: int = 70,
) -> None:

    tc_pr = (
        cell._tc.get_or_add_tcPr()
    )

    tc_mar = tc_pr.find(
        qn("w:tcMar")
    )

    if tc_mar is None:

        tc_mar = OxmlElement(
            "w:tcMar"
        )

        tc_pr.append(
            tc_mar
        )

    for edge in (
        "top",
        "left",
        "bottom",
        "right",
    ):

        node = tc_mar.find(
            qn(
                f"w:{edge}"
            )
        )

        if node is None:

            node = OxmlElement(
                f"w:{edge}"
            )

            tc_mar.append(
                node
            )

        node.set(
            qn("w:w"),
            str(twips),
        )

        node.set(
            qn("w:type"),
            "dxa",
        )


def set_no_wrap(
    cell,
    enabled: bool = True,
) -> None:

    tc_pr = (
        cell._tc.get_or_add_tcPr()
    )

    existing = tc_pr.find(
        qn("w:noWrap")
    )

    if (
        enabled
        and existing is None
    ):

        tc_pr.append(
            OxmlElement(
                "w:noWrap"
            )
        )

    elif (
        not enabled
        and existing is not None
    ):

        tc_pr.remove(
            existing
        )


def set_cell_width(
    cell,
    width_inches: float,
) -> None:

    twips = int(
        round(
            width_inches
            * 1440
        )
    )

    tc_pr = (
        cell._tc.get_or_add_tcPr()
    )

    tc_w = tc_pr.find(
        qn("w:tcW")
    )

    if tc_w is None:

        tc_w = OxmlElement(
            "w:tcW"
        )

        tc_pr.append(
            tc_w
        )

    tc_w.set(
        qn("w:w"),
        str(twips),
    )

    tc_w.set(
        qn("w:type"),
        "dxa",
    )

    cell.width = Inches(
        width_inches
    )


def set_table_fixed_width(
    table,
    widths: list[float],
) -> None:

    if len(
        table.columns
    ) != len(
        widths
    ):

        die(
            "Width specification has "
            f"{len(widths)} entries but "
            f"table has "
            f"{len(table.columns)} columns"
        )

    table.autofit = False

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    tbl_pr = (
        table._tbl.tblPr
    )

    layout = tbl_pr.find(
        qn("w:tblLayout")
    )

    if layout is None:

        layout = OxmlElement(
            "w:tblLayout"
        )

        tbl_pr.append(
            layout
        )

    layout.set(
        qn("w:type"),
        "fixed",
    )

    total_twips = int(
        round(
            sum(widths)
            * 1440
        )
    )

    tbl_w = tbl_pr.find(
        qn("w:tblW")
    )

    if tbl_w is None:

        tbl_w = OxmlElement(
            "w:tblW"
        )

        tbl_pr.append(
            tbl_w
        )

    tbl_w.set(
        qn("w:w"),
        str(
            total_twips
        ),
    )

    tbl_w.set(
        qn("w:type"),
        "dxa",
    )

    grid = (
        table._tbl.tblGrid
    )

    grid_cols = list(
        grid.gridCol_lst
    )

    for index, width in enumerate(
        widths
    ):

        if index < len(
            grid_cols
        ):

            grid_cols[
                index
            ].set(
                qn("w:w"),
                str(
                    int(
                        round(
                            width
                            * 1440
                        )
                    )
                ),
            )

    for row in table.rows:

        for index, cell in enumerate(
            row.cells
        ):

            set_cell_width(
                cell,
                widths[index],
            )


def apply_table_font_and_alignment(
    table,
    header_font: float,
    body_font: float,
    left_columns: set[int],
    nowrap_columns: set[int],
) -> None:

    for row_index, row in enumerate(
        table.rows
    ):

        for col_index, cell in enumerate(
            row.cells
        ):

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            set_cell_margin(
                cell,
                70,
            )

            set_no_wrap(
                cell,
                col_index
                in nowrap_columns,
            )

            for paragraph in (
                cell.paragraphs
            ):

                paragraph.paragraph_format.space_before = (
                    Pt(0)
                )

                paragraph.paragraph_format.space_after = (
                    Pt(0)
                )

                paragraph.paragraph_format.line_spacing = (
                    1.0
                )

                if (
                    col_index
                    in left_columns
                ):

                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                    )

                else:

                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                    )

                for run in (
                    paragraph.runs
                ):

                    run.font.name = (
                        "Times New Roman"
                    )

                    run.font.size = Pt(
                        header_font
                        if row_index == 0
                        else body_font
                    )

                    run.font.bold = (
                        row_index == 0
                    )

                    run.font.color.rgb = (
                        RGBColor(
                            0,
                            0,
                            0,
                        )
                    )

    set_repeat_header(
        table.rows[0]
    )


def replace_header_row(
    table,
    headers: list[str],
) -> None:

    if len(
        table.rows[0].cells
    ) != len(
        headers
    ):

        die(
            "Header count does not "
            "match table column count"
        )

    for index, text in enumerate(
        headers
    ):

        table.rows[
            0
        ].cells[
            index
        ].text = text


def transform_table1_values(
    table,
) -> None:

    for row in (
        table.rows[1:]
    ):

        for index in (
            5,
            6,
        ):

            text = (
                row.cells[
                    index
                ].text
            )

            if "; " in text:

                row.cells[
                    index
                ].text = (
                    text.replace(
                        "; ",
                        ";\n",
                        1,
                    )
                )

        for index in (
            7,
            8,
        ):

            value = (
                row.cells[
                    index
                ].text.strip().lower()
            )

            if value == "true":

                row.cells[
                    index
                ].text = "Yes"

            elif value == "false":

                row.cells[
                    index
                ].text = "No"


def transform_table2_values(
    table,
) -> None:

    pattern = re.compile(
        r"^([^\n]+?)\s+\((95% CI .+)\)$"
    )

    for row in (
        table.rows[1:]
    ):

        for index in (
            1,
            3,
        ):

            text = " ".join(
                row.cells[
                    index
                ].text.split()
            )

            match = pattern.match(
                text
            )

            if match:

                row.cells[
                    index
                ].text = (
                    f"{match.group(1)}\n"
                    f"({match.group(2)})"
                )


def clone_body_sectpr(
    doc: Document,
):

    return copy.deepcopy(
        doc.element.body.sectPr
    )


def set_section_type_next_page(
    sect_pr,
) -> None:

    node = sect_pr.find(
        qn("w:type")
    )

    if node is None:

        node = OxmlElement(
            "w:type"
        )

        sect_pr.insert(
            0,
            node,
        )

    node.set(
        qn("w:val"),
        "nextPage",
    )


def configure_portrait_sectpr(
    sect_pr,
) -> None:

    set_section_type_next_page(
        sect_pr
    )

    pg_sz = sect_pr.find(
        qn("w:pgSz")
    )

    if pg_sz is None:

        pg_sz = OxmlElement(
            "w:pgSz"
        )

        sect_pr.append(
            pg_sz
        )

    pg_sz.set(
        qn("w:w"),
        str(
            int(
                8.5
                * 1440
            )
        ),
    )

    pg_sz.set(
        qn("w:h"),
        str(
            int(
                11
                * 1440
            )
        ),
    )

    if (
        qn("w:orient")
        in pg_sz.attrib
    ):

        del pg_sz.attrib[
            qn("w:orient")
        ]

    ln = sect_pr.find(
        qn("w:lnNumType")
    )

    if ln is None:

        ln = OxmlElement(
            "w:lnNumType"
        )

        sect_pr.append(
            ln
        )

    ln.set(
        qn("w:countBy"),
        "1",
    )

    ln.set(
        qn("w:restart"),
        "continuous",
    )


def configure_landscape_sectpr(
    sect_pr,
) -> None:

    set_section_type_next_page(
        sect_pr
    )

    pg_sz = sect_pr.find(
        qn("w:pgSz")
    )

    if pg_sz is None:

        pg_sz = OxmlElement(
            "w:pgSz"
        )

        sect_pr.append(
            pg_sz
        )

    pg_sz.set(
        qn("w:w"),
        str(
            int(
                11
                * 1440
            )
        ),
    )

    pg_sz.set(
        qn("w:h"),
        str(
            int(
                8.5
                * 1440
            )
        ),
    )

    pg_sz.set(
        qn("w:orient"),
        "landscape",
    )

    pg_mar = sect_pr.find(
        qn("w:pgMar")
    )

    if pg_mar is None:

        pg_mar = OxmlElement(
            "w:pgMar"
        )

        sect_pr.append(
            pg_mar
        )

    for attr in (
        "top",
        "bottom",
        "left",
        "right",
    ):

        pg_mar.set(
            qn(
                f"w:{attr}"
            ),
            str(
                int(
                    0.60
                    * 1440
                )
            ),
        )

    pg_mar.set(
        qn("w:header"),
        str(
            int(
                0.35
                * 1440
            )
        ),
    )

    pg_mar.set(
        qn("w:footer"),
        str(
            int(
                0.35
                * 1440
            )
        ),
    )

    pg_mar.set(
        qn("w:gutter"),
        "0",
    )

    cols = sect_pr.find(
        qn("w:cols")
    )

    if cols is None:

        cols = OxmlElement(
            "w:cols"
        )

        sect_pr.append(
            cols
        )

    cols.set(
        qn("w:num"),
        "1",
    )

    ln = sect_pr.find(
        qn("w:lnNumType")
    )

    if ln is None:

        ln = OxmlElement(
            "w:lnNumType"
        )

        sect_pr.append(
            ln
        )

    ln.set(
        qn("w:countBy"),
        "1",
    )

    ln.set(
        qn("w:restart"),
        "continuous",
    )


def apply_sectpr_to_paragraph(
    paragraph,
    sect_pr,
) -> None:

    p_pr = (
        paragraph._p.get_or_add_pPr()
    )

    existing = p_pr.find(
        qn("w:sectPr")
    )

    if existing is not None:

        p_pr.remove(
            existing
        )

    p_pr.append(
        copy.deepcopy(
            sect_pr
        )
    )


def previous_body_paragraph(
    doc: Document,
    paragraph,
):

    node = (
        paragraph._p.getprevious()
    )

    while (
        node is not None
        and node.tag
        != qn("w:p")
    ):

        node = (
            node.getprevious()
        )

    if node is None:

        die(
            "No preceding body paragraph before "
            f"{paragraph.text[:60]!r}"
        )

    from docx.text.paragraph import Paragraph

    return Paragraph(
        node,
        doc,
    )


def format_title(
    paragraph,
) -> None:

    paragraph.style = (
        paragraph.part.document.styles[
            "Normal"
        ]
    )

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
    )

    paragraph.paragraph_format.keep_with_next = (
        True
    )

    paragraph.paragraph_format.space_before = (
        Pt(6)
    )

    paragraph.paragraph_format.space_after = (
        Pt(4)
    )

    paragraph.paragraph_format.line_spacing = (
        1.15
    )

    for run in (
        paragraph.runs
    ):

        run.font.name = (
            "Times New Roman"
        )

        run.font.size = (
            Pt(11)
        )

        run.font.bold = (
            True
        )

        run.font.color.rgb = (
            RGBColor(
                0,
                0,
                0,
            )
        )


def format_note_range(
    doc: Document,
    table,
    final_note,
    font_size: float = 10.0,
) -> None:

    body = list(
        doc.element.body.iterchildren()
    )

    table_index = (
        body.index(
            table._tbl
        )
    )

    final_index = (
        body.index(
            final_note._p
        )
    )

    for element in body[
        table_index + 1:
        final_index + 1
    ]:

        if element.tag != qn(
            "w:p"
        ):

            continue

        from docx.text.paragraph import Paragraph

        paragraph = Paragraph(
            element,
            doc,
        )

        paragraph.paragraph_format.space_before = (
            Pt(0)
        )

        paragraph.paragraph_format.space_after = (
            Pt(2)
        )

        paragraph.paragraph_format.line_spacing = (
            1.0
        )

        for run in (
            paragraph.runs
        ):

            run.font.name = (
                "Times New Roman"
            )

            run.font.size = (
                Pt(font_size)
            )

            run.font.color.rgb = (
                RGBColor(
                    0,
                    0,
                    0,
                )
            )


def insert_table1_legend(
    doc: Document,
    table,
) -> None:

    if any(
        paragraph.text.strip()
        == TABLE1_LEGEND
        for paragraph
        in doc.paragraphs
    ):

        return

    paragraph = doc.add_paragraph(
        TABLE1_LEGEND
    )

    paragraph.style = (
        doc.styles["Normal"]
    )

    paragraph.paragraph_format.space_before = (
        Pt(4)
    )

    paragraph.paragraph_format.space_after = (
        Pt(2)
    )

    paragraph.paragraph_format.line_spacing = (
        1.0
    )

    for run in (
        paragraph.runs
    ):

        run.font.name = (
            "Times New Roman"
        )

        run.font.size = (
            Pt(10)
        )

        run.font.color.rgb = (
            RGBColor(
                0,
                0,
                0,
            )
        )

    table._tbl.addnext(
        paragraph._p
    )


def package_state(
    path: Path,
) -> dict[str, object]:

    with zipfile.ZipFile(
        path
    ) as archive:

        names = (
            archive.namelist()
        )

        document_xml = (
            archive.read(
                "word/document.xml"
            ).decode(
                "utf-8",
                errors="replace",
            )
        )

        settings_xml = (
            archive.read(
                "word/settings.xml"
            ).decode(
                "utf-8",
                errors="replace",
            )
        )

    return {
        "comments": [
            name
            for name
            in names
            if (
                "comment"
                in name.lower()
                or "people"
                in name.lower()
                or "person"
                in name.lower()
            )
        ],
        "media": [
            name
            for name
            in names
            if name.startswith(
                "word/media/"
            )
        ],
        "tracked_insertions": len(
            re.findall(
                r"<w:ins\b",
                document_xml,
            )
        ),
        "tracked_deletions": len(
            re.findall(
                r"<w:del\b",
                document_xml,
            )
        ),
        "section_count": len(
            re.findall(
                r"<w:sectPr\b",
                document_xml,
            )
        ),
        "landscape_section_count": len(
            re.findall(
                (
                    r'<w:pgSz[^>]*'
                    r'w:orient="landscape"'
                ),
                document_xml,
            )
        ),
        "line_numbering_count": len(
            re.findall(
                r"<w:lnNumType\b",
                document_xml,
            )
        ),
        "protection": (
            "<w:documentProtection"
            in settings_xml
        ),
    }


def row_has_repeat_header(
    table,
) -> bool:

    tr_pr = (
        table.rows[0]._tr.trPr
    )

    return (
        tr_pr is not None
        and tr_pr.find(
            qn("w:tblHeader")
        )
        is not None
    )


def smoke_render(
    path: Path,
):

    if shutil.which(
        "soffice"
    ) is None:

        die(
            "soffice not found"
        )

    SMOKE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    PROFILE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    for pdf in (
        SMOKE_DIR.glob(
            "*.pdf"
        )
    ):

        pdf.unlink()

    abs_profile = (
        PROFILE_DIR.resolve()
    )

    environment = (
        os.environ.copy()
    )

    environment[
        "HOME"
    ] = str(
        abs_profile
    )

    result = subprocess.run(
        [
            "soffice",
            "--headless",
            (
                "-env:UserInstallation="
                f"file://{abs_profile}"
            ),
            "--convert-to",
            "pdf",
            "--outdir",
            str(
                SMOKE_DIR.resolve()
            ),
            str(
                path.resolve()
            ),
        ],
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        env=environment,
    )

    pdf = (
        SMOKE_DIR
        / f"{path.stem}.pdf"
    )

    success = (
        pdf.exists()
        and pdf.stat().st_size
        > 0
    )

    return (
        success,
        result.returncode,
        pdf,
        result.stdout,
    )


# ---------------------------------------------------------------------------
# Preflight
# ---------------------------------------------------------------------------

WORK.mkdir(
    parents=True,
    exist_ok=True,
)

OUT.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)

if (
    not SOURCE_MD.exists()
    or not FINAL_DOCX.exists()
):

    die(
        "Required source or "
        "clean DOCX is missing"
    )

source_sha = sha256(
    SOURCE_MD
)

input_sha = sha256(
    FINAL_DOCX
)

if (
    source_sha
    != EXPECTED_SOURCE_SHA
):

    die(
        "Scientific source SHA mismatch: "
        f"{source_sha}"
    )

if (
    input_sha
    != EXPECTED_INPUT_SHA
):

    die(
        "Input DOCX SHA mismatch: "
        f"{input_sha}"
    )

if BACKUP.exists():

    if (
        sha256(
            BACKUP
        )
        != EXPECTED_INPUT_SHA
    ):

        die(
            "Existing table-repair backup "
            "has unexpected SHA"
        )

else:

    shutil.copy2(
        FINAL_DOCX,
        BACKUP,
    )

backup_sha = sha256(
    BACKUP
)


# ---------------------------------------------------------------------------
# Transform candidate
# ---------------------------------------------------------------------------

source_doc = Document(
    BACKUP
)

if len(
    source_doc.tables
) != 2:

    die(
        "Expected two tables; "
        f"observed "
        f"{len(source_doc.tables)}"
    )

input_table1 = table_matrix(
    source_doc.tables[0]
)

input_table2 = table_matrix(
    source_doc.tables[1]
)

if (
    input_table1[
        0
    ][
        0
    ].strip()
    != "supplementary_table_section"
):

    die(
        "Unexpected Table 1 first header: "
        f"{input_table1[0][0]!r}"
    )


# Table 1: remove only internal provenance column.

remove_first_column(
    source_doc.tables[0]
)

table1 = (
    source_doc.tables[0]
)

replace_header_row(
    table1,
    TABLE1_HEADERS,
)

transform_table1_values(
    table1
)

set_table_fixed_width(
    table1,
    TABLE1_WIDTHS,
)

apply_table_font_and_alignment(
    table1,
    header_font=8.5,
    body_font=9.0,
    left_columns={
        1,
        9,
    },
    nowrap_columns={
        0,
        3,
        4,
        7,
        8,
        10,
    },
)


# Table 2: preserve all nine columns.

table2 = (
    source_doc.tables[1]
)

replace_header_row(
    table2,
    TABLE2_HEADERS,
)

transform_table2_values(
    table2
)

set_table_fixed_width(
    table2,
    TABLE2_WIDTHS,
)

apply_table_font_and_alignment(
    table2,
    header_font=8.5,
    body_font=9.0,
    left_columns={
        8,
    },
    nowrap_columns={
        0,
        2,
        4,
        5,
        6,
        7,
    },
)


# Concise titles and Table 1 legend.

title1 = prefix_paragraph(
    source_doc,
    "Table 1.",
)

title1.text = (
    TABLE1_TITLE
)

format_title(
    title1
)

insert_table1_legend(
    source_doc,
    table1,
)


title2 = prefix_paragraph(
    source_doc,
    "Table 2.",
)

title2.text = (
    TABLE2_TITLE
)

format_title(
    title2
)


# Dedicated landscape sections.

portrait_break = (
    clone_body_sectpr(
        source_doc
    )
)

configure_portrait_sectpr(
    portrait_break
)

landscape_break = (
    clone_body_sectpr(
        source_doc
    )
)

configure_landscape_sectpr(
    landscape_break
)


final_note1 = prefix_paragraph(
    source_doc,
    (
        "Table 1 reports fixed-module "
        "external projection."
    ),
)

final_note2 = prefix_paragraph(
    source_doc,
    (
        "Effects are Hodges-Lehmann "
        "bacterial-minus-viral location shifts"
    ),
)


apply_sectpr_to_paragraph(
    previous_body_paragraph(
        source_doc,
        title1,
    ),
    portrait_break,
)

apply_sectpr_to_paragraph(
    final_note1,
    landscape_break,
)

apply_sectpr_to_paragraph(
    previous_body_paragraph(
        source_doc,
        title2,
    ),
    portrait_break,
)

apply_sectpr_to_paragraph(
    final_note2,
    landscape_break,
)


format_note_range(
    source_doc,
    table1,
    final_note1,
    10.0,
)

format_note_range(
    source_doc,
    table2,
    final_note2,
    10.0,
)


if CANDIDATE.exists():

    CANDIDATE.unlink()

source_doc.save(
    CANDIDATE
)


# ---------------------------------------------------------------------------
# Reopen and independently inspect
# ---------------------------------------------------------------------------

candidate = Document(
    CANDIDATE
)

if len(
    candidate.tables
) != 2:

    die(
        "Candidate did not retain "
        "exactly two tables"
    )

candidate_t1 = table_matrix(
    candidate.tables[0]
)

candidate_t2 = table_matrix(
    candidate.tables[1]
)

candidate_sha = sha256(
    CANDIDATE
)

state = package_state(
    CANDIDATE
)


# Expected Table 1 transformation.

expected_t1_rows = []

for row_index, row in enumerate(
    input_table1
):

    reduced = list(
        row[1:]
    )

    if row_index == 0:

        reduced = (
            TABLE1_HEADERS[:]
        )

    else:

        for index in (
            5,
            6,
        ):

            if "; " in reduced[
                index
            ]:

                reduced[
                    index
                ] = reduced[
                    index
                ].replace(
                    "; ",
                    ";\n",
                    1,
                )

        for index in (
            7,
            8,
        ):

            value = reduced[
                index
            ].strip().lower()

            if value == "true":

                reduced[
                    index
                ] = "Yes"

            elif value == "false":

                reduced[
                    index
                ] = "No"

    expected_t1_rows.append(
        reduced
    )


# Expected Table 2 formatting transformation.

expected_t2_rows = []

pattern = re.compile(
    r"^([^\n]+?)\s+\((95% CI .+)\)$"
)

for row_index, row in enumerate(
    input_table2
):

    transformed = list(
        row
    )

    if row_index == 0:

        transformed = (
            TABLE2_HEADERS[:]
        )

    else:

        for index in (
            1,
            3,
        ):

            text = " ".join(
                transformed[
                    index
                ].split()
            )

            match = pattern.match(
                text
            )

            if match:

                transformed[
                    index
                ] = (
                    f"{match.group(1)}\n"
                    f"({match.group(2)})"
                )

    expected_t2_rows.append(
        transformed
    )


(
    smoke_ok,
    soffice_status,
    smoke_pdf,
    soffice_output,
) = smoke_render(
    CANDIDATE
)

smoke_bytes = (
    smoke_pdf.stat().st_size
    if smoke_pdf.exists()
    else 0
)


# ---------------------------------------------------------------------------
# QA
# ---------------------------------------------------------------------------

checks: list[
    dict[str, object]
] = []


def check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": (
                description
            ),
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected
            ),
        }
    )


check(
    "Scientific source SHA locked",
    source_sha
    == EXPECTED_SOURCE_SHA,
    source_sha,
    EXPECTED_SOURCE_SHA,
)

check(
    "Input DOCX SHA locked",
    input_sha
    == EXPECTED_INPUT_SHA,
    input_sha,
    EXPECTED_INPUT_SHA,
)

check(
    "Backup SHA locked",
    backup_sha
    == EXPECTED_INPUT_SHA,
    backup_sha,
    EXPECTED_INPUT_SHA,
)

check(
    "Candidate created",
    CANDIDATE.exists(),
    CANDIDATE.exists(),
    True,
)

check(
    "Two tables retained",
    len(
        candidate.tables
    ) == 2,
    len(
        candidate.tables
    ),
    2,
)

check(
    (
        "Table 1 internal provenance "
        "column removed"
    ),
    len(
        candidate.tables[
            0
        ].columns
    ) == 11,
    len(
        candidate.tables[
            0
        ].columns
    ),
    11,
)

check(
    "Table 2 retains nine columns",
    len(
        candidate.tables[
            1
        ].columns
    ) == 9,
    len(
        candidate.tables[
            1
        ].columns
    ),
    9,
)

check(
    (
        "Table 1 data preserved under "
        "reader-facing transformation"
    ),
    candidate_t1
    == expected_t1_rows,
    (
        "identical"
        if candidate_t1
        == expected_t1_rows
        else "different"
    ),
    "identical",
)

check(
    (
        "Table 2 data preserved under "
        "formatting transformation"
    ),
    candidate_t2
    == expected_t2_rows,
    (
        "identical"
        if candidate_t2
        == expected_t2_rows
        else "different"
    ),
    "identical",
)

candidate_title1 = prefix_paragraph(
    candidate,
    "Table 1.",
)

candidate_title2 = prefix_paragraph(
    candidate,
    "Table 2.",
)

check(
    "Table 1 concise title present",
    candidate_title1.text.strip()
    == TABLE1_TITLE,
    candidate_title1.text.strip(),
    TABLE1_TITLE,
)

legend_present = any(
    paragraph.text.strip()
    == TABLE1_LEGEND
    for paragraph
    in candidate.paragraphs
)

check(
    (
        "Table 1 methodological detail "
        "retained as legend below table"
    ),
    legend_present,
    (
        "present"
        if legend_present
        else "absent"
    ),
    "present",
)

check(
    "Table 2 title preserved",
    candidate_title2.text.strip()
    == TABLE2_TITLE,
    candidate_title2.text.strip(),
    TABLE2_TITLE,
)

check(
    (
        "Technical provenance token absent "
        "from manuscript Table 1"
    ),
    "supplementary_table_section"
    not in candidate_t1[0],
    candidate_t1[0],
    "token absent",
)

check(
    "Five manuscript sections encoded",
    state[
        "section_count"
    ] == 5,
    state[
        "section_count"
    ],
    5,
)

check(
    "Two landscape table sections encoded",
    state[
        "landscape_section_count"
    ] == 2,
    state[
        "landscape_section_count"
    ],
    2,
)

check(
    (
        "Continuous line-numbering "
        "properties encoded across sections"
    ),
    state[
        "line_numbering_count"
    ] >= 5,
    state[
        "line_numbering_count"
    ],
    ">=5",
)

check(
    "Table 1 header repeats across pages",
    row_has_repeat_header(
        candidate.tables[0]
    ),
    row_has_repeat_header(
        candidate.tables[0]
    ),
    True,
)

check(
    "Table 2 header repeats across pages",
    row_has_repeat_header(
        candidate.tables[1]
    ),
    row_has_repeat_header(
        candidate.tables[1]
    ),
    True,
)

check(
    "No embedded figures",
    len(
        state[
            "media"
        ]
    ) == 0,
    len(
        state[
            "media"
        ]
    ),
    0,
)

check(
    (
        "No comment-related "
        "package parts"
    ),
    len(
        state[
            "comments"
        ]
    ) == 0,
    (
        "|".join(
            state[
                "comments"
            ]
        )
        or "NONE"
    ),
    "NONE",
)

check(
    "No tracked insertions",
    state[
        "tracked_insertions"
    ] == 0,
    state[
        "tracked_insertions"
    ],
    0,
)

check(
    "No tracked deletions",
    state[
        "tracked_deletions"
    ] == 0,
    state[
        "tracked_deletions"
    ],
    0,
)

check(
    "Document remains unprotected",
    not state[
        "protection"
    ],
    state[
        "protection"
    ],
    False,
)


candidate_text = "\n".join(
    [
        paragraph.text
        for paragraph
        in candidate.paragraphs
    ]
    +
    [
        cell.text
        for table
        in candidate.tables
        for row
        in table.rows
        for cell
        in row.cells
    ]
)


anchors = [
    "Reuben S. Maghembe¹˒²*",
    "224 samples",
    "101 bacterial and 123 viral",
    (
        "23 definite bacterial "
        "and 28 definite viral"
    ),
    "29,826",
    "0.9940",
    "ChatGPT",
    "OpenAI",
    "plosone_revision_round1_2026",
    "Fig 1.",
    "Fig 2.",
    "Fig 3.",
    "S1 Fig.",
]


for anchor in anchors:

    check(
        f"Anchor preserved: {anchor}",
        anchor
        in candidate_text,
        (
            "present"
            if anchor
            in candidate_text
            else "absent"
        ),
        "present",
    )


check(
    "LibreOffice smoke render succeeded",
    smoke_ok,
    (
        f"status={soffice_status}; "
        f"bytes={smoke_bytes}"
    ),
    "non-empty PDF",
)


# ---------------------------------------------------------------------------
# Promote only after every pre-promotion check passes
# ---------------------------------------------------------------------------

pre_failures = sum(
    row[
        "pass"
    ] == "FALSE"
    for row
    in checks
)

promotion_ok = False

promoted_sha = (
    input_sha
)

if pre_failures == 0:

    shutil.copy2(
        CANDIDATE,
        FINAL_DOCX,
    )

    promoted_sha = sha256(
        FINAL_DOCX
    )

    promotion_ok = (
        promoted_sha
        == candidate_sha
    )


check(
    (
        "Candidate promoted to "
        "canonical clean DOCX"
    ),
    (
        pre_failures == 0
        and promotion_ok
    ),
    (
        promoted_sha
        if promotion_ok
        else "not promoted"
    ),
    (
        candidate_sha
        if pre_failures == 0
        else "promotion withheld"
    ),
)


passed = sum(
    row[
        "pass"
    ] == "TRUE"
    for row
    in checks
)

failed = (
    len(checks)
    - passed
)

gate = (
    "PASS"
    if failed == 0
    else "FAIL"
)

status = (
    "READY_FOR_TABLE_FOCUSED_FULL_RENDER_AND_VISUAL_QA"
    if failed == 0
    else "TABLE_VISUAL_REPAIR_REQUIRES_CORRECTION"
)


# ---------------------------------------------------------------------------
# Outputs
# ---------------------------------------------------------------------------

write_tsv(
    GATE,
    checks,
)


audit_rows = [
    {
        "item": "Table 1",
        "input_columns": len(
            input_table1[0]
        ),
        "output_columns": len(
            candidate_t1[0]
        ),
        "orientation": (
            "landscape"
        ),
        "title": (
            TABLE1_TITLE
        ),
        "reader_facing_change": (
            "Removed internal supplementary_table_section "
            "provenance column; preserved all scientific "
            "data columns"
        ),
    },
    {
        "item": "Table 2",
        "input_columns": len(
            input_table2[0]
        ),
        "output_columns": len(
            candidate_t2[0]
        ),
        "orientation": (
            "landscape"
        ),
        "title": (
            TABLE2_TITLE
        ),
        "reader_facing_change": (
            "Preserved all columns; compacted headers "
            "and effect-cell line wrapping"
        ),
    },
]


write_tsv(
    AUDIT,
    audit_rows,
)


write_tsv(
    SUMMARY,
    [
        {
            "scientific_source_sha256": (
                source_sha
            ),
            "input_docx_sha256": (
                input_sha
            ),
            "candidate_docx_sha256": (
                candidate_sha
            ),
            "promoted_docx_sha256": (
                promoted_sha
            ),
            "sections": (
                state[
                    "section_count"
                ]
            ),
            "landscape_sections": (
                state[
                    "landscape_section_count"
                ]
            ),
            "table1_columns": len(
                candidate.tables[
                    0
                ].columns
            ),
            "table2_columns": len(
                candidate.tables[
                    1
                ].columns
            ),
            "embedded_media_files": len(
                state[
                    "media"
                ]
            ),
            "comment_related_parts": len(
                state[
                    "comments"
                ]
            ),
            "smoke_render_pdf_bytes": (
                smoke_bytes
            ),
            "quality_checks": len(
                checks
            ),
            "quality_checks_passed": (
                passed
            ),
            "quality_checks_failed": (
                failed
            ),
            "quality_gate": (
                gate
            ),
            "final_status": (
                status
            ),
        }
    ],
)


REPORT.write_text(
    "\n".join(
        [
            "# PLOS ONE Table Visual Repair",
            "",
            (
                "Input clean DOCX SHA256: "
                f"`{input_sha}`"
            ),
            (
                "Candidate SHA256: "
                f"`{candidate_sha}`"
            ),
            (
                "Promoted DOCX SHA256: "
                f"`{promoted_sha}`"
            ),
            "",
            (
                "Visual QA identified Tables 1 and 2 as "
                "unreadable in portrait because Word "
                "compressed the wide editable tables "
                "into very narrow columns."
            ),
            "",
            (
                "Table 1 was converted to an 11-column "
                "reader-facing table by removing the "
                "internal supplementary_table_section "
                "provenance column. All scientific data "
                "columns were retained. The title was "
                "shortened and its methodological detail "
                "was retained as a legend below the table."
            ),
            "",
            (
                "Table 2 retained all nine columns. "
                "Both tables were assigned dedicated "
                "landscape sections, fixed reader-oriented "
                "column widths, repeating header rows, "
                "black Times New Roman table text, and "
                "compact within-cell spacing."
            ),
            "",
            (
                f"Checks passed: "
                f"{passed}/{len(checks)}"
            ),
            (
                f"Quality gate: "
                f"`{gate}`"
            ),
            (
                f"Final status: "
                f"`{status}`"
            ),
            "",
            (
                "A full render and page-by-page visual "
                "inspection remain required before the "
                "clean manuscript can be accepted."
            ),
            "",
        ]
    ),
    encoding="utf-8",
    newline="\n",
)


# ---------------------------------------------------------------------------
# Console
# ---------------------------------------------------------------------------

print(
    "===== PLOS ONE TABLE VISUAL REPAIR ====="
)

print(
    f"scientific_source_sha256\t"
    f"{source_sha}"
)

print(
    f"input_docx_sha256\t"
    f"{input_sha}"
)

print(
    f"candidate_docx_sha256\t"
    f"{candidate_sha}"
)

print(
    f"promoted_docx_sha256\t"
    f"{promoted_sha}"
)

print(
    f"sections\t"
    f"{state['section_count']}"
)

print(
    f"landscape_sections\t"
    f"{state['landscape_section_count']}"
)

print(
    f"table1_columns\t"
    f"{len(candidate.tables[0].columns)}"
)

print(
    f"table2_columns\t"
    f"{len(candidate.tables[1].columns)}"
)

print(
    f"embedded_media_files\t"
    f"{len(state['media'])}"
)

print(
    f"comment_related_parts\t"
    f"{len(state['comments'])}"
)

print(
    f"smoke_render_pdf_bytes\t"
    f"{smoke_bytes}"
)

print(
    f"soffice_status\t"
    f"{soffice_status}"
)

print(
    f"quality_checks_passed\t"
    f"{passed}/{len(checks)}"
)

print(
    f"quality_gate\t"
    f"{gate}"
)

print(
    f"final_status\t"
    f"{status}"
)

print(
    f"final_docx\t"
    f"{FINAL_DOCX}"
)

print(
    f"audit\t"
    f"{AUDIT}"
)

print(
    f"summary\t"
    f"{SUMMARY}"
)

print(
    f"quality_gate_file\t"
    f"{GATE}"
)

print(
    f"report\t"
    f"{REPORT}"
)

print()

print(
    "===== LIBREOFFICE SMOKE OUTPUT ====="
)

print(
    soffice_output.strip()
)


if failed:

    die(
        "Table visual repair failed "
        f"{failed} quality check(s)."
    )
