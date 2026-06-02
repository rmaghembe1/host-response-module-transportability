#!/usr/bin/env python3

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

src = Path("docs/complete_manuscript_draft_v0.8_cmi_compressed.md")
out = Path("submission/cmi_main_manuscript_v0.8.docx")
out.parent.mkdir(parents=True, exist_ok=True)

text = src.read_text()

# Remove internal draft/package labels that should not appear in the submission DOCX.
remove_lines = {
    "# Complete Manuscript Draft v0.8 — CMI-Compressed Draft",
    "Draft v0.8 compresses the Introduction, Methods and Discussion for CMI-facing length while preserving the CMI-compliant abstract, front matter, Results, supplementary table citations, submission route note and interpretation safeguards.",
    "---",
    "# CMI Final Front-Matter Statements",
    "## Purpose",
    "This document provides final draft front-matter statements for the CMI-facing manuscript package, including author information, funding, competing interests, ethics, author contributions, acknowledgements, data availability and code availability.",
    "# Manuscript Introduction Draft",
    "# Manuscript Methods Draft",
    "# Integrated Manuscript Results Section with Figure and Table Callouts",
    "# Manuscript Discussion Draft",
    "# Main Figure Captions",
    "# Polished Main Figure Captions",
    "# Table 1 Title and Footnotes",
    "# Polished Table 1 Title and Footnotes",
    "# Supplementary Tables Note",
    "# CMI Submission Route Note",
    "# Final Interpretation Boundary Reminder",
}

clean_lines = []
for line in text.splitlines():
    if line.strip() in remove_lines:
        continue
    clean_lines.append(line)

clean = "\n".join(clean_lines)

# Journal-facing heading normalization.
clean = clean.replace("## CMI-compliant structured abstract", "## Abstract")
clean = clean.replace("## APC / publication route note", "## Publication route note")
clean = clean.replace("## Front-matter interpretation boundary", "## Interpretation boundary")

# Remove helper section for constructing Table 1 columns, but keep table title/footnotes.
clean = re.sub(r"\n## Suggested table columns\n.*?(?=\n## Footnotes)", "\n", clean, flags=re.S)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
    style = doc.styles[style_name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

doc.styles["Heading 1"].font.size = Pt(14)
doc.styles["Heading 2"].font.size = Pt(13)
doc.styles["Heading 3"].font.size = Pt(12)

def add_formatted_paragraph(raw):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*)", raw)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

for line in clean.splitlines():
    s = line.strip()

    if not s:
        continue

    # Skip Markdown table separator lines.
    if re.match(r"^\|[\s:\-\|]+\|$", s):
        continue

    if s.startswith("### "):
        doc.add_heading(s[4:], level=3)
    elif s.startswith("## "):
        doc.add_heading(s[3:], level=2)
    elif s.startswith("# "):
        doc.add_heading(s[2:], level=1)
    elif s.startswith("- "):
        p = doc.add_paragraph("• " + s[2:])
        p.paragraph_format.left_indent = Inches(0.25)
    elif s.startswith("|"):
        # Keep simple markdown table rows as readable tab-separated text.
        row = " | ".join(part.strip() for part in s.strip("|").split("|"))
        p = doc.add_paragraph(row)
        p.paragraph_format.left_indent = Inches(0.15)
    else:
        add_formatted_paragraph(s)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Page ")

doc.save(out)
print(f"Wrote: {out}")
