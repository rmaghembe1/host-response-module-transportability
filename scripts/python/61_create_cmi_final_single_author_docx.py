#!/usr/bin/env python3

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

src = Path("docs/complete_manuscript_draft_v1.6_step7b_single_author.md")
out = Path("submission/cmi_main_manuscript_v1.6_single_author.docx")
out.parent.mkdir(parents=True, exist_ok=True)

text = src.read_text()

# Remove repository-only top draft heading and internal draft notes.
text = re.sub(r"^# Complete Manuscript Draft.*?\n\n", "", text, count=1, flags=re.S)

# Remove any residual internal draft/package notes that should not appear in submission DOCX.
internal_note_patterns = [
    r"Draft v0\.9 Step 1 removes internal submission-route and draft-boundary notes from the manuscript body while preserving scientific caution within the Methods, Results and Discussion\.\n*",
    r"Draft v[0-9.]+.*?\n*",
    r"^---\s*$",
]

for pattern in internal_note_patterns:
    text = re.sub(pattern, "", text, flags=re.MULTILINE)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Times New Roman"
    styles[style_name].font.size = Pt(12)
    styles[style_name].font.bold = True

def add_para(line="", bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(line)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p

def add_markdown_table(lines):
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        parts = [x.strip() for x in line.strip().strip("|").split("|")]
        # skip separator row
        if all(set(x) <= {"-"} for x in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

lines = text.splitlines()
i = 0
while i < len(lines):
    line = lines[i]
    s = line.strip()

    if not s:
        doc.add_paragraph("")
        i += 1
        continue

    # Markdown table block
    if s.startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i])
            i += 1
        add_markdown_table(table_lines)
        continue

    if s.startswith("# "):
        doc.add_heading(s[2:], level=1)
    elif s.startswith("## "):
        doc.add_heading(s[3:], level=2)
    elif s.startswith("### "):
        doc.add_heading(s[4:], level=3)
    elif s.startswith("- "):
        add_para("• " + s[2:])
    else:
        # Basic bold markdown cleanup
        s = s.replace("**", "")
        add_para(s)
    i += 1

# Add simple page-number placeholder intentionally omitted; journal system handles pagination.
doc.save(out)
print(f"Wrote: {out}")
