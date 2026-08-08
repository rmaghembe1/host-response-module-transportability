#!/usr/bin/env python3
from __future__ import annotations

import copy
import csv
import hashlib
import os
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def q(tag):
    return f"{{{W}}}{tag}"


CLEAN = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

MARKED = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_marked_up_revised_manuscript.docx"
)

SRC24 = Path(
    "docs/complete_manuscript_draft_v2.4_"
    "submission_candidate_ai_methods_compliance.md"
)

SRC23 = Path(
    "docs/complete_manuscript_draft_v2.3_"
    "submission_candidate_metadata_restored.md"
)

SHA_CLEAN = (
    "691f494f5f2335b1a96f5b8d009c815"
    "d733b3c74fab0f230eaa3f73274f6b38f"
)

SHA_MARKED = (
    "a67dafa915184b1b0823355d8447397a"
    "7d2d644edc46054c4c1e58f867b51871"
)

SHA_SRC24 = (
    "54aabce4263c581dd9685a1ae9d2fd1b"
    "05030d1c3fe072bff78512a293acd6bd"
)

SHA_SRC23 = (
    "f3b61e6ddb9f5d38c6211c6cfe0d869"
    "4e6ca3b761d52a3245d58df844ab5b2ae"
)

AI_HEAD = (
    "Declaration of generative AI and AI-assisted technologies "
    "in the manuscript preparation process"
)

AI_TEXT = (
    "During preparation and revision of this manuscript, the author used ChatGPT, "
    "an AI-assisted tool provided by OpenAI, to assist with editorial organisation, "
    "language refinement, workflow planning, code drafting and checking, formatting "
    "checks, and preparation of manuscript and submission-support materials. "
    "AI-generated suggestions were not treated as scientific evidence. Analysis scripts "
    "were executed against the stated public datasets, and numerical and graphical "
    "outputs were checked using the documented reproducibility, source-lock and "
    "quality-control procedures. The author reviewed and verified the analysis code, "
    "results, references, biological interpretation, figures, tables, manuscript text "
    "and submission materials and takes full responsibility for the final work."
)

METHODS_ANCHOR = (
    "Cross-cohort interpretation and reproducibility boundaries"
)

RESULTS = "Results"

AUTHOR = "Reuben S. Maghembe"


WORK = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E15B_ai_methods_docx_sync"
)

OUT = Path(
    "results/revision_round1/"
    "plosone_ai_methods_docx_sync_v2.4"
)

CLEAN_CAND = (
    WORK
    / "PONE-D-26-30583_clean_v2.4_candidate.docx"
)

MARKED_CAND = (
    WORK
    / "PONE-D-26-30583_marked_v2.4_candidate.docx"
)

ACCEPTED = (
    WORK
    / "PONE-D-26-30583_marked_v2.4_accepted_QA.docx"
)

CLEAN_BAK = (
    WORK
    / "PONE-D-26-30583_clean_pre_v2.4_backup.docx"
)

MARKED_BAK = (
    WORK
    / "PONE-D-26-30583_marked_pre_v2.4_backup.docx"
)

QUALITY = (
    OUT
    / "PLOS_ONE_ai_methods_docx_sync_quality_gate.tsv"
)

SUMMARY = (
    OUT
    / "PLOS_ONE_ai_methods_docx_sync_summary.tsv"
)

AUDIT = (
    OUT
    / "PLOS_ONE_ai_methods_docx_sync_structure_audit.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_ai_methods_docx_sync_report.md"
)


def sha(path):
    digest = hashlib.sha256()

    with path.open("rb") as handle:
        for chunk in iter(
            lambda: handle.read(
                1024 * 1024
            ),
            b"",
        ):
            digest.update(chunk)

    return digest.hexdigest()


def norm(value):
    return " ".join(
        value.replace(
            "\u00a0",
            " ",
        ).split()
    )


def ptext(paragraph):
    return norm(
        "".join(
            paragraph.xpath(
                ".//w:t/text() | .//w:delText/text()",
                namespaces=NS,
            )
        )
    )


def xml(root):
    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def top_ps(body):
    return [
        child
        for child in body
        if child.tag == q("p")
    ]


def matches(
    body,
    text,
):
    return [
        paragraph
        for paragraph in top_ps(body)
        if ptext(paragraph) == text
    ]


def one(
    body,
    text,
):
    hit = matches(
        body,
        text,
    )

    if len(hit) != 1:
        raise RuntimeError(
            f"Expected one top-level {text!r}; "
            f"found {len(hit)}"
        )

    return hit[0]


def one_after(
    body,
    text,
    anchor,
):
    anchor_index = body.index(
        anchor
    )

    hit = [
        paragraph
        for paragraph
        in matches(
            body,
            text,
        )
        if body.index(
            paragraph
        ) > anchor_index
    ]

    if len(hit) != 1:

        positions = [
            body.index(
                paragraph
            ) + 1
            for paragraph
            in matches(
                body,
                text,
            )
        ]

        raise RuntimeError(
            f"Expected one {text!r} after Methods anchor; "
            f"found {len(hit)}. "
            f"All positions={positions}"
        )

    return hit[0]


def write_docx(
    source,
    output,
    overrides,
):
    output.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with zipfile.ZipFile(
        source,
        "r",
    ) as zin:

        with zipfile.ZipFile(
            output,
            "w",
            zipfile.ZIP_DEFLATED,
        ) as zout:

            seen = set()

            for info in zin.infolist():

                name = info.filename

                zout.writestr(
                    name,
                    overrides.get(
                        name,
                        zin.read(
                            name
                        ),
                    ),
                )

                seen.add(
                    name
                )

            for name, payload in (
                overrides.items()
            ):

                if name not in seen:
                    zout.writestr(
                        name,
                        payload,
                    )


def relocate(
    source,
    output,
):
    with zipfile.ZipFile(
        source,
        "r",
    ) as zin:

        root = etree.fromstring(
            zin.read(
                "word/document.xml"
            )
        )

        settings = etree.fromstring(
            zin.read(
                "word/settings.xml"
            )
        )

    body = root.find(
        "w:body",
        namespaces=NS,
    )

    if body is None:
        raise RuntimeError(
            "Missing w:body"
        )

    ai_head = one(
        body,
        AI_HEAD,
    )

    ai_text = one(
        body,
        AI_TEXT,
    )

    methods_anchor = one(
        body,
        METHODS_ANCHOR,
    )

    results_before = matches(
        body,
        RESULTS,
    )

    main_results = one_after(
        body,
        RESULTS,
        methods_anchor,
    )

    donor_ppr = (
        methods_anchor.find(
            "w:pPr",
            namespaces=NS,
        )
    )

    if donor_ppr is None:
        raise RuntimeError(
            "Methods anchor has no w:pPr"
        )

    old_ppr = ai_head.find(
        "w:pPr",
        namespaces=NS,
    )

    if old_ppr is not None:
        ai_head.remove(
            old_ppr
        )

    ai_head.insert(
        0,
        copy.deepcopy(
            donor_ppr
        ),
    )

    body.remove(
        ai_head
    )

    body.remove(
        ai_text
    )

    insert_index = body.index(
        main_results
    )

    body.insert(
        insert_index,
        ai_head,
    )

    body.insert(
        insert_index + 1,
        ai_text,
    )

    paragraphs = top_ps(
        body
    )

    anchor_index = (
        paragraphs.index(
            methods_anchor
        )
    )

    head_index = (
        paragraphs.index(
            ai_head
        )
    )

    text_index = (
        paragraphs.index(
            ai_text
        )
    )

    results_index = (
        paragraphs.index(
            main_results
        )
    )

    if not (
        anchor_index
        < head_index
        < text_index
        < results_index
    ):
        raise RuntimeError(
            "AI block is not inside Methods "
            "before main Results"
        )

    if (
        text_index + 1
        != results_index
    ):
        raise RuntimeError(
            "AI declaration does not "
            "immediately precede main Results"
        )

    head_ppr = ai_head.find(
        "w:pPr",
        namespaces=NS,
    )

    anchor_ppr = (
        methods_anchor.find(
            "w:pPr",
            namespaces=NS,
        )
    )

    style_ok = (
        head_ppr is not None
        and anchor_ppr is not None
        and etree.tostring(
            head_ppr
        )
        == etree.tostring(
            anchor_ppr
        )
    )

    results_after = matches(
        body,
        RESULTS,
    )

    result_positions = [
        top_ps(
            body
        ).index(
            paragraph
        ) + 1
        for paragraph
        in results_after
    ]

    write_docx(
        source,
        output,
        {
            "word/document.xml": (
                xml(
                    root
                )
            ),
            "word/settings.xml": (
                xml(
                    settings
                )
            ),
        },
    )

    return {
        "results_heading_count_before": (
            len(
                results_before
            )
        ),
        "results_heading_count_after": (
            len(
                results_after
            )
        ),
        "results_positions_after": (
            ",".join(
                map(
                    str,
                    result_positions,
                )
            )
        ),
        "methods_anchor_paragraph": (
            anchor_index + 1
        ),
        "ai_heading_paragraph": (
            head_index + 1
        ),
        "ai_declaration_paragraph": (
            text_index + 1
        ),
        "main_results_heading_paragraph": (
            results_index + 1
        ),
        "heading_style_matches_methods_anchor": (
            style_ok
        ),
    }


def unwrap(
    element,
):
    parent = element.getparent()

    if parent is None:
        return

    index = parent.index(
        element
    )

    for child in list(
        element
    ):
        element.remove(
            child
        )

        parent.insert(
            index,
            child,
        )

        index += 1

    parent.remove(
        element
    )


def accept(
    marked,
    output,
):
    with zipfile.ZipFile(
        marked,
        "r",
    ) as zin:

        root = etree.fromstring(
            zin.read(
                "word/document.xml"
            )
        )

        settings = etree.fromstring(
            zin.read(
                "word/settings.xml"
            )
        )

    for element in list(
        root.xpath(
            ".//w:del",
            namespaces=NS,
        )
    ):

        parent = (
            element.getparent()
        )

        if parent is not None:
            parent.remove(
                element
            )

    for element in list(
        root.xpath(
            ".//w:ins",
            namespaces=NS,
        )
    ):
        unwrap(
            element
        )

    track = settings.find(
        "w:trackRevisions",
        namespaces=NS,
    )

    if track is not None:
        settings.remove(
            track
        )

    write_docx(
        marked,
        output,
        {
            "word/document.xml": (
                xml(
                    root
                )
            ),
            "word/settings.xml": (
                xml(
                    settings
                )
            ),
        },
    )


def audit(
    path,
):
    with zipfile.ZipFile(
        path,
        "r",
    ) as zin:

        names = zin.namelist()

        root = etree.fromstring(
            zin.read(
                "word/document.xml"
            )
        )

        settings = etree.fromstring(
            zin.read(
                "word/settings.xml"
            )
        )

    insertions = root.xpath(
        ".//w:ins",
        namespaces=NS,
    )

    deletions = root.xpath(
        ".//w:del",
        namespaces=NS,
    )

    return {
        "insertions": (
            len(
                insertions
            )
        ),
        "deletions": (
            len(
                deletions
            )
        ),
        "track": (
            settings.find(
                "w:trackRevisions",
                namespaces=NS,
            )
            is not None
        ),
        "authors": sorted({
            element.get(
                q(
                    "author"
                ),
                "",
            )
            for element
            in insertions
            + deletions
        }),
        "comments": [
            name
            for name
            in names
            if any(
                token
                in name.lower()
                for token
                in (
                    "comment",
                    "people",
                    "person",
                )
            )
        ],
        "media": [
            name
            for name
            in names
            if name.startswith(
                "word/media/"
            )
        ],
    }


def signature(
    doc,
):
    return (
        [
            paragraph.text
            for paragraph
            in doc.paragraphs
        ],
        [
            [
                [
                    cell.text
                    for cell
                    in row.cells
                ]
                for row
                in table.rows
            ]
            for table
            in doc.tables
        ],
    )


def orientations(
    doc,
):
    return [
        (
            "landscape"
            if section.page_width
            > section.page_height
            else "portrait"
        )
        for section
        in doc.sections
    ]


def render(
    docx,
    folder,
    profile,
):
    folder.mkdir(
        parents=True,
        exist_ok=True,
    )

    profile_dir = (
        WORK
        / "libreoffice_profile"
        / profile
    )

    profile_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    for old_pdf in folder.glob(
        "*.pdf"
    ):
        old_pdf.unlink()

    environment = (
        os.environ.copy()
    )

    environment[
        "HOME"
    ] = str(
        profile_dir.resolve()
    )

    result = subprocess.run(
        [
            "soffice",
            "--headless",
            (
                "-env:UserInstallation="
                f"file://{profile_dir.resolve()}"
            ),
            "--convert-to",
            "pdf",
            "--outdir",
            str(
                folder.resolve()
            ),
            str(
                docx.resolve()
            ),
        ],
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        env=environment,
    )

    pdf = (
        folder
        / f"{docx.stem}.pdf"
    )

    size = (
        pdf.stat().st_size
        if pdf.exists()
        else 0
    )

    return (
        result.returncode,
        size,
        result.stdout.strip(),
    )


def write_tsv(
    path,
    rows,
):
    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=list(
                rows[0]
            ),
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()
        writer.writerows(
            rows
        )


WORK.mkdir(
    parents=True,
    exist_ok=True,
)

OUT.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)


for path in (
    CLEAN,
    MARKED,
    SRC24,
    SRC23,
):
    if not path.exists():
        raise RuntimeError(
            f"Missing required input: {path}"
        )


if shutil.which(
    "soffice"
) is None:
    raise RuntimeError(
        "soffice unavailable"
    )


locks = {
    "clean": sha(
        CLEAN
    ),
    "marked": sha(
        MARKED
    ),
    "src24": sha(
        SRC24
    ),
    "src23": sha(
        SRC23
    ),
}

expected = {
    "clean": SHA_CLEAN,
    "marked": SHA_MARKED,
    "src24": SHA_SRC24,
    "src23": SHA_SRC23,
}


for key in locks:

    if (
        locks[
            key
        ]
        != expected[
            key
        ]
    ):
        raise RuntimeError(
            f"{key} SHA mismatch: "
            f"{locks[key]}"
        )


shutil.copy2(
    CLEAN,
    CLEAN_BAK,
)

shutil.copy2(
    MARKED,
    MARKED_BAK,
)


clean_structure = relocate(
    CLEAN,
    CLEAN_CAND,
)

marked_structure = relocate(
    MARKED,
    MARKED_CAND,
)


clean_audit = audit(
    CLEAN_CAND
)

marked_audit = audit(
    MARKED_CAND
)


accept(
    MARKED_CAND,
    ACCEPTED,
)


clean_doc = Document(
    CLEAN_CAND
)

accepted_doc = Document(
    ACCEPTED
)

original_clean = Document(
    CLEAN
)


clean_signature = signature(
    clean_doc
)

accepted_signature = signature(
    accepted_doc
)

original_tables = signature(
    original_clean
)[1]


clean_orientations = orientations(
    clean_doc
)

accepted_orientations = orientations(
    accepted_doc
)


(
    clean_render_status,
    clean_pdf_bytes,
    clean_render_output,
) = render(
    CLEAN_CAND,
    WORK / "clean_render",
    "clean",
)

(
    marked_render_status,
    marked_pdf_bytes,
    marked_render_output,
) = render(
    MARKED_CAND,
    WORK / "marked_render",
    "marked",
)


checks = []


def check(
    description,
    passed,
    observed,
    expected_value,
):
    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": (
                description
            ),
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected_value
            ),
        }
    )


for key, label in (
    (
        "clean",
        "Clean DOCX SHA locked",
    ),
    (
        "marked",
        "Marked DOCX SHA locked",
    ),
    (
        "src24",
        "v2.4 source SHA locked",
    ),
    (
        "src23",
        "v2.3 source remains locked",
    ),
):
    check(
        label,
        locks[key]
        == expected[key],
        locks[key],
        expected[key],
    )


check(
    "Clean has two legitimate top-level Results headings",
    clean_structure[
        "results_heading_count_before"
    ] == 2,
    clean_structure[
        "results_heading_count_before"
    ],
    2,
)

check(
    "Marked has two legitimate top-level Results headings",
    marked_structure[
        "results_heading_count_before"
    ] == 2,
    marked_structure[
        "results_heading_count_before"
    ],
    2,
)

check(
    "Clean retains two Results headings",
    clean_structure[
        "results_heading_count_after"
    ] == 2,
    clean_structure[
        "results_heading_count_after"
    ],
    2,
)

check(
    "Marked retains two Results headings",
    marked_structure[
        "results_heading_count_after"
    ] == 2,
    marked_structure[
        "results_heading_count_after"
    ],
    2,
)

check(
    "Clean AI block immediately precedes main Results",
    clean_structure[
        "main_results_heading_paragraph"
    ]
    == clean_structure[
        "ai_declaration_paragraph"
    ]
    + 1,
    clean_structure,
    "AI heading -> declaration -> main Results",
)

check(
    "Marked AI block immediately precedes main Results",
    marked_structure[
        "main_results_heading_paragraph"
    ]
    == marked_structure[
        "ai_declaration_paragraph"
    ]
    + 1,
    marked_structure,
    "AI heading -> declaration -> main Results",
)

check(
    "Clean AI heading uses Methods subsection formatting",
    bool(
        clean_structure[
            "heading_style_matches_methods_anchor"
        ]
    ),
    clean_structure[
        "heading_style_matches_methods_anchor"
    ],
    True,
)

check(
    "Marked AI heading uses Methods subsection formatting",
    bool(
        marked_structure[
            "heading_style_matches_methods_anchor"
        ]
    ),
    marked_structure[
        "heading_style_matches_methods_anchor"
    ],
    True,
)

check(
    "Clean has no tracked insertions",
    clean_audit[
        "insertions"
    ] == 0,
    clean_audit[
        "insertions"
    ],
    0,
)

check(
    "Clean has no tracked deletions",
    clean_audit[
        "deletions"
    ] == 0,
    clean_audit[
        "deletions"
    ],
    0,
)

check(
    "Marked retains 108 tracked insertions",
    marked_audit[
        "insertions"
    ] == 108,
    marked_audit[
        "insertions"
    ],
    108,
)

check(
    "Marked retains zero tracked deletions",
    marked_audit[
        "deletions"
    ] == 0,
    marked_audit[
        "deletions"
    ],
    0,
)

check(
    "Marked keeps Track Changes enabled",
    marked_audit[
        "track"
    ],
    marked_audit[
        "track"
    ],
    True,
)

check(
    "Marked revision author preserved",
    marked_audit[
        "authors"
    ]
    == [
        AUTHOR
    ],
    "|".join(
        marked_audit[
            "authors"
        ]
    ),
    AUTHOR,
)

check(
    "No comments in clean",
    not clean_audit[
        "comments"
    ],
    clean_audit[
        "comments"
    ]
    or "NONE",
    "NONE",
)

check(
    "No comments in marked",
    not marked_audit[
        "comments"
    ],
    marked_audit[
        "comments"
    ]
    or "NONE",
    "NONE",
)

check(
    "No embedded media in clean",
    not clean_audit[
        "media"
    ],
    len(
        clean_audit[
            "media"
        ]
    ),
    0,
)

check(
    "No embedded media in marked",
    not marked_audit[
        "media"
    ],
    len(
        marked_audit[
            "media"
        ]
    ),
    0,
)

check(
    "Accepted marked reproduces clean",
    accepted_signature
    == clean_signature,
    (
        "identical"
        if accepted_signature
        == clean_signature
        else "different"
    ),
    "identical",
)

check(
    "Accepted marked preserves orientations",
    accepted_orientations
    == clean_orientations,
    "|".join(
        accepted_orientations
    ),
    "|".join(
        clean_orientations
    ),
)

check(
    "Five sections retained",
    len(
        clean_doc.sections
    ) == 5,
    len(
        clean_doc.sections
    ),
    5,
)

check(
    "Two landscape sections retained",
    clean_orientations.count(
        "landscape"
    ) == 2,
    clean_orientations.count(
        "landscape"
    ),
    2,
)

check(
    "Two editable tables retained",
    len(
        clean_doc.tables
    ) == 2,
    len(
        clean_doc.tables
    ),
    2,
)

check(
    "Table cells unchanged",
    clean_signature[1]
    == original_tables,
    (
        "identical"
        if clean_signature[1]
        == original_tables
        else "different"
    ),
    "identical",
)

check(
    "Clean smoke render succeeds",
    (
        clean_render_status == 0
        and clean_pdf_bytes > 0
    ),
    (
        f"status={clean_render_status}; "
        f"bytes={clean_pdf_bytes}"
    ),
    "status=0 and non-empty PDF",
)

check(
    "Marked smoke render succeeds",
    (
        marked_render_status == 0
        and marked_pdf_bytes > 0
    ),
    (
        f"status={marked_render_status}; "
        f"bytes={marked_pdf_bytes}"
    ),
    "status=0 and non-empty PDF",
)

check(
    "Original clean immutable during QA",
    sha(
        CLEAN
    ) == SHA_CLEAN,
    sha(
        CLEAN
    ),
    SHA_CLEAN,
)

check(
    "Original marked immutable during QA",
    sha(
        MARKED
    ) == SHA_MARKED,
    sha(
        MARKED
    ),
    SHA_MARKED,
)


passed = sum(
    row[
        "pass"
    ] == "TRUE"
    for row
    in checks
)

failed = (
    len(
        checks
    )
    - passed
)

gate = (
    "PASS"
    if failed == 0
    else "FAIL"
)

status = (
    "READY_FOR_AI_METHODS_LOCATION_VISUAL_QA"
    if failed == 0
    else "AI_METHODS_DOCX_SYNC_REQUIRES_CORRECTION"
)


write_tsv(
    QUALITY,
    checks,
)

write_tsv(
    AUDIT,
    [
        {
            "document": "clean",
            **clean_structure,
        },
        {
            "document": "marked",
            **marked_structure,
        },
    ],
)


clean_new = sha(
    CLEAN_CAND
)

marked_new = sha(
    MARKED_CAND
)


write_tsv(
    SUMMARY,
    [
        {
            "source_v24_sha256": (
                locks[
                    "src24"
                ]
            ),
            "prior_clean_sha256": (
                locks[
                    "clean"
                ]
            ),
            "prior_marked_sha256": (
                locks[
                    "marked"
                ]
            ),
            "clean_candidate_sha256": (
                clean_new
            ),
            "marked_candidate_sha256": (
                marked_new
            ),
            "results_heading_count": (
                clean_structure[
                    "results_heading_count_after"
                ]
            ),
            "marked_tracked_insertions": (
                marked_audit[
                    "insertions"
                ]
            ),
            "marked_tracked_deletions": (
                marked_audit[
                    "deletions"
                ]
            ),
            "accepted_marked_matches_clean": (
                accepted_signature
                == clean_signature
            ),
            "sections": (
                len(
                    clean_doc.sections
                )
            ),
            "landscape_sections": (
                clean_orientations.count(
                    "landscape"
                )
            ),
            "tables": (
                len(
                    clean_doc.tables
                )
            ),
            "clean_render_pdf_bytes": (
                clean_pdf_bytes
            ),
            "marked_render_pdf_bytes": (
                marked_pdf_bytes
            ),
            "quality_checks": (
                len(
                    checks
                )
            ),
            "quality_checks_passed": (
                passed
            ),
            "quality_checks_failed": (
                failed
            ),
            "quality_gate": (
                gate
            ),
            "final_status": (
                status
            ),
        }
    ],
)


REPORT.write_text(
    "\n".join(
        [
            "# PLOS ONE AI Declaration DOCX Synchronisation",
            "",
            "Manuscript: PONE-D-26-30583",
            "",
            (
                "- The DOCX legitimately contains two "
                "top-level paragraphs named Results: "
                "the structured-abstract Results subheading "
                "and the main manuscript Results heading."
            ),
            (
                "- The main Results heading is selected as "
                "the unique Results paragraph occurring "
                "after the final Methods subsection anchor."
            ),
            (
                "- The AI declaration wording is moved "
                "unchanged into Methods immediately before "
                "the main Results section."
            ),
            (
                "- The marked-up manuscript retains its "
                "existing tracked-insertion markup."
            ),
            (
                "- Accepting all marked revisions reproduces "
                "the clean candidate exactly."
            ),
            (
                f"- Quality gate: `{gate}`"
            ),
            (
                f"- Final status: `{status}`"
            ),
            "",
        ]
    ),
    encoding="utf-8",
    newline="\n",
)


if failed:
    raise RuntimeError(
        "AI Methods DOCX sync failed "
        f"{failed} QA check(s)"
    )


shutil.copy2(
    CLEAN_CAND,
    CLEAN,
)

shutil.copy2(
    MARKED_CAND,
    MARKED,
)


print(
    "===== PLOS ONE AI METHODS DOCX SYNCHRONISATION ====="
)

print(
    f"source_v24_sha256\t"
    f"{locks['src24']}"
)

print(
    f"prior_clean_sha256\t"
    f"{locks['clean']}"
)

print(
    f"prior_marked_sha256\t"
    f"{locks['marked']}"
)

print(
    f"clean_candidate_sha256\t"
    f"{clean_new}"
)

print(
    f"marked_candidate_sha256\t"
    f"{marked_new}"
)

print(
    f"promoted_clean_sha256\t"
    f"{sha(CLEAN)}"
)

print(
    f"promoted_marked_sha256\t"
    f"{sha(MARKED)}"
)

print(
    f"results_heading_count\t"
    f"{clean_structure['results_heading_count_after']}"
)

print(
    f"marked_tracked_insertions\t"
    f"{marked_audit['insertions']}"
)

print(
    f"marked_tracked_deletions\t"
    f"{marked_audit['deletions']}"
)

print(
    "accepted_marked_matches_clean\t"
    f"{accepted_signature == clean_signature}"
)

print(
    f"sections\t"
    f"{len(clean_doc.sections)}"
)

print(
    f"landscape_sections\t"
    f"{clean_orientations.count('landscape')}"
)

print(
    f"tables\t"
    f"{len(clean_doc.tables)}"
)

print(
    f"clean_render_pdf_bytes\t"
    f"{clean_pdf_bytes}"
)

print(
    f"marked_render_pdf_bytes\t"
    f"{marked_pdf_bytes}"
)

print(
    f"quality_checks_passed\t"
    f"{passed}/{len(checks)}"
)

print(
    f"quality_gate\t"
    f"{gate}"
)

print(
    f"final_status\t"
    f"{status}"
)

print(
    f"structure_audit\t"
    f"{AUDIT}"
)

print(
    f"summary\t"
    f"{SUMMARY}"
)

print(
    f"quality_gate_file\t"
    f"{QUALITY}"
)

print(
    f"report\t"
    f"{REPORT}"
)

print(
    "\n===== CLEAN LIBREOFFICE OUTPUT ====="
)

print(
    clean_render_output
)

print(
    "\n===== MARKED LIBREOFFICE OUTPUT ====="
)

print(
    marked_render_output
)
