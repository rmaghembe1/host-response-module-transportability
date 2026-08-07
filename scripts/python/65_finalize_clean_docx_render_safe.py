#!/usr/bin/env python3

"""
65_finalize_clean_docx_render_safe.py

Purpose
-------
Repair the package-level LibreOffice incompatibility identified in the
clean PLOS ONE revised manuscript after empty Word-comment-package
cleanup.

Diagnostic evidence established that:

1. the cleaned DOCX is a valid ZIP package;
2. python-docx opens it correctly;
3. scientific text, tables, line numbering and page numbering are present;
4. LibreOffice rejects that cleaned package;
5. a plain python-docx load/save round-trip of the same DOCX renders
   successfully in LibreOffice;
6. the round-trip preserves all tested scientific anchors.

This script therefore performs only a render-safe python-docx
normalization of the already quality-gated clean manuscript.

It does NOT:
- change manuscript science;
- change manuscript wording;
- move figure captions;
- move tables;
- alter author metadata;
- alter references;
- alter module definitions;
- alter numerical results.

The five main figure/table placement corrections remain a later phase.
"""

from __future__ import annotations

import csv
import hashlib
import os
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document


# =============================================================================
# Locked inputs
# =============================================================================

SOURCE_MD = Path(
    "docs/"
    "complete_manuscript_draft_v2.3_submission_candidate_metadata_restored.md"
)

EXPECTED_SOURCE_MD_SHA256 = (
    "f3b61e6ddb9f5d38c6211c6cfe0d8694e6ca3b761d52a3245d58df844ab5b2ae"
)

FINAL_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

EXPECTED_PRE_FINALIZATION_SHA256 = (
    "299e0bd45efa1be622e361905279a915cd5896ddbe6df5343420345ad1eccac9"
)


# =============================================================================
# Outputs
# =============================================================================

WORK_DIR = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E13B_render_safe_finalization"
)

LOCKED_INPUT_COPY = (
    WORK_DIR
    / "PONE-D-26-30583_clean_revised_manuscript_pre_roundtrip.docx"
)

CANDIDATE_DOCX = (
    WORK_DIR
    / "PONE-D-26-30583_clean_revised_manuscript_render_safe_candidate.docx"
)

SMOKE_DIR = (
    WORK_DIR / "libreoffice_smoke_render"
)

PROFILE_DIR = (
    WORK_DIR / "libreoffice_profile"
)

OUT_DIR = Path(
    "results/revision_round1/"
    "plosone_clean_manuscript_render_safe_finalization"
)

QUALITY_GATE = (
    OUT_DIR
    / "PLOS_ONE_clean_manuscript_render_safe_quality_gate.tsv"
)

QUALITY_SUMMARY = (
    OUT_DIR
    / "PLOS_ONE_clean_manuscript_render_safe_summary.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_clean_manuscript_render_safe_finalization_report.md"
)


# =============================================================================
# Utilities
# =============================================================================

def fail(message: str) -> None:
    print(
        f"ERROR: {message}",
        file=sys.stderr,
    )
    sys.exit(1)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()

    with path.open("rb") as handle:
        for chunk in iter(
            lambda: handle.read(1024 * 1024),
            b"",
        ):
            digest.update(chunk)

    return digest.hexdigest()


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:

    if not rows:
        fail(
            f"No rows available for {path}"
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    fields = list(
        rows[0].keys()
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=fields,
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()
        writer.writerows(rows)


def collect_text(
    doc: Document,
) -> str:

    pieces: list[str] = []

    for paragraph in doc.paragraphs:
        pieces.append(
            paragraph.text
        )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.append(
                    cell.text
                )

    return "\n".join(
        pieces
    )


def inspect_package(
    path: Path,
) -> dict[str, object]:

    with zipfile.ZipFile(
        path,
        "r",
    ) as archive:

        names = archive.namelist()

        duplicate_names = sorted(
            {
                name
                for name in names
                if names.count(name) > 1
            }
        )

        document_xml = archive.read(
            "word/document.xml"
        ).decode(
            "utf-8",
            errors="replace",
        )

        settings_xml = archive.read(
            "word/settings.xml"
        ).decode(
            "utf-8",
            errors="replace",
        )

        footer_xml = "\n".join(
            archive.read(name).decode(
                "utf-8",
                errors="replace",
            )
            for name in names
            if (
                name.startswith(
                    "word/footer"
                )
                and name.endswith(
                    ".xml"
                )
            )
        )

        comment_related = sorted(
            name
            for name in names
            if (
                "comment" in name.lower()
                or "people" in name.lower()
                or "person" in name.lower()
            )
        )

        media = sorted(
            name
            for name in names
            if name.startswith(
                "word/media/"
            )
        )

    return {
        "zip_members": len(names),
        "duplicates": duplicate_names,
        "comments": comment_related,
        "media": media,
        "line_numbering": (
            "<w:lnNumType"
            in document_xml
        ),
        "page_number_field": (
            " PAGE "
            in footer_xml
        ),
        "tracked_insertions": len(
            re.findall(
                r"<w:ins\b",
                document_xml,
            )
        ),
        "tracked_deletions": len(
            re.findall(
                r"<w:del\b",
                document_xml,
            )
        ),
        "document_protection": (
            "<w:documentProtection"
            in settings_xml
        ),
    }


def libreoffice_smoke_render(
    docx_path: Path,
) -> tuple[
    bool,
    int,
    Path,
    str,
]:

    if shutil.which(
        "soffice"
    ) is None:
        fail(
            "LibreOffice soffice is not available on PATH."
        )

    SMOKE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    PROFILE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    for old_pdf in SMOKE_DIR.glob(
        "*.pdf"
    ):
        old_pdf.unlink()

    abs_input = (
        docx_path.resolve()
    )

    abs_output = (
        SMOKE_DIR.resolve()
    )

    abs_profile = (
        PROFILE_DIR.resolve()
    )

    command = [
        "soffice",
        "--headless",
        (
            "-env:UserInstallation="
            f"file://{abs_profile}"
        ),
        "--convert-to",
        "pdf",
        "--outdir",
        str(abs_output),
        str(abs_input),
    ]

    environment = os.environ.copy()

    environment["HOME"] = str(
        abs_profile
    )

    result = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        env=environment,
    )

    expected_pdf = (
        SMOKE_DIR
        / (
            docx_path.stem
            + ".pdf"
        )
    )

    success = (
        expected_pdf.exists()
        and expected_pdf.stat().st_size > 0
    )

    return (
        success,
        result.returncode,
        expected_pdf,
        result.stdout,
    )


# =============================================================================
# Preflight
# =============================================================================

WORK_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

OUT_DIR.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)

if not SOURCE_MD.exists():
    fail(
        f"Missing authoritative manuscript source: "
        f"{SOURCE_MD}"
    )

source_md_sha = sha256_file(
    SOURCE_MD
)

if (
    source_md_sha
    != EXPECTED_SOURCE_MD_SHA256
):
    fail(
        "Authoritative v2.3 Markdown SHA mismatch. "
        f"Observed {source_md_sha}; "
        f"expected {EXPECTED_SOURCE_MD_SHA256}."
    )

if not FINAL_DOCX.exists():
    fail(
        f"Missing clean DOCX: {FINAL_DOCX}"
    )

observed_pre_sha = sha256_file(
    FINAL_DOCX
)

if (
    observed_pre_sha
    != EXPECTED_PRE_FINALIZATION_SHA256
):
    fail(
        "Clean DOCX pre-finalization SHA mismatch. "
        f"Observed {observed_pre_sha}; "
        f"expected "
        f"{EXPECTED_PRE_FINALIZATION_SHA256}."
    )


# =============================================================================
# Preserve exact pre-finalization package
# =============================================================================

if LOCKED_INPUT_COPY.exists():

    existing_backup_sha = sha256_file(
        LOCKED_INPUT_COPY
    )

    if (
        existing_backup_sha
        != EXPECTED_PRE_FINALIZATION_SHA256
    ):
        fail(
            "Existing locked pre-roundtrip copy has "
            "an unexpected SHA256."
        )

else:

    shutil.copy2(
        FINAL_DOCX,
        LOCKED_INPUT_COPY,
    )


locked_input_sha = sha256_file(
    LOCKED_INPUT_COPY
)

if (
    locked_input_sha
    != EXPECTED_PRE_FINALIZATION_SHA256
):
    fail(
        "Locked input copy SHA mismatch."
    )


# =============================================================================
# Open validated source package
# =============================================================================

try:
    source_doc = Document(
        LOCKED_INPUT_COPY
    )
except Exception as exc:
    fail(
        "python-docx could not open the validated "
        f"input package: {type(exc).__name__}: {exc}"
    )

source_text = collect_text(
    source_doc
)

source_paragraphs = len(
    source_doc.paragraphs
)

source_tables = len(
    source_doc.tables
)

source_sections = len(
    source_doc.sections
)

source_package = inspect_package(
    LOCKED_INPUT_COPY
)


# =============================================================================
# Render-safe python-docx normalization
# =============================================================================

if CANDIDATE_DOCX.exists():
    CANDIDATE_DOCX.unlink()

source_doc.save(
    CANDIDATE_DOCX
)

if not CANDIDATE_DOCX.exists():
    fail(
        "python-docx round-trip did not create "
        "the render-safe candidate."
    )

try:
    candidate_doc = Document(
        CANDIDATE_DOCX
    )
except Exception as exc:
    fail(
        "Round-trip candidate could not be reopened: "
        f"{type(exc).__name__}: {exc}"
    )

candidate_text = collect_text(
    candidate_doc
)

candidate_package = inspect_package(
    CANDIDATE_DOCX
)

candidate_sha = sha256_file(
    CANDIDATE_DOCX
)


# =============================================================================
# LibreOffice smoke render
# =============================================================================

(
    render_success,
    soffice_status,
    smoke_pdf,
    soffice_output,
) = libreoffice_smoke_render(
    CANDIDATE_DOCX
)

smoke_pdf_size = (
    smoke_pdf.stat().st_size
    if smoke_pdf.exists()
    else 0
)


# =============================================================================
# QA
# =============================================================================

checks: list[
    dict[str, object]
] = []


def add_check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": (
                description
            ),
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected
            ),
        }
    )


add_check(
    "Authoritative v2.3 Markdown SHA matched",
    (
        source_md_sha
        == EXPECTED_SOURCE_MD_SHA256
    ),
    source_md_sha,
    EXPECTED_SOURCE_MD_SHA256,
)

add_check(
    "Pre-finalization clean DOCX SHA matched",
    (
        observed_pre_sha
        == EXPECTED_PRE_FINALIZATION_SHA256
    ),
    observed_pre_sha,
    EXPECTED_PRE_FINALIZATION_SHA256,
)

add_check(
    "Locked input copy SHA matched",
    (
        locked_input_sha
        == EXPECTED_PRE_FINALIZATION_SHA256
    ),
    locked_input_sha,
    EXPECTED_PRE_FINALIZATION_SHA256,
)

add_check(
    "Round-trip candidate was created",
    CANDIDATE_DOCX.exists(),
    CANDIDATE_DOCX.exists(),
    True,
)

add_check(
    "Paragraph count is preserved",
    (
        len(candidate_doc.paragraphs)
        == source_paragraphs
    ),
    len(candidate_doc.paragraphs),
    source_paragraphs,
)

add_check(
    "Table count is preserved",
    (
        len(candidate_doc.tables)
        == source_tables
    ),
    len(candidate_doc.tables),
    source_tables,
)

add_check(
    "Section count is preserved",
    (
        len(candidate_doc.sections)
        == source_sections
    ),
    len(candidate_doc.sections),
    source_sections,
)

add_check(
    "Document text is exactly preserved",
    candidate_text == source_text,
    (
        "identical"
        if candidate_text == source_text
        else "different"
    ),
    "identical",
)

add_check(
    "Exactly two editable manuscript tables remain",
    len(candidate_doc.tables) == 2,
    len(candidate_doc.tables),
    2,
)

table_title_count = sum(
    paragraph.text.strip()
    == "Title"
    for paragraph
    in candidate_doc.paragraphs
)

add_check(
    "Two table-local Title headings remain",
    table_title_count == 2,
    table_title_count,
    2,
)

editable_table_labels = sum(
    paragraph.text.strip()
    == "Editable table"
    for paragraph
    in candidate_doc.paragraphs
)

add_check(
    "Two Editable table labels remain",
    editable_table_labels == 2,
    editable_table_labels,
    2,
)

add_check(
    "Candidate has no comment-related package parts",
    len(
        candidate_package[
            "comments"
        ]
    ) == 0,
    (
        "|".join(
            candidate_package[
                "comments"
            ]
        )
        or "NONE"
    ),
    "NONE",
)

add_check(
    "Candidate has no duplicate ZIP members",
    len(
        candidate_package[
            "duplicates"
        ]
    ) == 0,
    (
        "|".join(
            candidate_package[
                "duplicates"
            ]
        )
        or "NONE"
    ),
    "NONE",
)

add_check(
    "Candidate retains continuous line numbering",
    bool(
        candidate_package[
            "line_numbering"
        ]
    ),
    candidate_package[
        "line_numbering"
    ],
    True,
)

add_check(
    "Candidate retains page-number field",
    bool(
        candidate_package[
            "page_number_field"
        ]
    ),
    candidate_package[
        "page_number_field"
    ],
    True,
)

add_check(
    "Candidate has zero tracked insertions",
    (
        int(
            candidate_package[
                "tracked_insertions"
            ]
        ) == 0
    ),
    candidate_package[
        "tracked_insertions"
    ],
    0,
)

add_check(
    "Candidate has zero tracked deletions",
    (
        int(
            candidate_package[
                "tracked_deletions"
            ]
        ) == 0
    ),
    candidate_package[
        "tracked_deletions"
    ],
    0,
)

add_check(
    "Candidate is not protected",
    not bool(
        candidate_package[
            "document_protection"
        ]
    ),
    candidate_package[
        "document_protection"
    ],
    False,
)

add_check(
    "Candidate contains no embedded manuscript figures",
    len(
        candidate_package[
            "media"
        ]
    ) == 0,
    len(
        candidate_package[
            "media"
        ]
    ),
    0,
)


# =============================================================================
# Scientific anchors
# =============================================================================

anchors = [
    (
        "Submitted author marker",
        "Reuben S. Maghembe¹˒²*",
    ),
    (
        "AfroBiomics affiliation",
        "AfroBiomics Co. Ltd.",
    ),
    (
        "Discovery cohort size",
        "224 samples",
    ),
    (
        "Discovery group counts",
        "101 bacterial and 123 viral",
    ),
    (
        "GSE72810 primary contrast",
        (
            "23 definite bacterial versus "
            "28 definite viral samples"
        ),
    ),
    (
        "Deletion variant count",
        "29,826",
    ),
    (
        "Minimum Pearson correlation",
        "0.9940",
    ),
    (
        "AI declaration",
        "ChatGPT",
    ),
    (
        "Public revision branch",
        "plosone_revision_round1_2026",
    ),
]

for label, anchor in anchors:

    present = (
        anchor in candidate_text
    )

    add_check(
        label + " is preserved",
        present,
        (
            "present"
            if present
            else "absent"
        ),
        "present",
    )


# =============================================================================
# Renderability gate
# =============================================================================

add_check(
    "LibreOffice candidate smoke render succeeded",
    render_success,
    (
        f"pdf={smoke_pdf.exists()}; "
        f"bytes={smoke_pdf_size}; "
        f"soffice_status={soffice_status}"
    ),
    "non-empty PDF",
)


# =============================================================================
# Promote only after all QA checks pass
# =============================================================================

quality_passed = sum(
    row["pass"] == "TRUE"
    for row in checks
)

quality_failed = (
    len(checks)
    - quality_passed
)

if quality_failed == 0:

    shutil.copy2(
        CANDIDATE_DOCX,
        FINAL_DOCX,
    )

    promoted_sha = sha256_file(
        FINAL_DOCX
    )

    promotion_success = (
        promoted_sha
        == candidate_sha
    )

else:

    promoted_sha = (
        observed_pre_sha
    )

    promotion_success = False


add_check(
    "Render-safe candidate promotion succeeded",
    (
        quality_failed == 0
        and promotion_success
    ),
    (
        promoted_sha
        if quality_failed == 0
        else "not promoted"
    ),
    (
        candidate_sha
        if quality_failed == 0
        else "candidate promotion withheld"
    ),
)


# Recalculate after promotion check.
quality_passed = sum(
    row["pass"] == "TRUE"
    for row in checks
)

quality_failed = (
    len(checks)
    - quality_passed
)

quality_gate = (
    "PASS"
    if quality_failed == 0
    else "FAIL"
)

final_status = (
    "READY_FOR_FULL_RENDER_AND_VISUAL_REVIEW"
    if quality_failed == 0
    else "RENDER_SAFE_FINALIZATION_REQUIRES_REVIEW"
)


# =============================================================================
# Outputs
# =============================================================================

write_tsv(
    QUALITY_GATE,
    checks,
)

write_tsv(
    QUALITY_SUMMARY,
    [
        {
            "authoritative_source_sha256": (
                source_md_sha
            ),
            "pre_finalization_docx_sha256": (
                observed_pre_sha
            ),
            "candidate_docx_sha256": (
                candidate_sha
            ),
            "final_promoted_docx_sha256": (
                promoted_sha
            ),
            "paragraphs": (
                len(candidate_doc.paragraphs)
            ),
            "tables": (
                len(candidate_doc.tables)
            ),
            "sections": (
                len(candidate_doc.sections)
            ),
            "comment_related_parts": (
                len(
                    candidate_package[
                        "comments"
                    ]
                )
            ),
            "embedded_media_files": (
                len(
                    candidate_package[
                        "media"
                    ]
                )
            ),
            "smoke_render_pdf_bytes": (
                smoke_pdf_size
            ),
            "quality_checks": (
                len(checks)
            ),
            "quality_checks_passed": (
                quality_passed
            ),
            "quality_checks_failed": (
                quality_failed
            ),
            "quality_gate": (
                quality_gate
            ),
            "final_status": (
                final_status
            ),
        }
    ],
)


report_lines = [
    "# PLOS ONE Clean Manuscript Render-Safe Finalization",
    "",
    "Manuscript: PONE-D-26-30583",
    "",
    "## Input",
    "",
    (
        f"- Pre-finalization DOCX SHA256: "
        f"`{observed_pre_sha}`"
    ),
    "",
    "## Render-safe normalization",
    "",
    (
        "- The validated clean DOCX was opened and resaved "
        "with python-docx."
    ),
    (
        "- Main-document text was required to remain exactly "
        "identical."
    ),
    (
        "- Paragraph, table and section counts were required "
        "to remain unchanged."
    ),
    (
        "- Line numbering and page-number fields were required "
        "to remain present."
    ),
    (
        "- Comment-related package parts remained absent."
    ),
    (
        "- LibreOffice smoke rendering was required to produce "
        "a non-empty PDF before promotion."
    ),
    "",
    "## Output",
    "",
    (
        f"- Candidate SHA256: "
        f"`{candidate_sha}`"
    ),
    (
        f"- Promoted clean DOCX SHA256: "
        f"`{promoted_sha}`"
    ),
    "",
    "## Remaining manuscript-layout work",
    "",
    (
        "- Figure 1, Figure 2, Figure 3, Table 1 and Table 2 "
        "still require PLOS read-order placement correction."
    ),
    (
        "- Figure S1 caption remains pending final supporting-"
        "information placement."
    ),
    (
        "- No placement change was performed in this phase."
    ),
    "",
    "## Quality gate",
    "",
    (
        f"- Checks passed: "
        f"{quality_passed}/{len(checks)}"
    ),
    f"- Quality gate: `{quality_gate}`",
    f"- Final status: `{final_status}`",
    "",
]

REPORT.write_text(
    "\n".join(report_lines),
    encoding="utf-8",
    newline="\n",
)


# =============================================================================
# Console
# =============================================================================

print(
    "===== PLOS ONE CLEAN DOCX RENDER-SAFE FINALIZATION ====="
)

print(
    f"source_sha256\t{source_md_sha}"
)

print(
    f"pre_finalization_sha256\t{observed_pre_sha}"
)

print(
    f"candidate_sha256\t{candidate_sha}"
)

print(
    f"promoted_sha256\t{promoted_sha}"
)

print(
    f"paragraphs\t{len(candidate_doc.paragraphs)}"
)

print(
    f"tables\t{len(candidate_doc.tables)}"
)

print(
    f"sections\t{len(candidate_doc.sections)}"
)

print(
    "comment_related_parts\t"
    f"{len(candidate_package['comments'])}"
)

print(
    "embedded_media_files\t"
    f"{len(candidate_package['media'])}"
)

print(
    f"smoke_render_pdf_bytes\t{smoke_pdf_size}"
)

print(
    f"soffice_status\t{soffice_status}"
)

print(
    f"quality_checks_passed\t"
    f"{quality_passed}/{len(checks)}"
)

print(
    f"quality_gate\t{quality_gate}"
)

print(
    f"final_status\t{final_status}"
)

print(
    f"final_docx\t{FINAL_DOCX}"
)

print(
    f"summary\t{QUALITY_SUMMARY}"
)

print(
    f"quality_gate_file\t{QUALITY_GATE}"
)

print(
    f"report\t{REPORT}"
)

print()
print(
    "===== LIBREOFFICE SMOKE OUTPUT ====="
)

print(
    soffice_output.strip()
)

if quality_failed:
    fail(
        f"Render-safe finalization failed "
        f"{quality_failed} quality check(s)."
    )
