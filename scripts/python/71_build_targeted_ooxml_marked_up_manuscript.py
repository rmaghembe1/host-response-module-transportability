#!/usr/bin/env python3

from __future__ import annotations

import copy
import csv
import hashlib
import os
import re
import shutil
import subprocess
import sys
import zipfile
from collections import defaultdict
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NS = {"w": W_NS}


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


# ---------------------------------------------------------------------------
# Locked inputs
# ---------------------------------------------------------------------------

CLEAN_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_clean_revised_manuscript.docx"
)

EXPECTED_CLEAN_SHA = (
    "691f494f5f2335b1a96f5b8d009c815"
    "d733b3c74fab0f230eaa3f73274f6b38f"
)

BASELINE_DOCX = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E14B_submitted_baseline_reconstruction/"
    "PONE-D-26-30583_reconstructed_submitted_baseline.docx"
)

EXPECTED_BASELINE_SHA = (
    "0e0b4a2bcec736d2735ecc903ad72e99"
    "a1a3897bc09e2f991a7249000852dd34"
)

SOURCE_MD = Path(
    "docs/complete_manuscript_draft_v2.3_"
    "submission_candidate_metadata_restored.md"
)

EXPECTED_SOURCE_SHA = (
    "f3b61e6ddb9f5d38c6211c6cfe0d869"
    "4e6ca3b761d52a3245d58df844ab5b2ae"
)

ALIGNMENT_TSV = Path(
    "results/revision_round1/"
    "plosone_controlled_redline_alignment_v2.3/"
    "PLOS_ONE_clean_to_submitted_paragraph_alignment.tsv"
)

TARGET_TSV = Path(
    "results/revision_round1/"
    "plosone_controlled_redline_alignment_v2.3/"
    "PLOS_ONE_reviewer_driven_redline_target_inventory.tsv"
)

ALIGNMENT_SUMMARY = Path(
    "results/revision_round1/"
    "plosone_controlled_redline_alignment_v2.3/"
    "PLOS_ONE_controlled_redline_alignment_summary.tsv"
)

SUBMITTED_TITLE = (
    "Cross-cohort transportability of bacterial- and viral-associated "
    "host-response modules in public infection transcriptomic datasets"
)

SUBMITTED_SHORT_TITLE = (
    "Short title: Transportability of infection host-response modules"
)

REVISION_AUTHOR = "Reuben S. Maghembe"
REVISION_DATE = "2026-08-07T18:58:00Z"


# ---------------------------------------------------------------------------
# Outputs
# ---------------------------------------------------------------------------

WORK = Path(
    "work/plosone_revision_round1_2026/"
    "phaseR1E14D_targeted_ooxml_redline"
)

MARKED_DOCX = Path(
    "submission/revision_round1/"
    "PONE-D-26-30583_marked_up_revised_manuscript.docx"
)

ACCEPTED_QA_DOCX = (
    WORK
    / "PONE-D-26-30583_marked_up_acceptance_QA.docx"
)

SMOKE_DIR = (
    WORK
    / "libreoffice_smoke_render"
)

PROFILE_DIR = (
    WORK
    / "libreoffice_profile"
)

OUT = Path(
    "results/revision_round1/"
    "plosone_targeted_ooxml_redline_v2.3"
)

MANIFEST = (
    OUT
    / "PLOS_ONE_targeted_redline_manifest.tsv"
)

QUALITY_GATE = (
    OUT
    / "PLOS_ONE_targeted_redline_quality_gate.tsv"
)

SUMMARY = (
    OUT
    / "PLOS_ONE_targeted_redline_summary.tsv"
)

REPORT = Path(
    "docs/revision_round1/"
    "PLOS_ONE_targeted_ooxml_redline_report.md"
)


# ---------------------------------------------------------------------------
# General helpers
# ---------------------------------------------------------------------------

def die(message: str) -> None:

    print(
        f"ERROR: {message}",
        file=sys.stderr,
    )

    raise SystemExit(1)


def sha256(path: Path) -> str:

    digest = hashlib.sha256()

    with path.open(
        "rb"
    ) as handle:

        for chunk in iter(
            lambda: handle.read(
                1024 * 1024
            ),
            b"",
        ):

            digest.update(
                chunk
            )

    return digest.hexdigest()


def write_tsv(
    path: Path,
    rows: list[dict[str, object]],
) -> None:

    if not rows:

        die(
            f"No rows available for {path}"
        )

    path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with path.open(
        "w",
        encoding="utf-8",
        newline="",
    ) as handle:

        writer = csv.DictWriter(
            handle,
            fieldnames=list(
                rows[0].keys()
            ),
            delimiter="\t",
            lineterminator="\n",
        )

        writer.writeheader()
        writer.writerows(
            rows
        )


def read_tsv(
    path: Path,
) -> list[dict[str, str]]:

    with path.open(
        "r",
        encoding="utf-8",
        newline="",
    ) as handle:

        return list(
            csv.DictReader(
                handle,
                delimiter="\t",
            )
        )


def xml_bytes(
    root: etree._Element,
) -> bytes:

    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def normalize(
    value: str,
) -> str:

    return re.sub(
        r"\s+",
        " ",
        value.replace(
            "\u00a0",
            " ",
        ),
    ).strip()


def excerpt(
    value: str,
    limit: int = 180,
) -> str:

    value = normalize(
        value
    )

    if len(value) <= limit:
        return value

    return (
        value[
            : limit - 3
        ]
        + "..."
    )


def docx_body_signature(
    doc: Document,
) -> tuple[
    list[str],
    list[list[list[str]]],
]:

    paragraphs = [
        paragraph.text
        for paragraph
        in doc.paragraphs
    ]

    tables = [
        [
            [
                cell.text
                for cell
                in row.cells
            ]
            for row
            in table.rows
        ]
        for table
        in doc.tables
    ]

    return (
        paragraphs,
        tables,
    )


def section_orientations(
    doc: Document,
) -> list[str]:

    return [
        (
            "landscape"
            if section.page_width
            > section.page_height
            else "portrait"
        )
        for section
        in doc.sections
    ]


def write_docx_with_overrides(
    source_docx: Path,
    output_docx: Path,
    overrides: dict[str, bytes],
) -> None:

    output_docx.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with zipfile.ZipFile(
        source_docx,
        "r",
    ) as zin:

        with zipfile.ZipFile(
            output_docx,
            "w",
            zipfile.ZIP_DEFLATED,
        ) as zout:

            seen = set()

            for info in zin.infolist():

                name = (
                    info.filename
                )

                if name in overrides:

                    zout.writestr(
                        name,
                        overrides[
                            name
                        ],
                    )

                else:

                    zout.writestr(
                        name,
                        zin.read(
                            name
                        ),
                    )

                seen.add(
                    name
                )

            for name, payload in (
                overrides.items()
            ):

                if name not in seen:

                    zout.writestr(
                        name,
                        payload,
                    )


# ---------------------------------------------------------------------------
# Tracked-change helpers
# ---------------------------------------------------------------------------

def enable_track_revisions(
    settings_root: etree._Element,
) -> None:

    if settings_root.find(
        "w:trackRevisions",
        namespaces=NS,
    ) is None:

        settings_root.insert(
            0,
            etree.Element(
                w(
                    "trackRevisions"
                )
            ),
        )


def next_change_id(
    doc_root: etree._Element,
) -> int:

    ids = []

    for element in doc_root.xpath(
        ".//*[@w:id]",
        namespaces=NS,
    ):

        value = element.get(
            w(
                "id"
            )
        )

        try:

            ids.append(
                int(
                    value
                )
            )

        except (
            TypeError,
            ValueError,
        ):

            pass

    return (
        max(ids) + 1
        if ids
        else 1
    )


def revision_wrapper(
    tag: str,
    change_id: int,
) -> etree._Element:

    element = etree.Element(
        w(
            tag
        )
    )

    element.set(
        w(
            "id"
        ),
        str(
            change_id
        ),
    )

    element.set(
        w(
            "author"
        ),
        REVISION_AUTHOR,
    )

    element.set(
        w(
            "date"
        ),
        REVISION_DATE,
    )

    return element


def first_run_properties(
    paragraph: etree._Element,
) -> etree._Element | None:

    result = paragraph.xpath(
        ".//w:r/w:rPr",
        namespaces=NS,
    )

    if not result:
        return None

    return copy.deepcopy(
        result[0]
    )


def make_run(
    text: str,
    run_properties: etree._Element | None,
    deleted: bool = False,
) -> etree._Element:

    run = etree.Element(
        w(
            "r"
        )
    )

    if run_properties is not None:

        run.append(
            copy.deepcopy(
                run_properties
            )
        )

    text_tag = (
        "delText"
        if deleted
        else "t"
    )

    text_node = etree.SubElement(
        run,
        w(
            text_tag
        ),
    )

    if (
        text.startswith(
            " "
        )
        or text.endswith(
            " "
        )
        or "  " in text
    ):

        text_node.set(
            f"{{{XML_NS}}}space",
            "preserve",
        )

    text_node.text = (
        text
    )

    return run


TOKEN_PATTERN = re.compile(
    r"\s+|[^\s]+"
)


def token_chunks(
    value: str,
) -> list[str]:

    return TOKEN_PATTERN.findall(
        value
    )


def paragraph_has_complex_content(
    paragraph: etree._Element,
) -> bool:

    allowed = {
        w(
            "pPr"
        ),
        w(
            "r"
        ),
    }

    return any(
        child.tag
        not in allowed
        for child
        in paragraph
    )


def replace_paragraph_with_tracked_diff(
    paragraph: etree._Element,
    old_text: str,
    new_text: str,
    change_id: int,
) -> tuple[
    int,
    int,
    int,
]:

    run_properties = (
        first_run_properties(
            paragraph
        )
    )

    ppr = paragraph.find(
        "w:pPr",
        namespaces=NS,
    )

    for child in list(
        paragraph
    ):

        if child is ppr:
            continue

        paragraph.remove(
            child
        )

    old_tokens = token_chunks(
        old_text
    )

    new_tokens = token_chunks(
        new_text
    )

    matcher = SequenceMatcher(
        None,
        old_tokens,
        new_tokens,
        autojunk=False,
    )

    insertions = 0
    deletions = 0

    for (
        tag,
        i1,
        i2,
        j1,
        j2,
    ) in matcher.get_opcodes():

        old_chunk = "".join(
            old_tokens[
                i1:i2
            ]
        )

        new_chunk = "".join(
            new_tokens[
                j1:j2
            ]
        )

        if tag == "equal":

            if new_chunk:

                paragraph.append(
                    make_run(
                        new_chunk,
                        run_properties,
                        deleted=False,
                    )
                )

            continue

        if (
            tag
            in {
                "delete",
                "replace",
            }
            and old_chunk
        ):

            wrapper = (
                revision_wrapper(
                    "del",
                    change_id,
                )
            )

            change_id += 1

            wrapper.append(
                make_run(
                    old_chunk,
                    run_properties,
                    deleted=True,
                )
            )

            paragraph.append(
                wrapper
            )

            deletions += 1

        if (
            tag
            in {
                "insert",
                "replace",
            }
            and new_chunk
        ):

            wrapper = (
                revision_wrapper(
                    "ins",
                    change_id,
                )
            )

            change_id += 1

            wrapper.append(
                make_run(
                    new_chunk,
                    run_properties,
                    deleted=False,
                )
            )

            paragraph.append(
                wrapper
            )

            insertions += 1

    return (
        change_id,
        insertions,
        deletions,
    )


def wrap_runs_in_container_as_insertions(
    container: etree._Element,
    change_id: int,
) -> tuple[
    int,
    int,
]:

    count = 0

    for child in list(
        container
    ):

        if child.tag == w(
            "pPr"
        ):
            continue

        if child.tag == w(
            "r"
        ):

            wrapper = (
                revision_wrapper(
                    "ins",
                    change_id,
                )
            )

            change_id += 1

            index = (
                container.index(
                    child
                )
            )

            container.remove(
                child
            )

            wrapper.append(
                child
            )

            container.insert(
                index,
                wrapper,
            )

            count += 1

            continue

        (
            change_id,
            nested,
        ) = (
            wrap_runs_in_container_as_insertions(
                child,
                change_id,
            )
        )

        count += (
            nested
        )

    return (
        change_id,
        count,
    )


def mark_paragraph_as_insertion(
    paragraph: etree._Element,
    change_id: int,
) -> tuple[
    int,
    int,
]:

    return (
        wrap_runs_in_container_as_insertions(
            paragraph,
            change_id,
        )
    )


def unwrap_element(
    element: etree._Element,
) -> None:

    parent = (
        element.getparent()
    )

    if parent is None:
        return

    index = (
        parent.index(
            element
        )
    )

    children = list(
        element
    )

    for child in children:

        element.remove(
            child
        )

        parent.insert(
            index,
            child,
        )

        index += 1

    parent.remove(
        element
    )


def accept_all_revisions(
    doc_root: etree._Element,
) -> etree._Element:

    root = copy.deepcopy(
        doc_root
    )

    for element in list(
        root.xpath(
            ".//w:del",
            namespaces=NS,
        )
    ):

        parent = (
            element.getparent()
        )

        if parent is not None:

            parent.remove(
                element
            )

    for element in list(
        root.xpath(
            ".//w:ins",
            namespaces=NS,
        )
    ):

        unwrap_element(
            element
        )

    for element in list(
        root.xpath(
            ".//w:moveFrom",
            namespaces=NS,
        )
    ):

        parent = (
            element.getparent()
        )

        if parent is not None:

            parent.remove(
                element
            )

    for element in list(
        root.xpath(
            ".//w:moveTo",
            namespaces=NS,
        )
    ):

        unwrap_element(
            element
        )

    return root


def settings_without_track_revisions(
    settings_root: etree._Element,
) -> etree._Element:

    root = copy.deepcopy(
        settings_root
    )

    node = root.find(
        "w:trackRevisions",
        namespaces=NS,
    )

    if node is not None:

        root.remove(
            node
        )

    return root


# ---------------------------------------------------------------------------
# Preflight
# ---------------------------------------------------------------------------

WORK.mkdir(
    parents=True,
    exist_ok=True,
)

OUT.mkdir(
    parents=True,
    exist_ok=True,
)

REPORT.parent.mkdir(
    parents=True,
    exist_ok=True,
)


for required in [
    CLEAN_DOCX,
    BASELINE_DOCX,
    SOURCE_MD,
    ALIGNMENT_TSV,
    TARGET_TSV,
    ALIGNMENT_SUMMARY,
]:

    if not required.exists():

        die(
            f"Required input missing: "
            f"{required}"
        )


clean_sha = sha256(
    CLEAN_DOCX
)

baseline_sha = sha256(
    BASELINE_DOCX
)

source_sha = sha256(
    SOURCE_MD
)


if (
    clean_sha
    != EXPECTED_CLEAN_SHA
):

    die(
        "Clean DOCX SHA mismatch: "
        f"{clean_sha}"
    )


if (
    baseline_sha
    != EXPECTED_BASELINE_SHA
):

    die(
        "Baseline DOCX SHA mismatch: "
        f"{baseline_sha}"
    )


if (
    source_sha
    != EXPECTED_SOURCE_SHA
):

    die(
        "Scientific source SHA mismatch: "
        f"{source_sha}"
    )


summary_rows = read_tsv(
    ALIGNMENT_SUMMARY
)


if (
    len(
        summary_rows
    ) != 1
    or summary_rows[
        0
    ].get(
        "quality_gate"
    ) != "PASS"
):

    die(
        "Controlled alignment summary "
        "is not PASS."
    )


alignment_rows = read_tsv(
    ALIGNMENT_TSV
)

target_rows = read_tsv(
    TARGET_TSV
)


alignment_by_index = {
    int(
        row[
            "clean_paragraph_index"
        ]
    ): row
    for row
    in alignment_rows
}


response_ids_by_index: dict[
    int,
    set[str],
] = defaultdict(
    set
)


for row in target_rows:

    response_id = (
        row[
            "response_id"
        ]
    )

    hits = row.get(
        "clean_paragraph_hits",
        "",
    )

    if (
        not hits
        or hits == "NONE"
    ):

        continue

    for token in hits.split(
        ","
    ):

        token = (
            token.strip()
        )

        if token:

            response_ids_by_index[
                int(
                    token
                )
            ].add(
                response_id
            )


clean_doc = Document(
    CLEAN_DOCX
)

baseline_doc = Document(
    BASELINE_DOCX
)

clean_paragraphs = list(
    clean_doc.paragraphs
)

baseline_paragraphs = list(
    baseline_doc.paragraphs
)

clean_signature = (
    docx_body_signature(
        clean_doc
    )
)

clean_orientations = (
    section_orientations(
        clean_doc
    )
)


# ---------------------------------------------------------------------------
# Select redline paragraphs
# ---------------------------------------------------------------------------

selected_indices = set(
    response_ids_by_index
)


# Title and short title are explicit tracked replacements.

selected_indices.update(
    {
        1,
        2,
    }
)

response_ids_by_index[
    1
].add(
    "TITLE_REFRAME"
)

response_ids_by_index[
    2
].add(
    "TITLE_REFRAME"
)


# Objectives and reviewer-driven caption/legend revisions.

for index, paragraph in enumerate(
    clean_paragraphs,
    start=1,
):

    text = normalize(
        paragraph.text
    )

    if (
        text == "Objectives"
        or text.startswith(
            "To determine whether bacterial-"
        )
    ):

        selected_indices.add(
            index
        )

        response_ids_by_index[
            index
        ].add(
            "REVISION_OBJECTIVES"
        )

    if re.match(
        r"^(Fig [123]\. |S1 Fig\. |Table [12]\. )",
        text,
    ):

        selected_indices.add(
            index
        )

        response_ids_by_index[
            index
        ].add(
            "RR12_LEGENDS"
        )

    if re.match(
        r"^\([ABC]\) ",
        text,
    ):

        selected_indices.add(
            index
        )

        response_ids_by_index[
            index
        ].add(
            "RR12_LEGENDS"
        )

    if text.startswith(
        "Table 1 reports fixed-module external projection"
    ):

        selected_indices.add(
            index
        )

        response_ids_by_index[
            index
        ].add(
            "RR12_LEGENDS"
        )

    if text.startswith(
        "Effects are Hodges-Lehmann bacterial-minus-viral"
    ):

        selected_indices.add(
            index
        )

        response_ids_by_index[
            index
        ].add(
            "RR12_LEGENDS"
        )


# Exclude reference entries accidentally captured by broad reviewer patterns.

for index in list(
    selected_indices
):

    if (
        index < 1
        or index
        > len(
            clean_paragraphs
        )
    ):

        selected_indices.discard(
            index
        )

        continue

    text = normalize(
        clean_paragraphs[
            index - 1
        ].text
    )

    if re.match(
        r"^\[\d+\]\s",
        text,
    ):

        selected_indices.discard(
            index
        )

        response_ids_by_index.pop(
            index,
            None,
        )


if not selected_indices:

    die(
        "No paragraphs selected "
        "for targeted redline."
    )


# ---------------------------------------------------------------------------
# Patch the clean DOCX package
# ---------------------------------------------------------------------------

with zipfile.ZipFile(
    CLEAN_DOCX,
    "r",
) as zin:

    doc_root = etree.fromstring(
        zin.read(
            "word/document.xml"
        )
    )

    settings_root = (
        etree.fromstring(
            zin.read(
                "word/settings.xml"
            )
        )
    )


body_paragraphs = doc_root.xpath(
    "/w:document/w:body/w:p",
    namespaces=NS,
)


if (
    len(
        body_paragraphs
    )
    != len(
        clean_paragraphs
    )
):

    die(
        "Top-level paragraph count mismatch "
        "between OOXML and python-docx: "
        f"xml={len(body_paragraphs)} "
        f"python_docx={len(clean_paragraphs)}"
    )


enable_track_revisions(
    settings_root
)


change_id = next_change_id(
    doc_root
)


manifest_rows: list[
    dict[str, object]
] = []


total_ins = 0
total_del = 0


for index in sorted(
    selected_indices
):

    paragraph = (
        body_paragraphs[
            index - 1
        ]
    )

    new_text = (
        clean_paragraphs[
            index - 1
        ].text
    )

    alignment = (
        alignment_by_index.get(
            index
        )
    )

    classification = (
        alignment.get(
            "classification",
            "UNALIGNED",
        )
        if alignment
        else "UNALIGNED"
    )

    best_baseline_index = (
        int(
            alignment[
                "best_submitted_paragraph_index"
            ]
        )
        if (
            alignment
            and alignment.get(
                "best_submitted_paragraph_index"
            )
        )
        else 0
    )

    similarity = (
        float(
            alignment.get(
                "similarity_score",
                "0",
            )
            or 0
        )
        if alignment
        else 0.0
    )

    sequence_ratio = (
        float(
            alignment.get(
                "sequence_ratio",
                "0",
            )
            or 0
        )
        if alignment
        else 0.0
    )


    old_text = ""


    if (
        1
        <= best_baseline_index
        <= len(
            baseline_paragraphs
        )
    ):

        old_text = (
            baseline_paragraphs[
                best_baseline_index
                - 1
            ].text
        )


    action = (
        "TRACKED_INSERT_PARAGRAPH"
    )

    ins_count = 0
    del_count = 0


    if index == 1:

        old_text = (
            SUBMITTED_TITLE
        )

        (
            change_id,
            ins_count,
            del_count,
        ) = (
            replace_paragraph_with_tracked_diff(
                paragraph,
                old_text,
                new_text,
                change_id,
            )
        )

        action = (
            "TRACKED_REPLACE_TITLE"
        )


    elif index == 2:

        old_text = (
            SUBMITTED_SHORT_TITLE
        )

        (
            change_id,
            ins_count,
            del_count,
        ) = (
            replace_paragraph_with_tracked_diff(
                paragraph,
                old_text,
                new_text,
                change_id,
            )
        )

        action = (
            "TRACKED_REPLACE_SHORT_TITLE"
        )


    else:

        old_words = max(
            1,
            len(
                old_text.split()
            ),
        )

        new_words = max(
            1,
            len(
                new_text.split()
            ),
        )

        length_ratio = (
            new_words
            / old_words
        )

        can_word_diff = (
            classification
            in {
                "NEAR_UNCHANGED",
                "MODIFIED_CANDIDATE",
            }
            and similarity
            >= 0.60
            and sequence_ratio
            >= 0.58
            and 0.45
            <= length_ratio
            <= 2.40
            and not (
                paragraph_has_complex_content(
                    paragraph
                )
            )
            and bool(
                old_text.strip()
            )
        )


        if can_word_diff:

            (
                change_id,
                ins_count,
                del_count,
            ) = (
                replace_paragraph_with_tracked_diff(
                    paragraph,
                    old_text,
                    new_text,
                    change_id,
                )
            )

            action = (
                "TRACKED_WORD_LEVEL_REVISION"
            )


        else:

            (
                change_id,
                ins_count,
            ) = (
                mark_paragraph_as_insertion(
                    paragraph,
                    change_id,
                )
            )

            action = (
                "TRACKED_INSERT_PARAGRAPH"
            )


    total_ins += (
        ins_count
    )

    total_del += (
        del_count
    )


    manifest_rows.append(
        {
            "clean_paragraph_index": (
                index
            ),
            "response_ids": (
                "|".join(
                    sorted(
                        response_ids_by_index.get(
                            index,
                            {
                                "UNSPECIFIED"
                            },
                        )
                    )
                )
            ),
            "alignment_class": (
                classification
            ),
            "best_submitted_paragraph_index": (
                best_baseline_index
                or "NONE"
            ),
            "similarity_score": (
                f"{similarity:.6f}"
            ),
            "sequence_ratio": (
                f"{sequence_ratio:.6f}"
            ),
            "redline_action": (
                action
            ),
            "tracked_insert_wrappers": (
                ins_count
            ),
            "tracked_delete_wrappers": (
                del_count
            ),
            "clean_excerpt": (
                excerpt(
                    new_text
                )
            ),
            "submitted_excerpt": (
                excerpt(
                    old_text
                )
            ),
        }
    )


# ---------------------------------------------------------------------------
# Build marked-up package
# ---------------------------------------------------------------------------

overrides = {
    "word/document.xml": (
        xml_bytes(
            doc_root
        )
    ),
    "word/settings.xml": (
        xml_bytes(
            settings_root
        )
    ),
}


write_docx_with_overrides(
    CLEAN_DOCX,
    MARKED_DOCX,
    overrides,
)


marked_sha = sha256(
    MARKED_DOCX
)


# ---------------------------------------------------------------------------
# Acceptance QA
# ---------------------------------------------------------------------------

accepted_doc_root = (
    accept_all_revisions(
        doc_root
    )
)

accepted_settings = (
    settings_without_track_revisions(
        settings_root
    )
)


write_docx_with_overrides(
    CLEAN_DOCX,
    ACCEPTED_QA_DOCX,
    {
        "word/document.xml": (
            xml_bytes(
                accepted_doc_root
            )
        ),
        "word/settings.xml": (
            xml_bytes(
                accepted_settings
            )
        ),
    },
)


accepted_doc = Document(
    ACCEPTED_QA_DOCX
)


accepted_signature = (
    docx_body_signature(
        accepted_doc
    )
)


accepted_orientations = (
    section_orientations(
        accepted_doc
    )
)


# ---------------------------------------------------------------------------
# Structural audit of marked-up package
# ---------------------------------------------------------------------------

with zipfile.ZipFile(
    MARKED_DOCX,
    "r",
) as zin:

    marked_doc_root = (
        etree.fromstring(
            zin.read(
                "word/document.xml"
            )
        )
    )

    marked_settings_root = (
        etree.fromstring(
            zin.read(
                "word/settings.xml"
            )
        )
    )

    names = (
        zin.namelist()
    )


tracked_insertions = len(
    marked_doc_root.xpath(
        ".//w:ins",
        namespaces=NS,
    )
)


tracked_deletions = len(
    marked_doc_root.xpath(
        ".//w:del",
        namespaces=NS,
    )
)


track_revisions_enabled = (
    marked_settings_root.find(
        "w:trackRevisions",
        namespaces=NS,
    )
    is not None
)


comment_parts = [
    name
    for name
    in names
    if (
        "comment"
        in name.lower()
        or "people"
        in name.lower()
        or "person"
        in name.lower()
    )
]


media_parts = [
    name
    for name
    in names
    if name.startswith(
        "word/media/"
    )
]


authors = sorted(
    {
        element.get(
            w(
                "author"
            ),
            "",
        )
        for element
        in marked_doc_root.xpath(
            ".//w:ins | .//w:del",
            namespaces=NS,
        )
    }
)


# ---------------------------------------------------------------------------
# Smoke render
# ---------------------------------------------------------------------------

SMOKE_DIR.mkdir(
    parents=True,
    exist_ok=True,
)


PROFILE_DIR.mkdir(
    parents=True,
    exist_ok=True,
)


for old_pdf in (
    SMOKE_DIR.glob(
        "*.pdf"
    )
):

    old_pdf.unlink()


if shutil.which(
    "soffice"
) is None:

    die(
        "LibreOffice soffice not found."
    )


environment = (
    os.environ.copy()
)


environment[
    "HOME"
] = str(
    PROFILE_DIR.resolve()
)


render_result = subprocess.run(
    [
        "soffice",
        "--headless",
        (
            "-env:UserInstallation="
            f"file://{PROFILE_DIR.resolve()}"
        ),
        "--convert-to",
        "pdf",
        "--outdir",
        str(
            SMOKE_DIR.resolve()
        ),
        str(
            MARKED_DOCX.resolve()
        ),
    ],
    stdout=subprocess.PIPE,
    stderr=subprocess.STDOUT,
    text=True,
    env=environment,
)


smoke_pdf = (
    SMOKE_DIR
    / f"{MARKED_DOCX.stem}.pdf"
)


smoke_ok = (
    smoke_pdf.exists()
    and smoke_pdf.stat().st_size
    > 0
)


smoke_bytes = (
    smoke_pdf.stat().st_size
    if smoke_pdf.exists()
    else 0
)


# ---------------------------------------------------------------------------
# Quality gate
# ---------------------------------------------------------------------------

checks: list[
    dict[str, object]
] = []


def check(
    description: str,
    passed: bool,
    observed: object,
    expected: object,
) -> None:

    checks.append(
        {
            "check_id": (
                f"Q{len(checks) + 1:02d}"
            ),
            "check_description": (
                description
            ),
            "pass": (
                "TRUE"
                if passed
                else "FALSE"
            ),
            "observed": str(
                observed
            ),
            "expected": str(
                expected
            ),
        }
    )


check(
    "Clean DOCX SHA locked",
    clean_sha
    == EXPECTED_CLEAN_SHA,
    clean_sha,
    EXPECTED_CLEAN_SHA,
)


check(
    "Submitted baseline SHA locked",
    baseline_sha
    == EXPECTED_BASELINE_SHA,
    baseline_sha,
    EXPECTED_BASELINE_SHA,
)


check(
    "Scientific source SHA locked",
    source_sha
    == EXPECTED_SOURCE_SHA,
    source_sha,
    EXPECTED_SOURCE_SHA,
)


check(
    "Marked-up DOCX created",
    MARKED_DOCX.exists(),
    MARKED_DOCX.exists(),
    True,
)


check(
    "Marked-up DOCX differs from clean DOCX",
    marked_sha
    != clean_sha,
    marked_sha,
    "different from clean SHA",
)


check(
    "Track Changes enabled",
    track_revisions_enabled,
    track_revisions_enabled,
    True,
)


check(
    "Tracked insertions exist",
    tracked_insertions
    > 0,
    tracked_insertions,
    ">0",
)


check(
    "Tracked deletions exist",
    tracked_deletions
    > 0,
    tracked_deletions,
    ">0",
)


check(
    "Tracked-change author is Reuben S. Maghembe",
    authors
    == [
        REVISION_AUTHOR
    ],
    "|".join(
        authors
    ),
    REVISION_AUTHOR,
)


check(
    "At least 20 revision-target paragraphs are marked",
    len(
        manifest_rows
    )
    >= 20,
    len(
        manifest_rows
    ),
    ">=20",
)


check(
    "Title replacement is tracked",
    any(
        row[
            "redline_action"
        ]
        == "TRACKED_REPLACE_TITLE"
        for row
        in manifest_rows
    ),
    "present",
    "present",
)


check(
    "Short-title replacement is tracked",
    any(
        row[
            "redline_action"
        ]
        == "TRACKED_REPLACE_SHORT_TITLE"
        for row
        in manifest_rows
    ),
    "present",
    "present",
)


represented_targets = {
    response_id
    for row
    in manifest_rows
    for response_id
    in str(
        row[
            "response_ids"
        ]
    ).split(
        "|"
    )
    if (
        response_id.startswith(
            "RR"
        )
        or response_id.startswith(
            "EDITOR"
        )
    )
}


check(
    "Reviewer/editor targets represented",
    len(
        represented_targets
    )
    >= 10,
    len(
        represented_targets
    ),
    ">=10",
)


check(
    "No comment-related package parts",
    len(
        comment_parts
    )
    == 0,
    (
        "|".join(
            comment_parts
        )
        or "NONE"
    ),
    "NONE",
)


check(
    "No embedded manuscript figures",
    len(
        media_parts
    )
    == 0,
    len(
        media_parts
    ),
    0,
)


check(
    "Accepted redline reproduces clean paragraph/table content",
    accepted_signature
    == clean_signature,
    (
        "identical"
        if accepted_signature
        == clean_signature
        else "different"
    ),
    "identical",
)


check(
    "Accepted redline preserves section orientations",
    accepted_orientations
    == clean_orientations,
    "|".join(
        accepted_orientations
    ),
    "|".join(
        clean_orientations
    ),
)


check(
    "Accepted QA copy has two editable tables",
    len(
        accepted_doc.tables
    )
    == 2,
    len(
        accepted_doc.tables
    ),
    2,
)


check(
    "Accepted QA copy has five sections",
    len(
        accepted_doc.sections
    )
    == 5,
    len(
        accepted_doc.sections
    ),
    5,
)


clean_sha_after = sha256(
    CLEAN_DOCX
)


check(
    "Canonical clean DOCX remained immutable",
    clean_sha_after
    == EXPECTED_CLEAN_SHA,
    clean_sha_after,
    EXPECTED_CLEAN_SHA,
)


check(
    "LibreOffice smoke render succeeded",
    smoke_ok,
    (
        f"status="
        f"{render_result.returncode}; "
        f"bytes="
        f"{smoke_bytes}"
    ),
    "non-empty PDF",
)


passed = sum(
    row[
        "pass"
    ]
    == "TRUE"
    for row
    in checks
)


failed = (
    len(
        checks
    )
    - passed
)


quality_gate = (
    "PASS"
    if failed == 0
    else "FAIL"
)


final_status = (
    "READY_FOR_MARKED_UP_RENDER_AND_VISUAL_QA"
    if failed == 0
    else "TARGETED_OOXML_REDLINE_REQUIRES_CORRECTION"
)


# ---------------------------------------------------------------------------
# Outputs
# ---------------------------------------------------------------------------

write_tsv(
    MANIFEST,
    manifest_rows,
)


write_tsv(
    QUALITY_GATE,
    checks,
)


write_tsv(
    SUMMARY,
    [
        {
            "clean_docx_sha256": (
                clean_sha
            ),
            "baseline_docx_sha256": (
                baseline_sha
            ),
            "scientific_source_sha256": (
                source_sha
            ),
            "marked_up_docx_sha256": (
                marked_sha
            ),
            "selected_paragraphs": len(
                manifest_rows
            ),
            "tracked_insertions": (
                tracked_insertions
            ),
            "tracked_deletions": (
                tracked_deletions
            ),
            "word_level_revision_paragraphs": sum(
                row[
                    "redline_action"
                ]
                == "TRACKED_WORD_LEVEL_REVISION"
                for row
                in manifest_rows
            ),
            "inserted_paragraphs": sum(
                row[
                    "redline_action"
                ]
                == "TRACKED_INSERT_PARAGRAPH"
                for row
                in manifest_rows
            ),
            "comments": len(
                comment_parts
            ),
            "embedded_media_files": len(
                media_parts
            ),
            "accepted_content_matches_clean": (
                accepted_signature
                == clean_signature
            ),
            "smoke_render_pdf_bytes": (
                smoke_bytes
            ),
            "quality_checks": len(
                checks
            ),
            "quality_checks_passed": (
                passed
            ),
            "quality_checks_failed": (
                failed
            ),
            "quality_gate": (
                quality_gate
            ),
            "final_status": (
                final_status
            ),
        }
    ],
)


REPORT.write_text(
    "\n".join(
        [
            (
                "# PLOS ONE Targeted OOXML "
                "Marked-Up Manuscript"
            ),
            "",
            (
                "Manuscript: "
                "PONE-D-26-30583"
            ),
            "",
            "## Strategy",
            "",
            (
                "The submitted and revised manuscripts "
                "differ substantially in section order "
                "and paragraph structure. A blanket "
                "sequential redline would therefore "
                "confound reviewer-driven revisions "
                "with relocation and formatting changes."
            ),
            "",
            (
                "This marked-up manuscript starts from "
                "the locked clean revised DOCX and "
                "applies true OOXML tracked changes only "
                "to reviewer/editor-targeted textual "
                "revisions, the revised title/short title, "
                "revision objectives, and revised "
                "figure/table legends."
            ),
            "",
            (
                "Clean paragraphs classified as revision "
                "additions are marked as tracked insertions. "
                "Where a reliable submitted paragraph "
                "match exists, modified text is represented "
                "by word-level tracked deletions and "
                "insertions. Section relocation and "
                "landscape-table formatting are not "
                "represented as textual deletions."
            ),
            "",
            "## Locked sources",
            "",
            (
                f"- Clean DOCX SHA256: "
                f"`{clean_sha}`"
            ),
            (
                "- Reconstructed submitted baseline "
                f"SHA256: `{baseline_sha}`"
            ),
            (
                f"- Scientific source SHA256: "
                f"`{source_sha}`"
            ),
            "",
            "## Output",
            "",
            (
                f"- Marked-up DOCX SHA256: "
                f"`{marked_sha}`"
            ),
            (
                f"- Marked paragraphs: "
                f"{len(manifest_rows)}"
            ),
            (
                "- Tracked insertion wrappers: "
                f"{tracked_insertions}"
            ),
            (
                "- Tracked deletion wrappers: "
                f"{tracked_deletions}"
            ),
            "",
            "## Acceptance QA",
            "",
            (
                "An internal acceptance copy was generated "
                "by accepting all OOXML revisions. Its "
                "paragraph text, editable table content "
                "and section orientations were required "
                "to match the locked clean manuscript "
                "exactly."
            ),
            "",
            "## Quality gate",
            "",
            (
                f"- Checks passed: "
                f"{passed}/{len(checks)}"
            ),
            (
                f"- Quality gate: "
                f"`{quality_gate}`"
            ),
            (
                f"- Final status: "
                f"`{final_status}`"
            ),
            "",
        ]
    ),
    encoding="utf-8",
    newline="\n",
)


# ---------------------------------------------------------------------------
# Console
# ---------------------------------------------------------------------------

print(
    "===== PLOS ONE TARGETED OOXML REDLINE ====="
)

print(
    f"clean_docx_sha256\t"
    f"{clean_sha}"
)

print(
    f"baseline_docx_sha256\t"
    f"{baseline_sha}"
)

print(
    f"scientific_source_sha256\t"
    f"{source_sha}"
)

print(
    f"marked_up_docx_sha256\t"
    f"{marked_sha}"
)

print(
    f"selected_paragraphs\t"
    f"{len(manifest_rows)}"
)

print(
    f"tracked_insertions\t"
    f"{tracked_insertions}"
)

print(
    f"tracked_deletions\t"
    f"{tracked_deletions}"
)

print(
    "accepted_content_matches_clean\t"
    f"{accepted_signature == clean_signature}"
)

print(
    f"smoke_render_pdf_bytes\t"
    f"{smoke_bytes}"
)

print(
    f"quality_checks_passed\t"
    f"{passed}/{len(checks)}"
)

print(
    f"quality_gate\t"
    f"{quality_gate}"
)

print(
    f"final_status\t"
    f"{final_status}"
)

print(
    f"marked_up_docx\t"
    f"{MARKED_DOCX}"
)

print(
    f"manifest\t"
    f"{MANIFEST}"
)

print(
    f"summary\t"
    f"{SUMMARY}"
)

print(
    f"quality_gate_file\t"
    f"{QUALITY_GATE}"
)

print(
    f"report\t"
    f"{REPORT}"
)

print()

print(
    "===== LIBREOFFICE SMOKE OUTPUT ====="
)

print(
    render_result.stdout.strip()
)


if failed:

    die(
        "Targeted OOXML redline failed "
        f"{failed} quality check(s)."
    )
