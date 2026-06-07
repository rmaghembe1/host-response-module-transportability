#!/usr/bin/env python3

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

src = Path("docs/cmi_cover_letter_final_single_author.md")
out = Path("submission/cmi_cover_letter_single_author.docx")
out.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

for style_name in ["Normal", "Heading 1"]:
    style = doc.styles[style_name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

text = src.read_text()

for line in text.splitlines():
    s = line.strip()
    if not s:
        doc.add_paragraph("")
    elif s.startswith("# "):
        p = doc.add_paragraph()
        r = p.add_run(s[2:])
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
    else:
        p = doc.add_paragraph()
        parts = s.split("**")
        bold = False
        for part in parts:
            r = p.add_run(part)
            r.bold = bold
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            bold = not bold

doc.save(out)
print(f"Wrote: {out}")
